# -*- coding: utf-8 -*-
"""Offline smoke tests for server-side export safety helpers."""
from __future__ import annotations

import json
import io
import re
import sys
import tempfile
import unittest
import zipfile
from pathlib import Path

from docx import Document

import generate_docx
import server
from audit_schema import REQUIRED_AUDIT_PATHS, nested

PAGE_ONLY_REPORT_STATUS_NEEDLES = (
    '產報前檢查',
    '附件適用狀態',
    '優先建議報告閱讀狀態',
    '優先閱讀',
    '報告閱讀狀態',
    '可作附件',
    '暫勿作附件',
    '頁面輔助',
    '公司內部整理計算附件',
    '不會寫入計算書',
    '不會寫入計算書或列印 PDF',
)


def fake_handler(*, method: str = 'GET', path: str = '/status', origin: str | None = None):
    handler = object.__new__(server.Handler)
    handler.command = method
    handler.path = path
    handler.headers = {}
    if origin is not None:
        handler.headers['Origin'] = origin
    return handler


def fake_json_reader(body: bytes, content_type: str = 'application/json'):
    handler = fake_handler(method='POST', path='/generate')
    handler.headers = {
        'Content-Type': content_type,
        'Content-Length': str(len(body)),
    }
    handler.rfile = io.BytesIO(body)
    return handler


def sample_case() -> dict:
    return {
        'name': '標準板',
        'type': 'pk_4h',
        'w': 870,
        'h': 800,
    }


def sample_result() -> dict:
    return {
        'allOK': False,
        'A': 0.6969,
        'G': 56.376,
        'Wp': 313.2,
        'Fph': 62.64,
        'PE': 43.64,
        'PEV': 6.76,
        'PW': 313.2,
        'FW': 313.2,
        'S': 2.1,
        'P': 313.2,
        'T': 156.6,
        'V': 156.6,
        'mc_h1': 0,
        'mc_h2': 0,
        'checks': [
            {'item': '插銷剪力', 'formula': 'V <= Va', 'calc': '156.6 <= 85', 'allow': '85', 'v': 156.6, 'a': 85, 'sense': 'max', 'dcr': 156.6 / 85, 'pass': False},
            {'item': '孔位提醒', 'formula': 'e >= emin', 'calc': '60 >= 25', 'allow': '25', 'v': 60, 'a': 25, 'sense': 'min', 'dcr': 25 / 60, 'status': 'warning', 'pass': False},
            {'item': '厚度限制', 'formula': 't >= tmin', 'calc': '30 >= 25', 'allow': '25', 'v': 30, 'a': 25, 'sense': 'min', 'dcr': 25 / 30, 'pass': True},
        ],
        'design': {'tensionPerPoint': 156.6},
    }


def v2_payload(include_results: bool = True) -> dict:
    payload = {
        'schema': 'stone-calc/v3',
        'version': 3,
        'inp': {'proj': 'Smoke Test'},
        'cases': [sample_case()],
        'meta': {
            'app_version': f'V{server.SERVER_VERSION}',
            'tool_html': server.TOOL_HTML,
            'calculator_build': '2026.04.23-spec-core11',
            'generator': 'stone-spec-direct-word-html-v2',
            'workflow_mode': 'full',
            'template': {
                'source': 'builtin',
                'key': 'pk_4h_t30',
                'version': 'templates-2026.04.25',
            },
            'review_summary': {
                'delivery_quality': {
                    'code': 'C',
                    'text': 'C 先修正',
                    'reasons': ['1 項檢核未通過'],
                },
                'counts': {
                    'cases': 1,
                    'checks': 3,
                    'failed': 1,
                    'warnings': 1,
                    'review_notes': 1,
                },
                'formula_source': {
                    'version': 'formula-registry-2026.04.25',
                    'total': 3,
                    'covered': 3,
                    'missing': 0,
                },
            },
        },
        'reviewNotes': {'1::插銷剪力': '改採加大插銷後覆核。'},
    }
    if include_results:
        payload['results'] = [sample_result()]
    return payload


def report_text(value) -> str:
    return json.dumps(value, ensure_ascii=False, sort_keys=True)


class ServerSmokeTests(unittest.TestCase):
    def assertAuditSchema(self, report: dict) -> None:
        for path in REQUIRED_AUDIT_PATHS:
            sentinel = object()
            value = nested(report, *path, default=sentinel)
            self.assertIsNot(value, sentinel, f"missing audit path: {'.'.join(path)}")
        self.assertEqual(report['report_type'], 'stone_calc_export_audit')
        self.assertEqual(report['summary']['dcr_method'], 'normalized_by_check_sense')
        self.assertIsInstance(report['review']['server_evaluation']['tool_html_match'], bool)
        self.assertGreater(len(str(report['trace']['payload_sha256'])), 16)
        self.assertGreaterEqual(report['output']['size_bytes'], 0)
        text = report_text(report)
        for needle in PAGE_ONLY_REPORT_STATUS_NEEDLES:
            self.assertNotIn(needle, text)

    def test_cors_allows_only_local_http_and_status_file_origin(self) -> None:
        self.assertTrue(fake_handler(origin='http://127.0.0.1:8765')._is_allowed_cors_origin())
        self.assertTrue(fake_handler(origin='http://localhost:8765')._is_allowed_cors_origin())
        self.assertTrue(fake_handler(method='GET', path='/status', origin='null')._is_allowed_cors_origin())

        self.assertFalse(fake_handler(method='POST', path='/generate', origin='null')._is_allowed_cors_origin())
        self.assertFalse(fake_handler(origin='https://example.com')._is_allowed_cors_origin())
        self.assertFalse(fake_handler(origin='http://127.0.0.1:9999')._is_allowed_cors_origin())

    def test_static_scope_blocks_parent_directory_targets(self) -> None:
        inside = server.TOOL_DIR / server.TOOL_HTML
        sibling = server.ROOT_DIR / 'outside.html'
        parent = server.TOOL_DIR.parent

        handler = fake_handler()
        self.assertTrue(handler._is_static_target_allowed(inside.resolve()))
        self.assertFalse(handler._is_static_target_allowed(sibling.resolve()))
        self.assertFalse(handler._is_static_target_allowed(parent.resolve()))

    def test_json_payload_reader_accepts_only_utf8_json_objects(self) -> None:
        handler = fake_json_reader(json.dumps({'ok': True}).encode('utf-8'))
        self.assertEqual(handler._read_json_payload(), {'ok': True})

        with self.assertRaisesRegex(server.PayloadError, 'application/json'):
            fake_json_reader(b'{}', content_type='text/plain')._read_json_payload()

        with self.assertRaisesRegex(server.PayloadError, 'JSON 根節點必須是物件'):
            fake_json_reader(b'[]')._read_json_payload()

        with self.assertRaisesRegex(server.PayloadError, 'JSON 格式錯誤'):
            fake_json_reader(b'{bad json')._read_json_payload()

    def test_v2_generate_fallback_is_blocked_without_results(self) -> None:
        with self.assertRaisesRegex(ValueError, 'V2 專案缺少有效的 results'):
            generate_docx.resolve_results(v2_payload(include_results=False), {}, [sample_case()])

    def test_frontend_result_source_is_inferred_for_matching_cases(self) -> None:
        payload = v2_payload()
        self.assertEqual(server.infer_result_source(payload, 'generate_legacy_compat'), 'frontend_results')
        results, source = generate_docx.resolve_results(payload, payload['inp'], payload['cases'])
        self.assertEqual(source, 'frontend_results')
        self.assertEqual(len(results), 1)

    def test_export_audit_records_summary_and_trace(self) -> None:
        payload = v2_payload()
        with tempfile.TemporaryDirectory() as tmp:
            out_path = Path(tmp) / 'smoke.docx'
            out_path.write_bytes(b'smoke-docx')
            audit_path = server.write_export_audit(
                payload,
                out_path,
                'generate_legacy_compat',
                result_source='frontend_results',
            )
            report = json.loads(audit_path.read_text(encoding='utf-8'))

        self.assertAuditSchema(report)
        self.assertTrue(report['ok'])
        self.assertEqual(report['tool']['tool_html'], server.TOOL_HTML)
        self.assertEqual(report['tool']['template']['key'], 'pk_4h_t30')
        self.assertEqual(report['trace']['result_source'], 'frontend_results')
        self.assertEqual(report['summary']['case_count'], 1)
        self.assertEqual(report['summary']['result_count'], 1)
        self.assertEqual(report['summary']['checks_total'], 3)
        self.assertEqual(report['summary']['checks_failed'], 1)
        self.assertEqual(report['summary']['warnings'], 1)
        self.assertEqual(report['summary']['review_notes'], 1)
        self.assertEqual(report['summary']['worst']['item'], '插銷剪力')
        self.assertAlmostEqual(report['summary']['worst']['dcr'], 156.6 / 85, places=6)
        self.assertEqual(report['summary']['worst_warning']['item'], '孔位提醒')
        self.assertEqual(report['summary']['dcr_method'], 'normalized_by_check_sense')
        self.assertEqual(report['summary']['dcr_sense_counts']['min'], 2)
        self.assertEqual(report['summary']['cases'][0]['failed_checks'], 1)
        self.assertEqual(report['review']['delivery_quality']['code'], 'C')
        self.assertEqual(report['review']['formula_source']['missing'], 0)
        self.assertFalse(report['review']['server_status_not_evaluated'])
        self.assertTrue(report['review']['server_evaluation']['tool_html_match'])

    def test_legacy_generate_report_outputs_structured_docx_and_audit_boundary(self) -> None:
        payload = v2_payload()
        review_summary = payload['meta']['review_summary']
        review_summary['priority_report_reading_status'] = {
            'title': '優先建議報告閱讀狀態',
            'decision': '暫勿作附件',
            'note': '頁面輔助，不會寫入計算書或列印 PDF。',
        }
        review_summary['delivery_quality']['reasons'].append('報告閱讀狀態：暫勿作附件')

        with tempfile.TemporaryDirectory() as tmp:
            payload_path = Path(tmp) / 'stone_payload.json'
            payload_path.write_text(json.dumps(payload, ensure_ascii=False), encoding='utf-8')

            out_path = Path(generate_docx.generate_report(str(payload_path)))
            self.assertEqual(out_path.suffix.lower(), '.docx')
            self.assertTrue(out_path.is_file())
            self.assertGreater(out_path.stat().st_size, 35000)

            doc = Document(str(out_path))
            paragraph_text = '\n'.join(para.text for para in doc.paragraphs)
            table_text = '\n'.join(
                cell.text
                for table in doc.tables
                for row in table.rows
                for cell in row.cells
            )
            combined_text = f'{paragraph_text}\n{table_text}'

            with zipfile.ZipFile(out_path) as archive:
                document_xml = archive.read('word/document.xml').decode('utf-8')

            paragraph_count = len(re.findall(r'<w:p(?:\s|>)', document_xml))
            table_count = len(re.findall(r'<w:tbl(?:\s|>)', document_xml))
            page_break_count = len(re.findall(r'<w:br[^>]+w:type="page"', document_xml))
            section_count = len(re.findall(r'[一二三四五六七八九十]+、', combined_text))

            self.assertGreaterEqual(len(combined_text), 2000)
            self.assertGreaterEqual(len(doc.paragraphs), 55)
            self.assertGreaterEqual(len(doc.tables), 4)
            self.assertGreaterEqual(paragraph_count, 110)
            self.assertGreaterEqual(table_count, 4)
            self.assertGreaterEqual(page_break_count, 2)
            self.assertGreaterEqual(section_count, 10)

            for needle in (
                '結　構　計　算　書',
                '石材外牆固定構件結構檢核',
                '一、設計基準',
                '五、各案例檢核彙整',
                '附件 1',
                '插銷剪力',
                '改採加大插銷後覆核。',
            ):
                self.assertIn(needle, combined_text)

            for needle in PAGE_ONLY_REPORT_STATUS_NEEDLES:
                self.assertNotIn(needle, combined_text)
                self.assertNotIn(needle, document_xml)

            audit_path = server.write_export_audit(
                payload,
                out_path,
                'generate_legacy_compat',
                result_source='frontend_results',
            )
            report = json.loads(audit_path.read_text(encoding='utf-8'))

            self.assertAuditSchema(report)
            self.assertEqual(report['output']['size_bytes'], out_path.stat().st_size)
            self.assertEqual(report['trace']['result_source'], 'frontend_results')
            self.assertEqual(report['review']['delivery_quality']['code'], 'C')
            self.assertEqual(report['summary']['checks_total'], 3)

    def test_export_audit_strips_page_only_report_reading_status(self) -> None:
        payload = v2_payload()
        review_summary = payload['meta']['review_summary']
        review_summary['priority_report_reading_status'] = {
            'title': '優先建議報告閱讀狀態',
            'decision': '暫勿作附件',
            'note': '頁面輔助，不會寫入計算書或列印 PDF。',
        }
        review_summary['delivery_quality']['reasons'].append('報告閱讀狀態：暫勿作附件')

        with tempfile.TemporaryDirectory() as tmp:
            out_path = Path(tmp) / 'smoke.docx'
            out_path.write_bytes(b'smoke-docx')
            audit_path = server.write_export_audit(payload, out_path, 'auto_word', result_source='frontend_results')
            report = json.loads(audit_path.read_text(encoding='utf-8'))

        self.assertAuditSchema(report)
        self.assertNotIn('priority_report_reading_status', report['review'])
        self.assertEqual(report['review']['delivery_quality']['code'], 'C')
        self.assertEqual(report['review']['formula_source']['missing'], 0)

    def test_export_audit_forces_quality_c_when_payload_html_differs(self) -> None:
        payload = v2_payload()
        payload['meta']['tool_html'] = '石材計算書產生器_規範版.html'
        payload['meta']['review_summary']['delivery_quality'] = {
            'code': 'A',
            'text': 'A 可交付',
            'class': 'ok',
            'reasons': [],
        }
        with tempfile.TemporaryDirectory() as tmp:
            out_path = Path(tmp) / 'smoke.docx'
            out_path.write_bytes(b'smoke-docx')
            audit_path = server.write_export_audit(payload, out_path, 'auto_word', result_source='frontend_results')
            report = json.loads(audit_path.read_text(encoding='utf-8'))

        self.assertAuditSchema(report)
        self.assertFalse(report['review']['server_status_not_evaluated'])
        self.assertFalse(report['review']['server_evaluation']['tool_html_match'])
        self.assertEqual(report['review']['delivery_quality']['code'], 'C')


if __name__ == '__main__':
    runner = unittest.TextTestRunner(stream=sys.stdout, verbosity=2)
    unittest.main(testRunner=runner)
