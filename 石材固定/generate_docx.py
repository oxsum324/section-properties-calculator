# -*- coding: utf-8 -*-
"""
╔════════════════════════════════════════════════════════════════════════╗
║  ⚠ LEGACY — DO NOT USE FOR FORMAL DELIVERABLES  ⚠                       ║
║                                                                          ║
║  此檔保留作為歷史對照用途。它包含獨立實作的計算邏輯（如 calc_case），    ║
║  與 V2 主流程使用的 js/calculator.spec.js (StoneCalculator) 不同步。      ║
║                                                                          ║
║  已知差異：                                                              ║
║    - 此檔 Fph 用簡化式 1.6 × SDS × Ip × Wp                               ║
║    - V2 對啟用耐震細算（C6 模組）案例改用                                ║
║      0.4 × ap × SDS × Ip × (1+2z/h) × Wp / Rp（受上下限控制）             ║
║    → 同案例兩管道輸出可能不同 Fph，違反「單一可信計算資料鏈」治理原則    ║
║                                                                          ║
║  正式 Word 匯出請使用：                                                  ║
║    auto_word.py  (HTML 主流程 → 列印 PDF → docx；不重算)                  ║
║                                                                          ║
║  本檔將於 calc-core 完全分離後正式 deprecate。維護者勿在此處修改公式。   ║
╚════════════════════════════════════════════════════════════════════════╝

石材外牆固定構件計算書產生器（legacy 路徑）
讀取 stone_report*.json，重算結構檢核，輸出格式化 .docx 並自動開啟。
"""

import warnings as _warnings
_warnings.warn(
    "generate_docx.py is LEGACY — use auto_word.py for formal Word output. "
    "See file header for details.",
    DeprecationWarning,
    stacklevel=2,
)

import sys, os, json, math, re
from pathlib import Path
from copy import deepcopy
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
import lxml.etree as etree

# ─────────────────────────────────────────────
#  類型對照表
# ─────────────────────────────────────────────
TYPE_N = {
    'bk_2': 2, 'bk_4h': 4, 'bk_6h': 6, 'bk_6v': 6,
    'pk_4h': 4, 'pk_6h': 6, 'pk_4v': 4, 'pk_6v': 6,
}

# 套管式膨脹錨栓廠商極限強度表（新生五金 SIOC 套管式膨脹錨栓，資料來源：新生五金工廠股份有限公司型錄）
ANCHOR_CATALOG = {
    'SH-440': { 'd': '1/2"',  'embedMm': 50.8, 'Nu': {210: 2800, 280: 3700}, 'Vu': {210: 3000, 280: 3000} },
    'SH-550': { 'd': '5/8"',  'embedMm': 63.5, 'Nu': {210: 3850, 280: 5200}, 'Vu': {210: 4600, 280: 5100} },
    'SH-560': { 'd': '5/8"',  'embedMm': 63.5, 'Nu': {210: 3850, 280: 5200}, 'Vu': {210: 4600, 280: 5100} },
    'SH-660': { 'd': '3/4"',  'embedMm': 76.2, 'Nu': {210: 5700, 280: 5700}, 'Vu': {210: 7200, 280: 7200} },
}

def anc_ta_detail(inp):
    """回傳內迫螺栓容許拉力的查表依據字串，供計算書主文引用。"""
    Ta = float(inp.get('m_anc_ta', 499.7))
    t  = inp.get('m_anc_type', 'custom')
    fc = int(float(inp.get('m_anc_fc', 280)))
    sf = float(inp.get('m_anc_sf', 5))
    spec = ANCHOR_CATALOG.get(t)
    if not spec or t == 'custom' or fc not in (spec['Nu'] if spec else {}):
        return {
            'short': f'{Ta:.1f} kgf（自訂值）',
            'full':  f'Ta = {Ta:.1f} kgf（自訂輸入，未套用廠商查表）',
            'source': '自訂',
        }
    Nu = spec['Nu'][fc]
    Vu = spec['Vu'][fc]
    return {
        'short': f'{Ta:.1f} kgf（{t}）',
        'full':  f'Ta = Nu ÷ SF = {Nu} ÷ {sf:g} = {Ta:.1f} kgf',
        'source': (f'查表 {t}（{spec["d"]}），埋入深度 ≥ {spec["embedMm"]} mm，錨入 fc\' = {fc} kgf/cm² 混凝土；'
                   f'極限抗拉 Nu = {Nu} kgf、極限抗剪 Vu = {Vu} kgf'
                   f'（資料來源：新生五金 SIOC 套管式膨脹錨栓型錄）；'
                   f'安全係數 SF = {sf:g}'),
    }
TYPE_LABEL = {
    'bk_2':  '背扣孔─上下各1點',
    'bk_4h': '背扣孔─上下各2點',
    'bk_6h': '背扣孔─上下各3點',
    'bk_6v': '背扣孔─左右各3點',
    'pk_4h': '插銷式─上下各2支',
    'pk_6h': '插銷式─上下各3支',
    'pk_4v': '插銷式─左右各2支',
    'pk_6v': '插銷式─左右各3支',
}
DEFAULT_SPEC_REFS = {
    'cc': '建築物耐風設計規範（107 年版）',
    'seismic': '建築物耐震設計規範及解說（113 年版）',
    'anchor': '錨栓製造商試驗報告（安全係數 SF=3）',
    'steel': '鋼構造建築物鋼結構設計技術規範（容許應力設計法）',
}

# ─────────────────────────────────────────────
#  結構計算核心
# ─────────────────────────────────────────────
def calc_case(cd, inp):
    w   = float(cd.get('w', 870))
    h   = float(cd.get('h', 800))
    N   = int(cd.get('N', 4))
    bh  = float(cd.get('bh', 10))
    d1  = float(cd.get('d1', 8))
    Lt  = float(cd.get('Lt', 0.5))
    LL  = float(cd.get('LL', 5))
    d0  = float(cd.get('d0', 1.2))

    h12_raw = cd.get('h12', '3.4,3.6')
    h12 = [float(x) for x in str(h12_raw).split(',')]
    mc_h1 = h12[0] if len(h12) > 0 else 3.4
    mc_h2 = h12[1] if len(h12) > 1 else 3.6

    Wp  = float(inp.get('st_gam', 2800)) * float(inp.get('st_t', 30)) / 1000  # kgf/m²
    A   = (w / 1000) * (h / 1000)                                               # m²
    G   = A * Wp                                                                 # kgf

    s_sds = float(inp.get('s_sds', 0.6))
    s_ip  = float(inp.get('s_ip', 1.0))
    Fph   = 1.6 * s_sds * s_ip * Wp                                             # kgf/m²
    PE    = Fph * A                                                              # kgf (地震水平力)
    PEV   = 0.5 * Fph * A                                                        # kgf (地震垂直力)

    w_pos = float(inp.get('w_pos', 426))
    w_neg = float(inp.get('w_neg', 341))
    w_cf  = float(inp.get('w_cf', 1.25))
    FW    = max(w_pos, w_neg) * w_cf                                             # kgf/m² (設計風壓，取正/負較大值)
    PW    = A * FW                                                               # kgf (設計風力)

    S = G + PEV                                                                  # kgf (垂直設計力)
    P = max(PE, PW)                                                              # kgf (水平設計力，控制)

    T   = P / N                                                                  # kgf/個 (每點水平力)
    V   = S / N                                                                  # kgf/個 (每點垂直力)
    Tu1 = V * bh / d1 if d1 else 0                                              # 垂直力引起拉力
    Tu2 = T * bh / d1 if d1 else 0                                              # 水平力引起拉力
    Tu  = max(Tu1, Tu2)                                                          # 控制拉力

    T1  = V * mc_h1 / mc_h2 if mc_h2 else 0                                     # 馬車螺栓拉力

    Au    = (LL - d0) * Lt                                                       # cm² 角鋼淨截面積
    S_sec = (1 / 6) * LL * Lt ** 2                                               # cm³ 角鋼斷面模數

    m_fy  = float(inp.get('m_fy', 2100))
    Va_L  = 0.4 * m_fy * Au                                                     # kgf 角鋼允許剪力

    M    = max(V * bh, T * d1)                                                   # kgf·cm 彎矩
    Fb   = 0.6 * m_fy * w_cf                                                    # kgf/cm² 允許彎應力
    Sreq = M / Fb if Fb else 0                                                   # cm³ 需要斷面模數

    pin_d  = float(inp.get('m_pin_d', 5))
    pin_fy = float(inp.get('m_pin_fy', 2100))
    pin_A  = math.pi * (pin_d / 10) ** 2 / 4                                    # cm²
    Va_pin = 1.25 * 0.4 * pin_fy * pin_A                                        # kgf

    has_mc = bool(cd.get('hasMC', True))

    m_screw_ta = float(inp.get('m_screw_ta', 255.3))
    m_anc_ta   = float(inp.get('m_anc_ta', 499.7))
    m_mc_va    = float(inp.get('m_mc_va', 567))
    m_mc_ta    = float(inp.get('m_mc_ta', 709))

    checks = [
        {'no': '①', 'item': '背扣螺絲 抗拔',
         'v': T, 'a': m_screw_ta, 'unit': 'kgf', 'pass': T <= m_screw_ta,
         'formula': f'T = P÷N = {P:.2f}÷{N} = {T:.2f} kgf'},
        {'no': '②', 'item': '膨脹螺栓 Tu1（垂直力）',
         'v': Tu1, 'a': m_anc_ta, 'unit': 'kgf', 'pass': Tu1 <= m_anc_ta,
         'formula': f'Tu1 = V×h÷d₁ = {V:.2f}×{bh}÷{d1} = {Tu1:.2f} kgf'},
        {'no': '③', 'item': '膨脹螺栓 Tu2（水平力）',
         'v': Tu2, 'a': m_anc_ta, 'unit': 'kgf', 'pass': Tu2 <= m_anc_ta,
         'formula': f'Tu2 = T×h÷d₁ = {T:.2f}×{bh}÷{d1} = {Tu2:.2f} kgf'},
    ]

    if has_mc:
        checks += [
            {'no': '④', 'item': '馬車螺栓 剪力',
             'v': T, 'a': m_mc_va, 'unit': 'kgf', 'pass': T <= m_mc_va,
             'formula': f'V = T = {T:.2f} kgf'},
            {'no': '⑤', 'item': '馬車螺栓 拉力 T₁',
             'v': T1, 'a': m_mc_ta, 'unit': 'kgf', 'pass': T1 <= m_mc_ta,
             'formula': f'T₁ = V×h₁÷h₂ = {V:.2f}×{mc_h1}÷{mc_h2} = {T1:.2f} kgf'},
        ]
        an = ('⑥', '⑦')
    else:
        checks.append(
            {'no': '④', 'item': '插銷 剪力',
             'v': Tu, 'a': Va_pin, 'unit': 'kgf', 'pass': Tu <= Va_pin,
             'formula': f'Tu = {Tu:.2f} kgf ≤ Va = {Va_pin:.1f} kgf'}
        )
        an = ('⑤', '⑥')

    checks += [
        {'no': an[0], 'item': '角鋼 剪力（淨截面）',
         'v': V, 'a': Va_L, 'unit': 'kgf', 'pass': V <= Va_L,
         'formula': f'V = {V:.2f} kgf，Va = 0.4×Fy×Au = {Va_L:.2f} kgf'},
        {'no': an[1], 'item': '角鋼 彎矩（斷面模數）',
         'v': Sreq, 'a': S_sec, 'unit': 'cm³', 'pass': Sreq <= S_sec,
         'formula': f'Sreq = M÷Fb = {M:.1f}÷{Fb:.1f} = {Sreq:.4f} cm³'},
    ]

    return {
        'A': A, 'G': G, 'Wp': Wp, 'Fph': Fph,
        'PE': PE, 'PEV': PEV, 'PW': PW, 'FW': FW,
        'S': S, 'P': P, 'T': T, 'V': V, 'gov_wind': PW >= PE,
        'Tu1': Tu1, 'Tu2': Tu2, 'Tu': Tu, 'T1': T1,
        'Au': Au, 'S_sec': S_sec, 'Va_L': Va_L, 'M': M, 'Fb': Fb, 'Sreq': Sreq,
        'Va_pin': Va_pin, 'mc_h1': mc_h1, 'mc_h2': mc_h2,
        'checks': checks, 'all_ok': all(c['pass'] for c in checks),
        'N': N, 'bh': bh, 'd1': d1, 'Lt': Lt, 'LL': LL, 'd0': d0,
        'pin_d': pin_d, 'pin_fy': pin_fy, 'pin_A': pin_A, 'Va_pin': Va_pin,
        'has_mc': has_mc, 'mc_h1': mc_h1, 'mc_h2': mc_h2,
    }


RESULT_KEYS = (
    'A', 'G', 'Wp', 'Fph', 'PE', 'PEV', 'PW', 'FW',
    'S', 'P', 'T', 'V', 'mc_h1', 'mc_h2',
)


def normalize_result(raw, cd, inp):
    if not isinstance(raw, dict):
        return None

    checks = raw.get('checks')
    if not isinstance(checks, list) or not checks:
        return None

    normalized = {}
    for key in RESULT_KEYS:
        value = raw.get(key)
        if value is None:
            return None
        try:
            normalized[key] = float(value)
        except (TypeError, ValueError):
            return None

    normalized['gov_wind'] = bool(raw.get('gov_wind', raw.get('govWind', False)))
    normalized['all_ok'] = bool(raw.get('all_ok', raw.get('allOK', False)))

    normalized_checks = []
    for idx, check in enumerate(checks, start=1):
        if not isinstance(check, dict):
            return None
        formula = check.get('formula')
        calc = check.get('calc')
        allow = check.get('allow')
        if formula is None or calc is None or allow is None:
            return None
        normalized_checks.append({
            'no': check.get('no', str(idx)),
            'item': check.get('item', f'檢核項目 {idx}'),
            'formula': str(formula),
            'calc': str(calc),
            'allow': str(allow),
            'v': float(check['v']) if check.get('v') is not None else None,
            'a': float(check['a']) if check.get('a') is not None else None,
            'unit': check.get('unit', ''),
            'pass': bool(check.get('pass', False)),
        })

    normalized['checks'] = normalized_checks

    for key in ('Tu1', 'Tu2', 'Tu', 'T1', 'Au', 'S_sec', 'Va_L', 'M', 'Fb', 'Sreq',
                'N', 'bh', 'd1', 'Lt', 'LL', 'd0', 'pin_d', 'pin_fy', 'pin_A', 'Va_pin'):
        value = raw.get(key)
        if value is not None:
            try:
                normalized[key] = float(value)
            except (TypeError, ValueError):
                pass
    normalized['has_mc'] = bool(raw.get('has_mc', raw.get('hasMC', cd.get('hasMC', True))))
    return normalized


def is_v2_payload(data):
    """V2 專案必須使用前端共用計算核心輸出的 results，避免靜默退回舊公式。"""
    if not isinstance(data, dict):
        return False
    meta = data.get('meta') if isinstance(data.get('meta'), dict) else {}
    schema = str(data.get('schema', ''))
    app_version = str(meta.get('app_version', ''))
    tool_html = str(meta.get('tool_html', ''))
    return (
        schema.startswith('stone-calc/')
        or app_version.upper().startswith('V2')
        or tool_html.endswith('規範版V2.html')
    )


def resolve_results(data, inp, cases):
    if is_v2_payload(data) and not cases:
        raise ValueError(
            'V2 專案沒有任何有效案例，已停止 /generate 舊版 Word 產生流程；'
            '請先在 V2 介面建立至少一個案例後再匯出。'
        )
    raw_results = data.get('results')
    if isinstance(raw_results, list) and len(raw_results) == len(cases):
        normalized = [normalize_result(raw, cd, inp) for raw, cd in zip(raw_results, cases)]
        if all(item is not None for item in normalized):
            return normalized, 'frontend_results'
    if is_v2_payload(data):
        raise ValueError(
            'V2 專案缺少有效的 results，已停止 /generate 舊版 Word 產生流程；'
            '請從 V2 介面重新儲存/匯出，或使用「一鍵匯出 Word」。'
        )
    return [calc_case(cd, inp) for cd in cases], 'legacy_python_fallback'


def resolve_meta(data):
    raw_meta = data.get('meta') if isinstance(data.get('meta'), dict) else {}
    spec_refs = deepcopy(DEFAULT_SPEC_REFS)
    if isinstance(raw_meta.get('spec_refs'), dict):
        spec_refs.update(raw_meta['spec_refs'])
    generated_at = raw_meta.get('generated_at')
    if not generated_at:
        generated_at = datetime.now().isoformat(timespec='seconds')
    return {
        'app_version': raw_meta.get('app_version', ''),
        'tool_html': raw_meta.get('tool_html', ''),
        'workflow_mode': raw_meta.get('workflow_mode', ''),
        'calculator_build': raw_meta.get('calculator_build', data.get('calculator_version', 'legacy')),
        'generator': raw_meta.get('generator', 'generate_docx.py'),
        'generated_at': generated_at,
        'input_hash': str(raw_meta.get('input_hash', '')),
        'result_source': raw_meta.get('result_source', ''),
        'spec_refs': spec_refs,
    }


def build_audit_text(meta):
    build = meta.get('calculator_build') or 'legacy'
    short_hash = (meta.get('input_hash') or '—')[:8]
    result_source = meta.get('result_source') or 'unknown'
    return f'工具 {build}｜結果 {result_source}｜雜湊 {short_hash}'


# ─────────────────────────────────────────────
#  XML 工具函式
# ─────────────────────────────────────────────
FONT_NAME = '標楷體'

def _set_font(run, size_pt=12, bold=False, color=None):
    """設定字型（標楷體），含中文 eastAsia 屬性。"""
    run.font.name = FONT_NAME
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    # 強制設定 eastAsia
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), FONT_NAME)
    rFonts.set(qn('w:ascii'), FONT_NAME)
    rFonts.set(qn('w:hAnsi'), FONT_NAME)


def _add_run(para, text, size_pt=12, bold=False, color=None, underline=False):
    run = para.add_run(text)
    _set_font(run, size_pt, bold, color)
    if underline:
        run.font.underline = True
    return run


def _set_para_font(para, size_pt=12):
    """把段落的預設字型設定為標楷體。"""
    pPr = para._p.get_or_add_pPr()
    rPr = pPr.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        pPr.append(rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), FONT_NAME)
    rFonts.set(qn('w:ascii'), FONT_NAME)
    rFonts.set(qn('w:hAnsi'), FONT_NAME)
    sz = rPr.find(qn('w:sz'))
    if sz is None:
        sz = OxmlElement('w:sz')
        rPr.append(sz)
    sz.set(qn('w:val'), str(int(size_pt * 2)))
    szCs = rPr.find(qn('w:szCs'))
    if szCs is None:
        szCs = OxmlElement('w:szCs')
        rPr.append(szCs)
    szCs.set(qn('w:val'), str(int(size_pt * 2)))


def _add_para_border_bottom(para, sz=6, color='AAAAAA'):
    """在段落下方加分隔線。"""
    pPr = para._p.get_or_add_pPr()
    pBdr = pPr.find(qn('w:pBdr'))
    if pBdr is None:
        pBdr = OxmlElement('w:pBdr')
        pPr.append(pBdr)
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(sz))
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)


def _add_para_border_top(para, sz=4, color='DDDDDD'):
    """在段落上方加分隔線。"""
    pPr = para._p.get_or_add_pPr()
    pBdr = pPr.find(qn('w:pBdr'))
    if pBdr is None:
        pBdr = OxmlElement('w:pBdr')
        pPr.append(pBdr)
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), str(sz))
    top.set(qn('w:space'), '1')
    top.set(qn('w:color'), color)
    pBdr.append(top)


def _set_para_spacing(para, before=0, after=0, line=None):
    pPr = para._p.get_or_add_pPr()
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = OxmlElement('w:spacing')
        pPr.append(spacing)
    if before is not None:
        spacing.set(qn('w:before'), str(int(before)))
    if after is not None:
        spacing.set(qn('w:after'), str(int(after)))
    if line:
        spacing.set(qn('w:line'), str(int(line)))
        spacing.set(qn('w:lineRule'), 'auto')


def _set_indent(para, left_cm=0, first_cm=0):
    pPr = para._p.get_or_add_pPr()
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        ind = OxmlElement('w:ind')
        pPr.append(ind)
    if left_cm:
        ind.set(qn('w:left'), str(int(Cm(left_cm).twips)))
    if first_cm:
        ind.set(qn('w:firstLine'), str(int(Cm(first_cm).twips)))


def _set_tab_stop(para, pos_cm, align='right'):
    """在段落加 tab 停格（右對齊）。"""
    pPr = para._p.get_or_add_pPr()
    tabs = pPr.find(qn('w:tabs'))
    if tabs is None:
        tabs = OxmlElement('w:tabs')
        pPr.append(tabs)
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), align)
    tab.set(qn('w:pos'), str(int(Cm(pos_cm).twips)))
    tabs.append(tab)


def _add_field(para, instr):
    """插入 Word 欄位（如頁碼）。"""
    run = para.add_run()
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar_begin)

    run2 = para.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = instr
    run2._r.append(instrText)

    run3 = para.add_run()
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    run3._r.append(fldChar_end)


def _set_page_number_start(section, start=1):
    """設定節起始頁碼。"""
    sectPr = section._sectPr
    pgNumType = sectPr.find(qn('w:pgNumType'))
    if pgNumType is None:
        pgNumType = OxmlElement('w:pgNumType')
        sectPr.append(pgNumType)
    pgNumType.set(qn('w:start'), str(start))


def _set_cell_shading(cell, fill_hex):
    """設定表格儲存格背景色。"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tcPr.append(shd)
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)


def _set_cell_borders(cell, color='CCCCCC', sz=4):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for side in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), str(sz))
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)
        tcBorders.append(el)


def _cell_para(cell, text, size_pt=11, bold=False, color=None,
               align=WD_ALIGN_PARAGRAPH.LEFT, indent_cm=0):
    """清除儲存格既有段落，新增格式化文字。"""
    cell.paragraphs[0].clear()
    para = cell.paragraphs[0]
    para.alignment = align
    _set_para_spacing(para, before=40, after=40)
    if indent_cm:
        _set_indent(para, left_cm=indent_cm)
    _add_run(para, text, size_pt=size_pt, bold=bold, color=color)
    return para


def _make_table(doc, headers, rows, col_widths_cm, header_bg='1A3A5C'):
    """
    建立表格。
    headers: list[str]
    rows: list[list[str | (str, dict)]]   dict 可含 bold/color/align
    col_widths_cm: list[float]
    """
    n_col = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n_col)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 設定欄寬
    for i, w_cm in enumerate(col_widths_cm):
        for row in table.rows:
            cell = row.cells[i]
            cell.width = Cm(w_cm)

    # 標頭列
    hdr_row = table.rows[0]
    for i, hdr in enumerate(headers):
        cell = hdr_row.cells[i]
        _set_cell_shading(cell, header_bg)
        _set_cell_borders(cell, color='FFFFFF', sz=4)
        _cell_para(cell, hdr, size_pt=11, bold=True,
                   color=RGBColor(0xFF, 0xFF, 0xFF),
                   align=WD_ALIGN_PARAGRAPH.CENTER)

    # 資料列
    for r_idx, row_data in enumerate(rows):
        tr = table.rows[r_idx + 1]
        bg = 'F5F5F5' if r_idx % 2 == 0 else 'FFFFFF'
        for c_idx, cell_data in enumerate(row_data):
            cell = tr.cells[c_idx]
            _set_cell_shading(cell, bg)
            _set_cell_borders(cell, color='CCCCCC', sz=4)
            if isinstance(cell_data, tuple):
                text, opts = cell_data
            else:
                text, opts = str(cell_data), {}
            _cell_para(cell, text,
                       size_pt=opts.get('size', 11),
                       bold=opts.get('bold', False),
                       color=opts.get('color', None),
                       align=opts.get('align', WD_ALIGN_PARAGRAPH.LEFT))

    return table


# ─────────────────────────────────────────────
#  頁首頁尾
# ─────────────────────────────────────────────
def _set_header_footer(section, proj_name, section_title, page_start=None, audit_text=''):
    """設定頁首（左：工程名稱，右：節標題）與頁尾（置中頁碼＋版本資訊）。"""
    # 頁首
    header = section.header
    header.is_linked_to_previous = False
    for p in header.paragraphs:
        p._element.getparent().remove(p._element)

    hdr_para = header.add_paragraph()
    hdr_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_tab_stop(hdr_para, pos_cm=14.0, align='right')
    _set_para_spacing(hdr_para, before=0, after=0)
    _add_run(hdr_para, proj_name, size_pt=10)
    _add_run(hdr_para, '\t', size_pt=10)
    _add_run(hdr_para, section_title, size_pt=10)
    _add_para_border_bottom(hdr_para, sz=12, color='333333')

    # 頁尾
    footer = section.footer
    footer.is_linked_to_previous = False
    for p in footer.paragraphs:
        p._element.getparent().remove(p._element)

    ftr_para = footer.add_paragraph()
    ftr_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_para_spacing(ftr_para, before=0, after=0)
    _add_para_border_top(ftr_para, sz=4, color='DDDDDD')
    _add_run(ftr_para, '第 ', size_pt=10)
    _add_field(ftr_para, ' PAGE ')
    _add_run(ftr_para, ' 頁', size_pt=10)
    if audit_text:
        _add_run(ftr_para, '　｜　', size_pt=9, color=RGBColor(0x66, 0x66, 0x66))
        _add_run(ftr_para, audit_text, size_pt=9, color=RGBColor(0x66, 0x66, 0x66))

    if page_start is not None:
        _set_page_number_start(section, page_start)


# ─────────────────────────────────────────────
#  封面頁
# ─────────────────────────────────────────────
def build_cover(doc, inp, meta):
    # 封面使用不同第一頁頁首/尾（空白）
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    # first_page_header/footer 保持空白（預設）

    # 大標題
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_para_spacing(title_para, before=1440, after=240)   # 72pt before
    _add_run(title_para, '結　構　計　算　書', size_pt=36, bold=True)

    # 副標題
    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_para_spacing(sub_para, before=0, after=720)
    _add_run(sub_para, '石材外牆固定構件結構檢核', size_pt=18, bold=False,
             color=RGBColor(0x33, 0x33, 0x33))

    # 資訊表格
    info_items = [
        ('工程名稱', inp.get('proj', '')),
        ('委託單位', inp.get('client', '')),
        ('校核單位', inp.get('checker', '弘一工程顧問有限公司')),
        ('施工地點', inp.get('loc', '')),
        ('高度範圍', inp.get('hrange', '')),
        ('工具介面', meta.get('app_version') or meta.get('tool_html') or 'legacy'),
        ('計算工具版本', meta.get('calculator_build', 'legacy')),
        ('計算結果來源', meta.get('result_source') or 'legacy_python_fallback'),
        ('輸入雜湊', (meta.get('input_hash', '') or '—')[:12]),
    ]

    info_table = doc.add_table(rows=len(info_items), cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    col_w = [3.5, 9.0]
    for r_i, (label, val) in enumerate(info_items):
        row = info_table.rows[r_i]
        row.cells[0].width = Cm(col_w[0])
        row.cells[1].width = Cm(col_w[1])
        for cell in row.cells:
            _set_cell_shading(cell, 'FFFFFF')
            _set_cell_borders(cell, color='DDDDDD', sz=2)

        lbl_para = row.cells[0].paragraphs[0]
        lbl_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _set_para_spacing(lbl_para, before=60, after=60)
        _add_run(lbl_para, label + '：', size_pt=14, bold=True,
                 color=RGBColor(0x1A, 0x3A, 0x5C))

        val_para = row.cells[1].paragraphs[0]
        _set_para_spacing(val_para, before=60, after=60)
        _add_run(val_para, val, size_pt=14)

    # 日期
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_para_spacing(date_para, before=480, after=0)
    _add_para_border_top(date_para, sz=6, color='888888')
    date_str = inp.get('date', datetime.today().strftime('%Y 年 %m 月 %d 日'))
    _add_run(date_para, date_str, size_pt=14)

    spec_refs = meta.get('spec_refs', DEFAULT_SPEC_REFS)
    audit_para = doc.add_paragraph()
    audit_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_para_spacing(audit_para, before=180, after=0)
    _add_run(
        audit_para,
        f'採用規範：風壓 {spec_refs.get("cc", "—")}；錨栓 {spec_refs.get("anchor", "—")}；鋼材 {spec_refs.get("steel", "—")}',
        size_pt=10,
        color=RGBColor(0x66, 0x66, 0x66),
    )
    audit_para2 = doc.add_paragraph()
    audit_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_para_spacing(audit_para2, before=0, after=0)
    _add_run(
        audit_para2,
        f'報表產生時間：{meta.get("generated_at", "")}　｜　產生器：{meta.get("generator", "")}',
        size_pt=9,
        color=RGBColor(0x77, 0x77, 0x77),
    )

    # 備註
    cnote = inp.get('cnote', '')
    if cnote:
        note_para = doc.add_paragraph()
        note_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_para_spacing(note_para, before=120, after=0)
        _add_run(note_para, cnote, size_pt=11, color=RGBColor(0x55, 0x55, 0x55))

    doc.add_page_break()


# ─────────────────────────────────────────────
#  目錄
# ─────────────────────────────────────────────
def build_toc(doc, cases, has_wind_app=False, extra_ann_title=''):
    # 目錄標題
    toc_title = doc.add_paragraph()
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_para_spacing(toc_title, before=0, after=240)
    _add_run(toc_title, '目　　錄', size_pt=22, bold=True)
    _add_para_border_bottom(toc_title, sz=8, color='1A3A5C')

    toc_sections = [
        ('一、', '設計基準'),
        ('二、', '材料性質'),
        ('三、', '風力設計'),
        ('四、', '地震設計'),
        ('五、', '各案例檢核彙整'),
    ]

    for i, (no, name) in enumerate(toc_sections):
        p = doc.add_paragraph()
        _set_tab_stop(p, pos_cm=14.0, align='right')
        _set_para_spacing(p, before=60, after=60)
        _add_run(p, no + name, size_pt=12, bold=True)
        _add_run(p, '\t', size_pt=12)
        _add_run(p, str(i + 1), size_pt=12)

    # 附件清單
    sep = doc.add_paragraph()
    _set_para_spacing(sep, before=120, after=60)
    _add_run(sep, '附件', size_pt=12, bold=True)
    _add_para_border_bottom(sep, sz=4, color='BBBBBB')

    ann_offset = 0
    if has_wind_app:
        p = doc.add_paragraph()
        _set_tab_stop(p, pos_cm=14.0, align='right')
        _set_para_spacing(p, before=40, after=40)
        _set_indent(p, left_cm=1.0)
        _add_run(p, '附件 1　耐風設計計算（C&C）', size_pt=11)
        _add_run(p, '\t', size_pt=11)
        _add_run(p, '6', size_pt=11)
        ann_offset = 1

    for idx, cd in enumerate(cases):
        n = idx + 1
        name = cd.get('name', f'案例{n}')
        tp   = cd.get('type', '')
        label = TYPE_LABEL.get(tp, tp)
        ann_n = n + ann_offset
        p = doc.add_paragraph()
        _set_tab_stop(p, pos_cm=14.0, align='right')
        _set_para_spacing(p, before=40, after=40)
        _set_indent(p, left_cm=1.0)
        _add_run(p, f'附件 {ann_n}　{name}（{label}）', size_pt=11)
        _add_run(p, '\t', size_pt=11)
        _add_run(p, str(5 + ann_n), size_pt=11)

    if extra_ann_title:
        extra_n = len(cases) + ann_offset + 1
        p = doc.add_paragraph()
        _set_tab_stop(p, pos_cm=14.0, align='right')
        _set_para_spacing(p, before=40, after=40)
        _set_indent(p, left_cm=1.0)
        _add_run(p, f'附件 {extra_n}　{extra_ann_title}', size_pt=11)
        _add_run(p, '\t', size_pt=11)
        _add_run(p, str(5 + extra_n), size_pt=11)

    doc.add_page_break()


# ─────────────────────────────────────────────
#  主文（一～五節）
# ─────────────────────────────────────────────
def _heading(doc, text, size_pt=14):
    p = doc.add_paragraph()
    _set_para_spacing(p, before=200, after=80)
    p.paragraph_format.keep_with_next = True
    _add_run(p, text, size_pt=size_pt, bold=True)
    _add_para_border_bottom(p, sz=6, color='AAAAAA')
    return p


def _body(doc, text, indent_cm=1.0, size_pt=12, before=40, after=40):
    p = doc.add_paragraph()
    _set_para_spacing(p, before=before, after=after)
    if indent_cm:
        _set_indent(p, left_cm=indent_cm)
    _add_run(p, text, size_pt=size_pt)
    return p


def build_main_body(doc, inp, cases, results, meta):
    proj = inp.get('proj', '')
    spec_refs = meta.get('spec_refs', DEFAULT_SPEC_REFS)

    # 主文標題
    main_title = doc.add_paragraph()
    main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_para_spacing(main_title, before=0, after=300)
    _add_run(main_title, '石材外牆固定構件檢核計算', size_pt=16, bold=True)
    _add_para_border_bottom(main_title, sz=10, color='1A3A5C')

    # ── 一、設計基準 ──────────────────────────
    _heading(doc, '一、設計基準')
    _body(doc, f'本計算書依據建築物外牆石材施工規範及工程慣例，採用 {spec_refs.get("cc", DEFAULT_SPEC_REFS["cc"])} 進行風壓檢核、'
               f'{spec_refs.get("seismic", DEFAULT_SPEC_REFS["seismic"])} 進行地震力檢核、'
               f'{spec_refs.get("anchor", DEFAULT_SPEC_REFS["anchor"])} 進行錨栓檢核，以及 {spec_refs.get("steel", DEFAULT_SPEC_REFS["steel"])} 進行鋼構件強度檢核。')
    _body(doc, f'設計風速　　Vz = {inp.get("w_v", 37.5)} m/s（地面粗糙度 {inp.get("w_exp", "B")} 類）')
    _body(doc, f'重要性係數　Iw = {inp.get("w_i", 1.0)}')
    _body(doc, f'負壓風壓　　p⁻ = {inp.get("w_neg", 341):.0f} kgf/m²（設計用正壓　p⁺ = {inp.get("w_pos", 426):.0f} kgf/m²）')
    _body(doc, f'安全係數　　φ = {inp.get("w_cf", 1.25)}')
    _body(doc, f'設計譜加速度　SDS = {inp.get("s_sds", 0.6)}，重要性係數 Ip = {inp.get("s_ip", 1.0)}')

    # ── 二、材料性質 ──────────────────────────
    _heading(doc, '二、材料性質')
    mat_rows = [
        ('石材容積重', f'γ = {inp.get("st_gam", 2800):.0f} kgf/m³'),
        ('石材厚度', f't = {inp.get("st_t", 30):.0f} mm'),
        ('背扣角鋼（SS400）', f'Fy = {inp.get("m_fy", 2100):.0f} kgf/cm²'),
        ('A36 鋼材', f'Fy = {inp.get("m_fya36", 2500):.0f} kgf/cm²'),
        ('背扣螺絲允許拔出力', f'Ta = {inp.get("m_screw_ta", 255.3):.1f} kgf'),
        ('膨脹螺栓允許拉力', anc_ta_detail(inp)['full']),
        ('　查表依據', anc_ta_detail(inp)['source']),
        ('馬車螺栓允許剪力', f'Va = {inp.get("m_mc_va", 567):.0f} kgf'),
        ('馬車螺栓允許拉力', f'Ta = {inp.get("m_mc_ta", 709):.0f} kgf'),
        ('插銷直徑', f'd = {inp.get("m_pin_d", 5):.0f} mm，Fy = {inp.get("m_pin_fy", 2100):.0f} kgf/cm²'),
    ]
    _make_table(doc, ['項目', '規格／允許值'], mat_rows, [5.5, 8.5])

    # ── 三、風力設計 ──────────────────────────
    _heading(doc, '三、風力設計')
    w_pos  = float(inp.get('w_pos', 426))
    w_neg  = float(inp.get('w_neg', 341))
    w_cf   = float(inp.get('w_cf', 1.25))
    w_ctrl = max(w_pos, w_neg)
    _body(doc, f'正風壓 p⁺ = {w_pos:.0f} kgf/m²，負風壓 |p⁻| = {w_neg:.0f} kgf/m²')
    _body(doc, f'設計控制值 = max({w_pos:.0f}, {w_neg:.0f}) = {w_ctrl:.0f} kgf/m²，乘以載重組合係數 {w_cf}：')
    _body(doc, f'設計風壓 FW = {w_ctrl:.0f} × {w_cf} = {w_ctrl * w_cf:.1f} kgf/m²')
    _body(doc, f'設計風力 PW = FW × A（各案例分別計算，詳附件）')

    # ── 四、地震設計 ──────────────────────────
    _heading(doc, '四、地震設計')
    s_sds = float(inp.get('s_sds', 0.6))
    s_ip  = float(inp.get('s_ip', 1.0))
    Fph_val = 1.6 * s_sds * s_ip
    _body(doc, f'非結構構件地震設計力（依{spec_refs.get("seismic", DEFAULT_SPEC_REFS["seismic"])}，取包絡上限值 1.6·SDS·Ip·Wp）：')
    _body(doc, f'Fph = 1.6 × SDS × Ip × Wp = 1.6 × {s_sds} × {s_ip} × Wp = {Fph_val:.3f} Wp')
    _body(doc, f'地震水平力 PE = Fph × A，地震垂直力 PEV = 0.5 × Fph × A')

    # ── 五、各案例檢核彙整 ──────────────────────
    _heading(doc, '五、各案例檢核彙整')

    summary_headers = ['編號', '說明', '固定形式', '與結構體固定方式', '檢核結果']
    summary_rows = []
    for idx, (cd, res) in enumerate(zip(cases, results)):
        n = idx + 1
        name  = cd.get('name', f'案例{n}')
        tp    = cd.get('type', '')
        label = TYPE_LABEL.get(tp, tp)
        anc   = cd.get('anc_spec', cd.get('anc_key', ''))
        ok    = res.get('all_ok', False)
        ok_text = '符合' if ok else '不符合'
        ok_color = RGBColor(0x15, 0x52, 0x15) if ok else RGBColor(0xAA, 0x00, 0x00)
        summary_rows.append([
            (str(n), {'align': WD_ALIGN_PARAGRAPH.CENTER}),
            name,
            label,
            anc,
            (ok_text, {'bold': True, 'color': ok_color, 'align': WD_ALIGN_PARAGRAPH.CENTER}),
        ])
    _make_table(doc, summary_headers, summary_rows, [1.2, 3.5, 4.0, 4.0, 1.5])

    # 結語
    all_pass = all(r.get('all_ok', False) for r in results)
    conclusion_para = doc.add_paragraph()
    conclusion_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_para_spacing(conclusion_para, before=360, after=120)
    _add_para_border_top(conclusion_para, sz=8, color='333333')
    _add_para_border_bottom(conclusion_para, sz=8, color='333333')
    if all_pass:
        concl_text = '各案例固定構件結構檢核均符合規範要求，可安全施工。'
    else:
        concl_text = '部分案例固定構件結構檢核不符合規範要求，請依附件說明進行調整。'
    _add_run(conclusion_para, concl_text, size_pt=13, bold=True)


# ─────────────────────────────────────────────
#  附件 1：耐風設計計算（C&C）
# ─────────────────────────────────────────────
TERRAIN_NAMES = {'A': '地況 A（大城市市中心）', 'B': '地況 B（都市郊區）', 'C': '地況 C（開闊平坦地）'}
TERRAIN_ALPHA = {'A': 0.32, 'B': 0.25, 'C': 0.15}
TERRAIN_ZG    = {'A': 500,  'B': 400,  'C': 300}
ENCL_NAMES    = {'enclosed': '封閉式', 'partial': '部分封閉式'}

def build_wind_appendix(doc, cc, inp):
    """cc = inp['cc_report'] dict"""
    city    = cc.get('city', '')
    dist    = cc.get('dist', '')
    V       = cc.get('V', 32.5)
    I       = cc.get('I', 1.0)
    terrain = cc.get('terrain', 'B')
    h       = cc.get('h', 15.0)
    z       = cc.get('z', h)
    A       = cc.get('A', 1.0)
    kzt     = cc.get('kzt', 1.0)
    zone    = cc.get('zone', 'zone4')
    encl    = cc.get('encl', 'enclosed')
    qh      = cc.get('qh', 0.0)
    qz      = cc.get('qz', 0.0)
    GCp_pos = cc.get('GCp_pos', 0.0)
    GCp_neg = cc.get('GCp_neg', 0.0)
    GCpi    = cc.get('GCpi', 0.0)
    p_pos   = cc.get('p_pos', 0.0)
    p_neg   = cc.get('p_neg', 0.0)
    isLE18  = cc.get('isLE18', h <= 18)

    alpha = TERRAIN_ALPHA.get(terrain, 0.25)
    zg    = TERRAIN_ZG.get(terrain, 400)

    # 標題
    app_title = doc.add_paragraph()
    app_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_para_spacing(app_title, before=0, after=200)
    _add_run(app_title, '附件 1　耐風設計計算（外牆構件與外覆材 C&C）', size_pt=14, bold=True)
    _add_para_border_bottom(app_title, sz=10, color='1A3A5C')

    loc_str = f'{city}{dist}' if city else inp.get('loc', '')
    _body(doc, f'依據「建築物耐風設計規範及解說」（107年版）§3，採構件與外覆材（C&C）設計風壓。', indent_cm=0)
    _body(doc, f'施工地點：{loc_str}', indent_cm=0)

    # 一、設計條件
    _heading(doc, '一、設計條件', size_pt=13)
    cond_rows = [
        ('基本設計風速',    f'V₁₀(C) = {V} m/s'),
        ('用途係數',        f'I = {I}'),
        ('地況種類',        TERRAIN_NAMES.get(terrain, terrain)),
        ('地形係數',        f'Kzt = {kzt}'),
        ('屋頂高度',        f'h = {h} m'),
        ('構材標高',        f'z = {z} m'),
        ('有效受風面積',    f'A = {A} m²'),
        ('外牆計算區域',    'Zone 4（一般外牆）' if zone=='zone4' else 'Zone 5（轉角）'),
        ('建築圍閉種類',    ENCL_NAMES.get(encl, encl)),
        ('h ≤ 18m 分類',   '是（圖3.1）' if isLE18 else '否（圖3.2）'),
    ]
    _make_table(doc, ['條件', '數值'], cond_rows, [5.0, 9.0])

    # 二、速度壓計算
    _heading(doc, '二、速度壓計算（Eq 2.6）', size_pt=13)
    zUse_h = max(h, 9.0 if terrain=='B' else (18.0 if terrain=='A' else 4.5))
    Kz_h   = 2.774 * (zUse_h / zg) ** (2 * alpha)
    zUse_z = max(z, 9.0 if terrain=='B' else (18.0 if terrain=='A' else 4.5))
    Kz_z   = 2.774 * (zUse_z / zg) ** (2 * alpha)

    calc_lines = [
        f'K(h) = 2.774 × (h/zg)^(2α) = 2.774 × ({zUse_h}/{zg})^(2×{alpha}) = {Kz_h:.4f}',
        f'q_h  = 0.06 × K(h) × Kzt × (I×V)² = 0.06 × {Kz_h:.4f} × {kzt} × ({I}×{V})²'
        f'  = {qh:.2f} kgf/m²',
    ]
    if not isLE18:
        calc_lines += [
            f'K(z) = 2.774 × ({zUse_z}/{zg})^(2×{alpha}) = {Kz_z:.4f}',
            f'q_z  = 0.06 × {Kz_z:.4f} × {kzt} × ({I}×{V})² = {qz:.2f} kgf/m²（h>18m 牆面用）',
        ]
    for ln in calc_lines:
        _body(doc, ln, indent_cm=1.0, size_pt=11, before=30, after=30)

    # 三、GCp 係數
    _heading(doc, '三、外壓係數 GCp（C&C）', size_pt=13)
    fig_ref = '圖 3.1(a)' if isLE18 else '圖 3.2'
    zone_lbl = 'Zone 4' if zone=='zone4' else 'Zone 5'
    gcpi_str = f'±{GCpi:.3f}' if GCpi else '0.000'
    coef_rows = [
        ('參考圖表',    f'{fig_ref}，{zone_lbl}'),
        ('有效受風面積', f'A = {A} m²（對數插值）'),
        ('GCp（正）',   f'+{GCp_pos:.2f}'),
        ('GCp（負）',   f'{GCp_neg:.2f}'),
        ('GCpi（內壓）', gcpi_str),
    ]
    _make_table(doc, ['項目', '數值'], coef_rows, [5.0, 9.0])

    # 四、設計風壓
    _heading(doc, '四、設計風壓（Eq 3.1/3.2）', size_pt=13)
    qPos = qz if (not isLE18) else qh
    _body(doc, f'p⁺ = q × GCp(+) + q_h × GCpi'
               f' = {qPos:.2f} × {GCp_pos:.2f} + {qh:.2f} × {GCpi:.3f}'
               f' = {p_pos:.1f} kgf/m²', indent_cm=1.0, size_pt=11, before=30, after=30)
    _body(doc, f'p⁻ = q_h × GCp(-) − q_h × GCpi'
               f' = {qh:.2f} × {GCp_neg:.2f} − {qh:.2f} × {GCpi:.3f}'
               f' = {p_neg:.1f} kgf/m²', indent_cm=1.0, size_pt=11, before=30, after=30)

    result_para = doc.add_paragraph()
    result_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_para_spacing(result_para, before=240, after=120)
    _add_para_border_top(result_para, sz=8, color='1A3A5C')
    _add_para_border_bottom(result_para, sz=8, color='1A3A5C')
    _add_run(result_para,
             f'採用設計正壓 p⁺ = {p_pos:.1f} kgf/m²，設計負壓 |p⁻| = {abs(p_neg):.1f} kgf/m²',
             size_pt=13, bold=True)


# ─────────────────────────────────────────────
#  附件（各案例詳細計算）
# ─────────────────────────────────────────────
def build_case_appendix(doc, cd, res, n, inp):
    """n: 附件編號（從1起）"""
    name  = cd.get('name', f'案例{n}')
    tp    = cd.get('type', '')
    label = TYPE_LABEL.get(tp, tp)
    proj  = inp.get('proj', '')

    # ── 附件標題 ──────────────────────────────
    app_title = doc.add_paragraph()
    app_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_para_spacing(app_title, before=0, after=200)
    _add_run(app_title, f'附件 {n}　{name}（{label}）', size_pt=14, bold=True)
    _add_para_border_bottom(app_title, sz=10, color='1A3A5C')

    # ── 一、設計參數 ──────────────────────────
    _heading(doc, '一、設計參數', size_pt=13)

    w   = float(cd.get('w', 870))
    h   = float(cd.get('h', 800))
    N   = int(cd.get('N', 4))
    bh  = float(cd.get('bh', 10))
    d1  = float(cd.get('d1', 8))
    Lt  = float(cd.get('Lt', 0.5))
    LL  = float(cd.get('LL', 5))
    d0  = float(cd.get('d0', 1.2))
    has_mc = bool(cd.get('hasMC', True))
    anc_spec  = cd.get('anc_spec', cd.get('anc_key', ''))
    conn_spec = cd.get('conn_spec', '')
    mc_h1, mc_h2 = res['mc_h1'], res['mc_h2']

    param_rows = [
        ('石材板面尺寸',      f'W = {w:.0f} mm，H = {h:.0f} mm'),
        ('石材面積',          f'A = {res["A"]:.4f} m²'),
        ('固定點數',          f'N = {N} 點'),
        ('背扣組件高度',      f'h = {bh} cm'),
        ('膨脹螺栓間距',      f'd₁ = {d1} cm'),
        ('角鋼尺寸',          f'LL = {LL} cm，Lt = {Lt} cm，d₀ = {d0} cm（開孔徑）'),
        ('與結構體固定方式',  anc_spec),
        ('石材固定螺絲規格',  conn_spec),
    ]
    if has_mc:
        param_rows.append(('馬車螺栓臂長',  f'h₁ = {mc_h1} cm，h₂ = {mc_h2} cm'))

    _make_table(doc, ['項目', '數值'], param_rows, [4.5, 9.5])

    # ── 二、載重計算 ──────────────────────────
    _heading(doc, '二、載重計算', size_pt=13)

    load_lines = [
        f'石材面重　　Wp = γ × t = {inp.get("st_gam", 2800):.0f} × {inp.get("st_t", 30):.0f}/1000 = {res["Wp"]:.2f} kgf/m²',
        f'石材自重　　G = A × Wp = {res["A"]:.4f} × {res["Wp"]:.2f} = {res["G"]:.2f} kgf',
        '',
        f'【風力】',
        f'設計風壓　　FW = p⁺ × φ = {inp.get("w_pos", 426):.0f} × {inp.get("w_cf", 1.25)} = {res["FW"]:.2f} kgf/m²',
        f'設計風力　　PW = A × FW = {res["A"]:.4f} × {res["FW"]:.2f} = {res["PW"]:.2f} kgf',
        '',
        f'【地震力】',
        f'地震加速度　Fph = 1.6 × SDS × Ip × Wp = 1.6 × {inp.get("s_sds", 0.6)} × {inp.get("s_ip", 1.0)} × {res["Wp"]:.2f} = {res["Fph"]:.2f} kgf/m²',
        f'地震水平力　PE = Fph × A = {res["Fph"]:.2f} × {res["A"]:.4f} = {res["PE"]:.2f} kgf',
        f'地震垂直力　PEV = 0.5 × PE = {res["PEV"]:.2f} kgf',
        '',
        f'【設計力】',
        f'垂直設計力　S = G + PEV = {res["G"]:.2f} + {res["PEV"]:.2f} = {res["S"]:.2f} kgf',
        f'水平設計力　P = max(PE, PW) = max({res["PE"]:.2f}, {res["PW"]:.2f}) = {res["P"]:.2f} kgf'
        + ('（風力控制）' if res['gov_wind'] else '（地震控制）'),
        '',
        f'每點垂直力　V = S ÷ N = {res["S"]:.2f} ÷ {N} = {res["V"]:.2f} kgf',
        f'每點水平力　T = P ÷ N = {res["P"]:.2f} ÷ {N} = {res["T"]:.2f} kgf',
    ]

    for line in load_lines:
        if line == '':
            doc.add_paragraph()
        else:
            _body(doc, line, indent_cm=1.0, size_pt=11, before=30, after=30)

    # ── 三、結構檢核 ──────────────────────────
    _heading(doc, '三、結構檢核', size_pt=13)

    for ck in res['checks']:
        p = doc.add_paragraph()
        _set_tab_stop(p, pos_cm=13.0, align='right')
        _set_para_spacing(p, before=40, after=40)
        _set_indent(p, left_cm=0.5)

        ok_text  = '符合' if ck['pass'] else '不符合'
        ok_color = RGBColor(0x15, 0x52, 0x15) if ck['pass'] else RGBColor(0xAA, 0x00, 0x00)

        _add_run(p, f'{ck["no"]} {ck["item"]}　', size_pt=11, bold=True)
        _add_run(p, ck['formula'], size_pt=10,
                 color=RGBColor(0x33, 0x33, 0x33))
        _add_run(p, f'\t【{ok_text}】', size_pt=11, bold=True, color=ok_color)

    # 附件結語
    all_ok = res.get('all_ok', False)
    concl  = doc.add_paragraph()
    concl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_para_spacing(concl, before=240, after=120)
    _add_para_border_top(concl, sz=4, color='CCCCCC')
    if all_ok:
        txt = f'本案例（{name}）各項固定構件結構檢核均符合規範要求。'
    else:
        txt = f'本案例（{name}）部分固定構件結構檢核不符合規範要求，請調整設計。'
    _add_run(concl, txt, size_pt=12, bold=True,
             color=RGBColor(0x15, 0x52, 0x15) if all_ok else RGBColor(0xAA, 0x00, 0x00))


# ─────────────────────────────────────────────
#  附圖附件（圖說與試驗資料）
# ─────────────────────────────────────────────
def build_extra_appendix(doc, n, title, note, images, inp):
    import base64, io, tempfile

    app_title = doc.add_paragraph()
    app_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_para_spacing(app_title, before=0, after=200)
    _add_run(app_title, f'附件 {n}　{title}', size_pt=14, bold=True)
    _add_para_border_bottom(app_title, sz=10, color='1A3A5C')

    if note:
        _body(doc, note, indent_cm=0)

    if not images:
        p = doc.add_paragraph()
        _set_para_spacing(p, before=400, after=400)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(p, '（本附件無圖片）', size_pt=12,
                 color=RGBColor(0xAA, 0xAA, 0xAA))
        return

    tmp_files = []
    try:
        # 每兩張並排
        for i in range(0, len(images), 2):
            pair = images[i:i+2]
            tbl  = doc.add_table(rows=2, cols=len(pair))
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            col_w = 14.0 / len(pair)

            for j, img_data in enumerate(pair):
                src = img_data.get('src', '')
                cap = img_data.get('caption', '')

                # 解碼 base64 dataURL  → 臨時檔
                if ',' in src:
                    _, b64 = src.split(',', 1)
                else:
                    b64 = src
                raw = base64.b64decode(b64)
                tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
                tmp.write(raw)
                tmp.close()
                tmp_files.append(tmp.name)

                # 圖片 cell
                img_cell = tbl.rows[0].cells[j]
                img_cell.width = Cm(col_w)
                _set_cell_borders(img_cell, color='FFFFFF', sz=0)
                ip = img_cell.paragraphs[0]
                ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _set_para_spacing(ip, before=60, after=40)
                try:
                    run = ip.add_run()
                    run.add_picture(tmp.name, width=Cm(col_w - 0.5))
                except Exception:
                    _add_run(ip, f'[圖片 {i+j+1} 無法嵌入]', size_pt=10,
                             color=RGBColor(0xAA, 0x00, 0x00))

                # 圖說 cell
                cap_cell = tbl.rows[1].cells[j]
                cap_cell.width = Cm(col_w)
                _set_cell_borders(cap_cell, color='FFFFFF', sz=0)
                cp = cap_cell.paragraphs[0]
                cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _set_para_spacing(cp, before=0, after=80)
                if cap:
                    _add_run(cp, cap, size_pt=11,
                             color=RGBColor(0x33, 0x33, 0x33))

            doc.add_paragraph()  # 圖組間距

    finally:
        for f in tmp_files:
            try:
                import os; os.unlink(f)
            except Exception:
                pass


# ─────────────────────────────────────────────
#  頁面邊界
# ─────────────────────────────────────────────
def _set_margins(section):
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.0)
    section.top_margin    = Cm(2.2)
    section.bottom_margin = Cm(1.6)


# ─────────────────────────────────────────────
#  主產生流程
# ─────────────────────────────────────────────
def generate_report(json_path: str) -> str:
    with open(json_path, encoding='utf-8') as f:
        data = json.load(f)

    inp   = data.get('inp', {})
    cases = data.get('cases', [])
    meta = resolve_meta(data)

    # 優先採用前端共用計算核心輸出的結果；若舊版 JSON 缺資料，再退回 Python fallback。
    results, result_source = resolve_results(data, inp, cases)
    meta['result_source'] = result_source
    audit_text = build_audit_text(meta)

    doc = Document()

    # 設定預設字型（英文部分）
    doc.styles['Normal'].font.name = FONT_NAME
    doc.styles['Normal'].font.size = Pt(12)
    rPr_el = doc.styles['Normal']._element.get_or_add_rPr()
    rFonts_el = rPr_el.find(qn('w:rFonts'))
    if rFonts_el is None:
        rFonts_el = OxmlElement('w:rFonts')
        rPr_el.insert(0, rFonts_el)
    rFonts_el.set(qn('w:eastAsia'), FONT_NAME)

    # 首節邊界
    section0 = doc.sections[0]
    _set_margins(section0)

    # 封面
    build_cover(doc, inp, meta)

    # 設定主文頁首頁尾（different_first_page 後的預設頁首）
    _set_header_footer(section0, inp.get('proj', ''), '主　文', page_start=1, audit_text=audit_text)

    # 判斷是否有風力附件
    cc_report = inp.get('cc_report') or {}
    has_wind_app = bool(cc_report)

    # 目錄
    extra_title = inp.get('extra_ann_title', '') if inp.get('extra_ann_on') else ''
    build_toc(doc, cases, has_wind_app=has_wind_app, extra_ann_title=extra_title)

    # 主文
    build_main_body(doc, inp, cases, results, meta)

    # 附件 1：耐風設計計算（若有）
    proj = inp.get('proj', '')
    ann_offset = 0
    if has_wind_app:
        wind_sec = doc.add_section(WD_SECTION.NEW_PAGE)
        _set_margins(wind_sec)
        _set_header_footer(wind_sec, proj, '附件 1　耐風設計計算', audit_text=audit_text)
        build_wind_appendix(doc, cc_report, inp)
        ann_offset = 1

    # 附件（每個案例獨立節）
    for idx, (cd, res) in enumerate(zip(cases, results)):
        n     = idx + 1
        ann_n = n + ann_offset
        name  = cd.get('name', f'案例{n}')

        new_section = doc.add_section(WD_SECTION.NEW_PAGE)
        _set_margins(new_section)
        section_title = f'附件 {ann_n}　{name}'
        _set_header_footer(new_section, proj, section_title, audit_text=audit_text)

        build_case_appendix(doc, cd, res, ann_n, inp)

    # 附圖附件（若啟用）
    if inp.get('extra_ann_on'):
        extra_n     = len(cases) + ann_offset + 1
        extra_title = inp.get('extra_ann_title', '相關圖說與試驗數據')
        extra_note  = inp.get('extra_ann_note',  '')
        extra_imgs  = inp.get('extra_ann_images', [])

        ex_sec = doc.add_section(WD_SECTION.NEW_PAGE)
        _set_margins(ex_sec)
        _set_header_footer(ex_sec, proj, f'附件 {extra_n}　{extra_title}', audit_text=audit_text)
        build_extra_appendix(doc, extra_n, extra_title, extra_note, extra_imgs, inp)

    # 輸出路徑
    json_p = Path(json_path)
    out_name = json_p.stem + '_計算書.docx'
    out_path = json_p.parent / out_name
    doc.save(str(out_path))
    return str(out_path)


# ─────────────────────────────────────────────
#  進入點
# ─────────────────────────────────────────────
if __name__ == '__main__':
    script_dir = Path(__file__).parent

    if len(sys.argv) > 1:
        json_path = Path(sys.argv[1])
    else:
        # 搜尋順序：腳本資料夾 → 使用者 Downloads 資料夾
        downloads = Path.home() / 'Downloads'
        candidates = (sorted(script_dir.glob('stone_report*.json')) +
                      sorted(script_dir.glob('*.json')) +
                      sorted(downloads.glob('stone_report*.json')))
        if not candidates:
            print('找不到 JSON，請先在工具中按「匯出 Word」')
            sys.exit(1)
        json_path = max(candidates, key=lambda p: p.stat().st_mtime)
        print(f'使用：{json_path}')

    print(f'讀取：{json_path}')
    try:
        out = generate_report(str(json_path))
        print(f'完成：{out}')
        os.startfile(out)
    except Exception as e:
        import traceback
        traceback.print_exc()
        print(f'\n錯誤：{e}')
        sys.exit(1)
