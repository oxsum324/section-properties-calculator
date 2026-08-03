# -*- coding: utf-8 -*-
"""End-to-end artifact checks for the formal auto_word export path."""
from __future__ import annotations

import json
import hashlib
import logging
import os
import shutil
import subprocess
import sys
import tempfile
import time
import unittest
import urllib.error
import urllib.request
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from pypdf import PdfReader

import auto_word
import server
from audit_schema import REQUIRED_AUDIT_PATHS, nested
from server_smoke_test import PAGE_ONLY_REPORT_STATUS_NEEDLES, report_text, v2_payload


def file_sha256(file_path: Path) -> str:
    return hashlib.sha256(file_path.read_bytes()).hexdigest()


ROOT = Path(__file__).resolve().parent
ROOT_URL = f'http://127.0.0.1:{server.PORT}'
GOLDEN_CASE_ID = 'case_01_standard_safe'
GOLDEN_CASE_PATH = ROOT / 'tests' / 'golden' / f'{GOLDEN_CASE_ID}.json'
GOLDEN_INPUT_PATH = ROOT / 'tests' / 'golden' / 'inputs' / f'{GOLDEN_CASE_ID}.json'


def local_server_status() -> dict | None:
    try:
        with urllib.request.urlopen(f'{ROOT_URL}/status', timeout=2) as response:
            if response.status != 200:
                return None
            return json.loads(response.read().decode('utf-8'))
    except (OSError, urllib.error.URLError, json.JSONDecodeError):
        return None


def wait_for_server(timeout_sec: float = 10.0) -> bool:
    deadline = time.time() + timeout_sec
    while time.time() < deadline:
        status = local_server_status()
        if status and status.get('server_version') == server.SERVER_VERSION:
            return True
        time.sleep(0.25)
    return False


def formal_payload() -> dict:
    payload = v2_payload()
    payload['inp'].update({
        'proj': 'Auto Word Formal Artifact',
        'review_summary_on': True,
    })
    payload['cases'][0]['id'] = 1
    payload['reviewNotes'] = {
        '1::插銷 剪力': '正式路徑註記：已加大插銷後覆核。',
    }
    payload['meta'].update({
        'calculator_version': '2026.04.23-spec-core11',
        'calc_source_hash': 'sha256:' + 'a' * 64,
        'input_hash': 'sha256:' + 'b' * 64,
        'result_hash': 'sha256:' + 'c' * 64,
    })
    review_summary = payload['meta']['review_summary']
    review_summary['priority_report_reading_status'] = {
        'title': '優先建議報告閱讀狀態',
        'decision': '暫勿作附件',
        'note': '頁面輔助，不會寫入計算書或列印 PDF。',
    }
    review_summary['delivery_quality']['reasons'].append('報告閱讀狀態：暫勿作附件')
    return payload


def replay_golden_formal_payload() -> tuple[dict, dict]:
    from playwright.sync_api import sync_playwright

    golden = json.loads(GOLDEN_CASE_PATH.read_text(encoding='utf-8'))
    payload = json.loads(json.dumps(golden['input']['raw'], ensure_ascii=False))
    payload['inp'].update({
        'proj': 'Stone Golden Replay Formal Artifact',
        'review_summary_on': True,
    })
    payload['cases'][0]['id'] = 1
    payload['reviewNotes'] = {
        '1::背擴孔深度': 'Golden replay：已核對控制強度與細部檢核。',
    }

    with sync_playwright() as playwright:
        browser = playwright.chromium.launch(headless=True)
        try:
            page = browser.new_page(viewport={'width': 1280, 'height': 1800})
            page.goto(auto_word.TOOL_URL, wait_until='domcontentloaded', timeout=60000)
            page.evaluate(
                """(args) => {
                  localStorage.clear();
                  localStorage.setItem(args.key, JSON.stringify(args.data));
                }""",
                {'key': auto_word.STORAGE_KEY, 'data': payload},
            )
            page.goto(auto_word.TOOL_URL, wait_until='networkidle', timeout=60000)
            page.wait_for_function(
                """() => document.querySelectorAll('#preview-sheets .a4').length > 0""",
                timeout=30000,
            )
            replayed = page.evaluate("""async () => {
              const hashes = await v2RefreshDeliveryHashes();
              const payload = await buildProjectPayload();
              payload.meta.normalized_input_hash = hashes?.normalized_input_hash || '';
              payload.meta.result_hash = hashes?.result_hash || '';
              payload.meta.calc_source_hash = window.__CALC_SOURCE_HASH || '';
              payload.meta.calculator_version = payload.meta.calculator_build || '';
              return payload;
            }""")
        finally:
            browser.close()

    replayed['meta']['review_summary']['priority_report_reading_status'] = {
        'title': '優先建議報告閱讀狀態',
        'decision': '暫勿作附件',
        'note': '頁面輔助，不會寫入計算書或列印 PDF。',
    }
    replayed['meta']['review_summary']['delivery_quality']['reasons'].append('報告閱讀狀態：暫勿作附件')
    return replayed, golden


def extract_docx_text(docx_path: Path) -> tuple[str, Document]:
    document = Document(str(docx_path))
    paragraph_text = '\n'.join(paragraph.text for paragraph in document.paragraphs)
    table_text = '\n'.join(
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    )
    return f'{paragraph_text}\n{table_text}', document


def rendered_evidence_dir() -> Path | None:
    override = os.environ.get('STONE_RENDERED_EVIDENCE_DIR', '').strip()
    if override:
        return Path(override).resolve()
    run_dir = os.environ.get('PREFLIGHT_RUN_DIR', '').strip()
    if os.environ.get('PREFLIGHT_RELEASE') == '1' and run_dir:
        return Path(run_dir).resolve() / 'rendered-delivery-evidence' / 'stone-formal'
    return None


def persist_rendered_evidence(
    pdf_path: Path,
    docx_path: Path,
    audit_path: Path,
    *,
    page_count: int,
    pdf_text_length: int,
    docx_text_length: int,
    result_reconciliation: dict,
) -> Path | None:
    output_dir = rendered_evidence_dir()
    if output_dir is None:
        return None
    output_dir.mkdir(parents=True, exist_ok=True)
    pdf_copy = output_dir / 'stone-auto-word-formal.pdf'
    docx_copy = output_dir / 'stone-auto-word-formal.docx'
    audit_copy = output_dir / 'stone-auto-word-formal.audit.json'
    shutil.copy2(pdf_path, pdf_copy)
    shutil.copy2(docx_path, docx_copy)
    shutil.copy2(audit_path, audit_copy)
    summary = {
        'schemaVersion': 1,
        'family': 'stone-formal',
        'pass': True,
        'generatedAt': datetime.now(timezone.utc).isoformat(),
        'records': [{
            'key': 'stone-fixing',
            'artifact': pdf_copy.name,
            'artifactBytes': pdf_copy.stat().st_size,
            'artifactSha256': file_sha256(pdf_copy),
            'document': docx_copy.name,
            'documentBytes': docx_copy.stat().st_size,
            'documentSha256': file_sha256(docx_copy),
            'evidence': audit_copy.name,
            'evidenceBytes': audit_copy.stat().st_size,
            'evidenceSha256': file_sha256(audit_copy),
            'pageCount': page_count,
            'pdfTextLength': pdf_text_length,
            'docxTextLength': docx_text_length,
            'resultReconciliation': result_reconciliation,
        }],
    }
    summary_path = output_dir / 'rendered-delivery-evidence-summary.json'
    summary_path.write_text(json.dumps(summary, ensure_ascii=False, indent=2) + '\n', encoding='utf-8')
    return summary_path


class AutoWordArtifactTests(unittest.TestCase):
    server_process: subprocess.Popen | None = None

    @classmethod
    def setUpClass(cls) -> None:
        status = local_server_status()
        if status:
            if status.get('tool_html') != server.TOOL_HTML:
                raise RuntimeError(f'Port {server.PORT} is serving an unexpected stone tool: {status}')
            return
        cls.server_process = subprocess.Popen(
            [sys.executable, 'server.py'],
            cwd=ROOT,
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
        )
        if not wait_for_server():
            cls.tearDownClass()
            raise RuntimeError('local stone server did not start for auto_word artifact test')

    @classmethod
    def tearDownClass(cls) -> None:
        if cls.server_process and cls.server_process.poll() is None:
            cls.server_process.terminate()
            try:
                cls.server_process.wait(timeout=3)
            except subprocess.TimeoutExpired:
                cls.server_process.kill()
        cls.server_process = None

    def assertAuditSchema(self, report: dict) -> None:
        for path in REQUIRED_AUDIT_PATHS:
            sentinel = object()
            value = nested(report, *path, default=sentinel)
            self.assertIsNot(value, sentinel, f"missing audit path: {'.'.join(path)}")

    def test_review_summary_import_is_true_opt_in_without_toc_ghost(self) -> None:
        from playwright.sync_api import sync_playwright

        with sync_playwright() as playwright:
            browser = playwright.chromium.launch(headless=True)
            try:
                for enabled in (False, True):
                    payload = formal_payload()
                    payload['inp']['review_summary_on'] = enabled
                    page = browser.new_page(viewport={'width': 1280, 'height': 1800})
                    try:
                        page.goto(auto_word.TOOL_URL, wait_until='domcontentloaded', timeout=60000)
                        page.evaluate(
                            """(args) => localStorage.setItem(args.key, JSON.stringify(args.data))""",
                            {'key': auto_word.STORAGE_KEY, 'data': payload},
                        )
                        page.goto(auto_word.TOOL_URL, wait_until='networkidle', timeout=60000)
                        page.wait_for_function(
                            """() => document.querySelectorAll('#preview-sheets .a4').length > 0""",
                            timeout=30000,
                        )
                        page.wait_for_timeout(500)
                        state = page.evaluate(
                            """() => ({
                              checked: document.getElementById('review_summary_on')?.checked || false,
                              text: document.querySelector('#preview-sheets')?.innerText || '',
                              toc: document.querySelector('.toc-page')?.innerText || '',
                            })"""
                        )
                    finally:
                        page.close()

                    self.assertEqual(state['checked'], enabled)
                    if enabled:
                        self.assertIn('送審速覽', state['text'])
                        self.assertIn('送審速覽', state['toc'])
                        self.assertIn('設計者註記', state['text'])
                        self.assertIn('正式路徑註記：已加大插銷後覆核。', state['text'])
                    else:
                        self.assertNotIn('送審速覽', state['text'])
                        self.assertNotIn('送審速覽', state['toc'])
            finally:
                browser.close()

    def test_formal_auto_word_outputs_pdf_docx_with_review_notes_and_boundaries(self) -> None:
        payload, golden = replay_golden_formal_payload()
        replayed_result = payload['results'][0]
        expected = golden['result']['summary']
        governing = next(
            check for check in replayed_result.get('checks', [])
            if check.get('item') == expected['governing_check']
        )
        self.assertAlmostEqual(replayed_result['Fph'], expected['Fph'], places=6)
        self.assertAlmostEqual(replayed_result['T'], expected['Tmax'], places=6)
        self.assertAlmostEqual(replayed_result['V'], expected['Vmax'], places=6)
        self.assertAlmostEqual(governing['dcr'], expected['DCR_governing'], places=4)
        self.assertEqual(governing['item'], expected['governing_check'])
        self.assertEqual(payload['cases'][0]['name'], expected['governing_case'])
        self.assertRegex(payload['meta']['input_hash'], r'^(?:sha256:)?[0-9a-f]{64}$')
        self.assertRegex(payload['meta']['result_hash'], r'^(?:sha256:)?[0-9a-f]{64}$')
        self.assertRegex(payload['meta']['calc_source_hash'], r'^sha256:[0-9a-f]{64}$')
        validation = auto_word.validate_payload_for_formal_output(payload)
        self.assertTrue(validation['is_formal'], validation)

        with tempfile.TemporaryDirectory() as tmp:
            out_dir = Path(tmp)
            previous_disable_level = logging.root.manager.disable
            logging.disable(logging.INFO)
            try:
                docx_path = auto_word.auto_export_word(payload, out_dir, 'stone_auto_word_formal')
            finally:
                logging.disable(previous_disable_level)
            pdf_path = out_dir / 'stone_auto_word_formal.pdf'

            self.assertTrue(pdf_path.is_file())
            self.assertTrue(docx_path.is_file())
            self.assertGreater(pdf_path.stat().st_size, 1_000_000)
            self.assertGreater(docx_path.stat().st_size, 120_000)

            reader = PdfReader(str(pdf_path))
            pdf_text = '\n'.join(page.extract_text() or '' for page in reader.pages)
            docx_text, document = extract_docx_text(docx_path)

            self.assertGreaterEqual(len(reader.pages), 10)
            self.assertGreaterEqual(len(document.paragraphs), 100)
            self.assertGreaterEqual(len(document.tables), 20)
            self.assertGreaterEqual(len(pdf_text), 8_000)
            self.assertGreaterEqual(len(docx_text), 8_000)

            for needle in (
                '送審速覽',
                '設計者註記',
                'Golden replay：已核對控制強度與細部檢核。',
                '石材外牆固定構件',
                expected['governing_check'],
            ):
                self.assertIn(needle, pdf_text)
                self.assertIn(needle, docx_text)

            for text in (pdf_text, docx_text):
                self.assertNotIn('草稿 DRAFT', text)
                for needle in PAGE_ONLY_REPORT_STATUS_NEEDLES:
                    self.assertNotIn(needle, text)

            audit_path = server.write_export_audit(
                payload,
                docx_path,
                'auto_word',
                result_source='frontend_results',
            )
            report = json.loads(audit_path.read_text(encoding='utf-8'))
            self.assertAuditSchema(report)
            self.assertEqual(report['mode'], 'auto_word')
            self.assertEqual(report['output']['size_bytes'], docx_path.stat().st_size)
            self.assertEqual(report['trace']['result_source'], 'frontend_results')
            self.assertTrue(report['review']['server_evaluation']['tool_html_match'])
            self.assertNotIn('priority_report_reading_status', report['review'])
            self.assertEqual(report['trace']['payload_sha256'], server.payload_sha256(payload))
            self.assertEqual(report['trace']['input_hash'], payload['meta']['input_hash'])
            self.assertEqual(report['trace']['result_hash'], payload['meta']['result_hash'])
            self.assertEqual(report['trace']['calc_source_hash'], payload['meta']['calc_source_hash'])
            for needle in PAGE_ONLY_REPORT_STATUS_NEEDLES:
                self.assertNotIn(needle, report_text(report))

            result_reconciliation = {
                'schemaVersion': 1,
                'strategy': 'stone-golden-replay-to-pdf-docx-hash',
                'caseId': GOLDEN_CASE_ID,
                'goldenCaseSha256': file_sha256(GOLDEN_CASE_PATH),
                'goldenInputSha256': file_sha256(GOLDEN_INPUT_PATH),
                'sourcePayloadSha256': report['trace']['payload_sha256'],
                'inputHash': report['trace']['input_hash'],
                'resultHash': report['trace']['result_hash'],
                'calcSourceHash': report['trace']['calc_source_hash'],
                'verifiedAssertionCount': 6,
                'pdfSha256': file_sha256(pdf_path),
                'docxSha256': file_sha256(docx_path),
                'auditSha256': file_sha256(audit_path),
                'pass': True,
            }

            evidence_summary = persist_rendered_evidence(
                pdf_path,
                docx_path,
                audit_path,
                page_count=len(reader.pages),
                pdf_text_length=len(pdf_text),
                docx_text_length=len(docx_text),
                result_reconciliation=result_reconciliation,
            )
            if rendered_evidence_dir() is not None:
                self.assertIsNotNone(evidence_summary)
                self.assertTrue(evidence_summary.is_file())


if __name__ == '__main__':
    runner = unittest.TextTestRunner(stream=sys.stdout, verbosity=2)
    unittest.main(testRunner=runner)
