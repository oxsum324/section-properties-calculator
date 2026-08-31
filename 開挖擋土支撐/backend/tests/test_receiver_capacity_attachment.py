from __future__ import annotations

from copy import deepcopy
from datetime import datetime, timezone
import hashlib
from pathlib import Path
from tempfile import TemporaryDirectory
from types import SimpleNamespace
import unittest
from unittest.mock import patch
from xml.etree import ElementTree
from zipfile import ZipFile

from docx import Document
from fastapi.testclient import TestClient
from pypdf import PdfReader

from backend.app.main import app
from backend.app.receiver_capacity import calculate_reshore_member_capacity
from backend.app.receiver_capacity_attachment import (
    ATTACHMENT_KIND,
    attachment_download_name_allowed,
    build_reshore_capacity_attachment,
)
from backend.tests import test_receiver_capacity as receiver_capacity_tests


W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def document_text(path: Path) -> str:
    document = Document(path)
    values = [paragraph.text for paragraph in document.paragraphs]
    values.extend(
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    )
    return "\n".join(values)


class ReshoreCapacityAttachmentTests(unittest.TestCase):
    def setUp(self) -> None:
        fixture = receiver_capacity_tests.ReshoreMemberCapacityTests()
        self.handoff = fixture.prepared_handoff()
        self.transfer_id = self.handoff["transfers"][0]["transferId"]
        self.input = fixture.valid_input(**fixture.valid_end_bearing_overrides())
        self.response = calculate_reshore_member_capacity(
            self.handoff,
            self.transfer_id,
            self.input,
            generated_at="2026-09-01T01:02:03Z",
        )

    def build(self, directory: Path, report_kind: str, response: dict | None = None):
        with patch(
            "backend.app.receiver_capacity_attachment.get_settings",
            return_value=SimpleNamespace(reports_dir=directory),
        ):
            return build_reshore_capacity_attachment(
                response or self.response,
                report_kind=report_kind,
                output_at=datetime(2026, 9, 1, 1, 3, 4, tzinfo=timezone.utc),
            )

    def api_payload(self, *, response: dict | None = None, report_kind: str = "docx") -> dict:
        return {
            "handoff": self.handoff,
            "transfer_id": self.transfer_id,
            "calculation_input": self.input.model_dump(),
            "calculation_response": response or self.response,
            "report_kind": report_kind,
            "formal_attachment_confirmed": True,
        }

    def test_builds_verifiable_pdf_and_compact_reference_docx(self) -> None:
        with TemporaryDirectory() as temporary:
            directory = Path(temporary)
            pdf = self.build(directory, "pdf")
            docx = self.build(directory, "docx")

            for artifact in (pdf, docx):
                self.assertTrue(artifact.path.is_file())
                self.assertEqual(artifact.attachment_size_bytes, artifact.path.stat().st_size)
                self.assertEqual(artifact.attachment_file_sha256, sha256(artifact.path))
                self.assertEqual(artifact.document_status, "formal-attachment")
                self.assertEqual(
                    artifact.rsc_calculation_fingerprint,
                    self.response["calculation"]["calculationFingerprint"],
                )
                self.assertEqual(
                    artifact.rsb_calculation_fingerprint,
                    self.response["bearingEvidence"]["documentReference"],
                )

            pdf_text = "\n".join(page.extract_text() or "" for page in PdfReader(pdf.path).pages)
            self._assert_pdf_embeds_font(pdf.path)
            word_text = document_text(docx.path)
            for value in (
                self.response["calculation"]["calculationFingerprint"],
                self.response["bearingEvidence"]["documentReference"],
                self.response["evidence"]["fileSha256"],
                "0.298816",
                "0.717457",
                "0.223352",
                "PASS",
                self.response["bearingEvidence"]["fileName"],
                "https://ej.aisc.org/index.php/engj/article/view/214",
            ):
                self.assertIn(value, pdf_text)
                self.assertIn(value, word_text)

            self._assert_docx_preset(docx.path)

    def _assert_pdf_embeds_font(self, path: Path) -> None:
        reader = PdfReader(path)
        embedded = False
        font_names: set[str] = set()
        for page in reader.pages:
            fonts = page["/Resources"].get("/Font", {}).get_object()
            for font_reference in fonts.values():
                font = font_reference.get_object()
                font_names.add(str(font.get("/BaseFont", "")))
                descriptors = []
                descriptor = font.get("/FontDescriptor")
                if descriptor is not None:
                    descriptors.append(descriptor.get_object())
                descendants = font.get("/DescendantFonts")
                if descendants is not None:
                    for descendant in descendants.get_object():
                        child_descriptor = descendant.get_object().get("/FontDescriptor")
                        if child_descriptor is not None:
                            descriptors.append(child_descriptor.get_object())
                embedded = embedded or any(
                    any(key in descriptor for key in ("/FontFile", "/FontFile2", "/FontFile3"))
                    for descriptor in descriptors
                )
        self.assertTrue(embedded, f"PDF must embed its CJK font; fonts={sorted(font_names)}")
        self.assertTrue(all("STSong" not in name for name in font_names))

    def _assert_docx_preset(self, path: Path) -> None:
        with ZipFile(path) as archive:
            document_root = ElementTree.fromstring(archive.read("word/document.xml"))
            styles_root = ElementTree.fromstring(archive.read("word/styles.xml"))
            numbering_root = ElementTree.fromstring(archive.read("word/numbering.xml"))

        self.assertEqual(document_root.findall(f".//{W}pBdr"), [])
        section = document_root.find(f".//{W}sectPr")
        self.assertIsNotNone(section)
        self.assertEqual(section.find(f"{W}pgSz").attrib[f"{W}w"], "12240")
        self.assertEqual(section.find(f"{W}pgSz").attrib[f"{W}h"], "15840")
        margins = section.find(f"{W}pgMar").attrib
        for key in ("top", "bottom", "left", "right"):
            self.assertEqual(margins[f"{W}{key}"], "1440")
        self.assertEqual(margins[f"{W}header"], "708")
        self.assertEqual(margins[f"{W}footer"], "708")

        normal = next(
            style
            for style in styles_root.findall(f"{W}style")
            if style.attrib.get(f"{W}styleId") == "Normal"
        )
        self.assertEqual(normal.find(f"{W}rPr/{W}rFonts").attrib[f"{W}ascii"], "Calibri")
        self.assertEqual(normal.find(f"{W}rPr/{W}sz").attrib[f"{W}val"], "22")
        self.assertEqual(normal.find(f"{W}pPr/{W}spacing").attrib[f"{W}line"], "300")

        tables = document_root.findall(f".//{W}tbl")
        self.assertGreaterEqual(len(tables), 6)
        for table in tables:
            properties = table.find(f"{W}tblPr")
            self.assertEqual(properties.find(f"{W}tblW").attrib[f"{W}w"], "9360")
            self.assertEqual(properties.find(f"{W}tblInd").attrib[f"{W}w"], "120")
            grid_widths = [
                int(column.attrib[f"{W}w"])
                for column in table.findall(f"{W}tblGrid/{W}gridCol")
            ]
            self.assertEqual(sum(grid_widths), 9360)
            for row in table.findall(f"{W}tr"):
                self.assertIsNotNone(row.find(f"{W}trPr/{W}cantSplit"))
                widths = [
                    int(cell.find(f"{W}tcPr/{W}tcW").attrib[f"{W}w"])
                    for cell in row.findall(f"{W}tc")
                ]
                self.assertEqual(widths, grid_widths)
                for cell in row.findall(f"{W}tc"):
                    margins = cell.find(f"{W}tcPr/{W}tcMar")
                    self.assertEqual(margins.find(f"{W}top").attrib[f"{W}w"], "80")
                    self.assertEqual(margins.find(f"{W}bottom").attrib[f"{W}w"], "80")
                    self.assertEqual(margins.find(f"{W}left").attrib[f"{W}w"], "120")
                    self.assertEqual(margins.find(f"{W}right").attrib[f"{W}w"], "120")

        fills = [node.attrib.get(f"{W}fill") for node in document_root.findall(f".//{W}shd")]
        self.assertIn("E8EEF5", fills)
        bullet_indents = [
            level.find(f"{W}pPr/{W}ind").attrib
            for level in numbering_root.findall(f".//{W}lvl")
            if level.find(f"{W}numFmt") is not None
            and level.find(f"{W}numFmt").attrib.get(f"{W}val") == "bullet"
            and level.find(f"{W}pPr/{W}ind") is not None
        ]
        self.assertTrue(
            any(
                indent.get(f"{W}left") == "540"
                and indent.get(f"{W}hanging") == "270"
                for indent in bullet_indents
            )
        )

    def test_ng_remains_finite_and_visible(self) -> None:
        weak_input = self.input.model_copy(
            update={"bottom_end_plate_thickness_cm": 1.9}
        )
        response = calculate_reshore_member_capacity(
            self.handoff,
            self.transfer_id,
            weak_input,
            generated_at="2026-09-01T01:02:04Z",
        )
        end = response["calculation"]["results"]["endBearing"]["bottom"]
        self.assertEqual(end["status"], "failed")
        self.assertGreater(end["plateBendingUtilizationRatio"], 1)
        with TemporaryDirectory() as temporary:
            artifact = self.build(Path(temporary), "docx", response)
            text = document_text(artifact.path)
            self.assertIn("NG", text)
            self.assertIn(str(end["plateBendingUtilizationRatio"]), text)

    def test_api_replays_exact_existing_response_and_rejects_tamper(self) -> None:
        client = TestClient(app)
        with TemporaryDirectory() as temporary, patch(
            "backend.app.receiver_capacity_attachment.get_settings",
            return_value=SimpleNamespace(reports_dir=Path(temporary)),
        ):
            result = client.post(
                "/api/removal-transfer/reshore-member-capacity/attachment",
                json=self.api_payload(),
            )
            self.assertEqual(result.status_code, 200, result.text)
            payload = result.json()
            self.assertEqual(payload["attachment_kind"], ATTACHMENT_KIND)
            self.assertTrue(payload["exact_calculation_response_verified"])
            self.assertEqual(payload["adopted_calculation_generated_at"], "2026-09-01T01:02:03Z")
            self.assertEqual(payload["rsc_evidence_file_sha256"], self.response["evidence"]["fileSha256"])
            self.assertIsNone(payload["canonical_evidence_url"])
            self.assertIsNone(payload["formal_source_bundle_url"])

            tampered = deepcopy(self.response)
            tampered["evidence"]["fileSha256"] = "0" * 64
            rejected = client.post(
                "/api/removal-transfer/reshore-member-capacity/attachment",
                json=self.api_payload(response=tampered),
            )
            self.assertEqual(rejected.status_code, 400, rejected.text)
            self.assertIn("精確重播", rejected.json()["detail"])

    def test_pdf_packaging_failure_removes_artifact_and_partial_evidence(self) -> None:
        def fail_after_partial_evidence(pdf_path: Path) -> Path:
            partial = pdf_path.with_name(f"{pdf_path.stem}.canonical-render.evidence.json")
            partial.write_text("partial", encoding="utf-8")
            raise RuntimeError("fixture render failure")

        client = TestClient(app)
        with TemporaryDirectory() as temporary, patch(
            "backend.app.receiver_capacity_attachment.get_settings",
            return_value=SimpleNamespace(reports_dir=Path(temporary)),
        ), patch(
            "backend.app.main.build_pdf_canonical_render_evidence",
            side_effect=fail_after_partial_evidence,
        ):
            result = client.post(
                "/api/removal-transfer/reshore-member-capacity/attachment",
                json=self.api_payload(report_kind="pdf"),
            )
            self.assertEqual(result.status_code, 500, result.text)
            self.assertEqual(list(Path(temporary).iterdir()), [])

    def test_non_bearing_response_and_unrelated_download_names_are_rejected(self) -> None:
        fixture = receiver_capacity_tests.ReshoreMemberCapacityTests()
        input_without_bearing = fixture.valid_input()
        response = calculate_reshore_member_capacity(
            self.handoff,
            self.transfer_id,
            input_without_bearing,
            generated_at="2026-09-01T01:02:05Z",
        )
        with TemporaryDirectory() as temporary:
            with self.assertRaisesRegex(ValueError, "只有啟用"):
                self.build(Path(temporary), "pdf", response)
            self.assertEqual(list(Path(temporary).iterdir()), [])

        expected_name = (
            "rsc-v4-rsb-" + "A" * 20 + "-" + "B" * 20 + "-20260901010203000000.pdf"
        )
        self.assertTrue(attachment_download_name_allowed(expected_name))
        self.assertTrue(
            attachment_download_name_allowed(
                expected_name.removesuffix(".pdf") + ".canonical-render.evidence.json"
            )
        )
        self.assertFalse(attachment_download_name_allowed("../" + expected_name))
        self.assertFalse(attachment_download_name_allowed("excavation-report.pdf"))


if __name__ == "__main__":
    unittest.main()
