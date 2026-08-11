from __future__ import annotations

import json
import re
import unittest
from datetime import datetime, timezone
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from pypdf import PdfReader

from backend.app.calculations import calculate_project
from backend.app.reporting import (
    _analysis_mapping_lines,
    _concise_metric_text,
    build_report,
    build_word_report,
    calculation_fingerprint,
)
from backend.app.schemas import AnalysisForceCase
from backend.app.workbook_loader import load_default_project
from backend.tests.handoff_fixtures import make_stage_adoption, make_verified_handoff_source


CALCULATION_BOOK_CONTENT_BOUNDARY = json.loads(
    (Path(__file__).resolve().parents[3] / "結構工具箱" / "tools" / "calculation-book-content-boundary.json").read_text(encoding="utf-8")
)
PAGE_ONLY_REPORT_STATUS_NEEDLES = tuple(dict.fromkeys(
    needle
    for category in CALCULATION_BOOK_CONTENT_BOUNDARY["forbiddenCategories"].values()
    for needle in category
))


class ReportingTests(unittest.TestCase):
    _default_word_artifact: dict[str, object] | None = None

    @classmethod
    def tearDownClass(cls) -> None:
        artifact = cls._default_word_artifact or {}
        report_path = artifact.get("path")
        if report_path is not None:
            report_path.unlink(missing_ok=True)

    @classmethod
    def default_word_artifact(cls) -> dict[str, object]:
        if cls._default_word_artifact is None:
            project = load_default_project().model_copy(deep=True)
            project.calculation_results = calculate_project(project)
            report_path = build_word_report(project, concise_mode=False)
            document = Document(str(report_path))
            text = "\n".join(paragraph.text for paragraph in document.paragraphs)
            table_text = "\n".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)
            combined_text = text + "\n" + table_text
            section = document.sections[0]
            header_text = "\n".join(paragraph.text for paragraph in section.header.paragraphs)
            footer_text = "\n".join(paragraph.text for paragraph in section.footer.paragraphs)
            first_page_header_text = "\n".join(paragraph.text for paragraph in section.first_page_header.paragraphs)
            first_page_footer_text = "\n".join(paragraph.text for paragraph in section.first_page_footer.paragraphs)
            with ZipFile(report_path) as archive:
                archive_names = archive.namelist()
                document_xml = archive.read("word/document.xml").decode("utf-8")
                footer_files = [name for name in archive.namelist() if name.startswith("word/footer")]
                footer_xml = "\n".join(archive.read(name).decode("utf-8") for name in footer_files)
            cls._default_word_artifact = {
                "project": project,
                "path": report_path,
                "document": document,
                "text": text,
                "table_text": table_text,
                "combined_text": combined_text,
                "section": section,
                "header_text": header_text,
                "footer_text": footer_text,
                "first_page_header_text": first_page_header_text,
                "first_page_footer_text": first_page_footer_text,
                "document_xml": document_xml,
                "footer_xml": footer_xml,
                "section_heading_count": len(re.findall(r"[一二三四五六七八九十]+、", combined_text)),
                "page_break_count": document_xml.count('w:type="page"'),
                "drawing_count": document_xml.count("<w:drawing"),
                "media_count": len([name for name in archive_names if name.startswith("word/media/")]),
            }
        return cls._default_word_artifact

    def test_word_report_is_generated(self) -> None:
        report_path = self.default_word_artifact()["path"]

        self.assertTrue(report_path.exists())
        self.assertEqual(report_path.suffix.lower(), ".docx")
        self.assertGreater(report_path.stat().st_size, 0)

    def test_word_report_has_substantial_attachment_structure(self) -> None:
        artifact = self.default_word_artifact()
        document = artifact["document"]

        self.assertGreater(len(artifact["combined_text"]), 25000)
        self.assertGreaterEqual(len(document.paragraphs), 500)
        self.assertGreaterEqual(len(document.tables), 10)
        self.assertGreaterEqual(artifact["section_heading_count"], 8)
        self.assertGreaterEqual(artifact["page_break_count"], 2)
        self.assertGreaterEqual(artifact["drawing_count"], 1)
        self.assertGreaterEqual(artifact["media_count"], 1)

    def test_word_report_includes_cover_title_and_project_metadata(self) -> None:
        combined_text = self.default_word_artifact()["combined_text"]

        self.assertIn("擋土支撐檢核計算書", combined_text)
        self.assertIn("文件狀態：內部審閱", combined_text)
        self.assertIn("工程名稱", combined_text)
        self.assertIn("工程位置", combined_text)
        self.assertIn("產出工具\n開挖擋土支撐計算書", combined_text)
        self.assertIn("工具版本\nv0.1.0", combined_text)
        self.assertIn("輸出時間", combined_text)
        self.assertNotIn("核可時間", combined_text)
        self.assertIn("計算指紋", combined_text)
        self.assertIn(calculation_fingerprint(self.default_word_artifact()["project"]), combined_text)
        self.assertIn("內容提要", combined_text)
        self.assertIn("報告說明", combined_text)
        self.assertIn("六、結構計算結果", combined_text)

    def test_word_report_excludes_page_only_status_overview(self) -> None:
        combined_text = self.default_word_artifact()["combined_text"]

        for needle in PAGE_ONLY_REPORT_STATUS_NEEDLES:
            self.assertNotIn(needle, combined_text)

    def test_pdf_report_uses_formal_section_structure(self) -> None:
        project = load_default_project().model_copy(deep=True)
        project.calculation_results = calculate_project(project)

        report_path = build_report(project)
        reader = PdfReader(str(report_path))
        first_page_text = reader.pages[0].extract_text() or ""
        page_texts = [page.extract_text() or "" for page in reader.pages]
        full_text = "\n".join(page_texts)
        text = "\n".join(page.extract_text() or "" for page in reader.pages[:4])

        self.assertGreaterEqual(len(reader.pages), 20)
        self.assertGreater(len(full_text), 25000)
        self.assertEqual(sum(1 for page_text in page_texts if len(page_text.strip()) > 100), len(reader.pages))
        self.assertGreaterEqual(len(re.findall(r"[一二三四五六七八九十]+、", full_text)), 8)
        self.assertIn("一、摘要", text)
        self.assertIn("二、設計依據", text)
        self.assertIn("三、結構分析使用之電腦程式", text)
        self.assertIn("四、材料性質", text)
        self.assertIn("五、輸入基本資料", text)
        self.assertIn("六、結構計算結果", text)
        self.assertIn("上、下層支撐資料均採手動輸入", text)
        self.assertIn("內容提要", text)
        self.assertIn("附件一", text)
        self.assertIn("主要控制項目彙整", text)
        self.assertIn("計算指紋", text)
        self.assertIn(calculation_fingerprint(project), text)
        self.assertEqual(first_page_text.count("擋土支撐檢核計算書"), 1)
        self.assertIn("文件狀態：內部審閱", first_page_text)
        report_path.unlink(missing_ok=True)

    def test_approved_reports_are_formal_attachments_even_with_blank_optional_identity(self) -> None:
        project = load_default_project().model_copy(deep=True)
        project.metadata.name = ""
        project.metadata.designer = ""
        project.metadata.project_code = "EXC-APPROVED-001"
        project.calculation_results = calculate_project(project)
        output_at = datetime(2026, 8, 8, 15, 30, 0, tzinfo=timezone.utc)

        word_path = build_word_report(project, approved=True, output_at=output_at)
        pdf_path = build_report(project, approved=True, output_at=output_at)
        try:
            document = Document(str(word_path))
            word_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
            word_table_text = "\n".join(
                cell.text for table in document.tables for row in table.rows for cell in row.cells
            )
            combined_word_text = word_text + "\n" + word_table_text
            pdf_text = "\n".join(page.extract_text() or "" for page in PdfReader(str(pdf_path)).pages)

            for text in (combined_word_text, pdf_text):
                self.assertIn("文件狀態：正式附件", text)
                self.assertIn("產出工具", text)
                self.assertIn("開挖擋土支撐計算書", text)
                self.assertIn("工具版本", text)
                self.assertIn("v0.1.0", text)
                self.assertIn("輸出時間", text)
                self.assertIn("2026-08-08T15:30:00+00:00", text)
                self.assertIn("核可時間", text)
                self.assertNotIn("文件狀態：內部審閱", text)
            self.assertIn("本計算書係就『本案』", combined_word_text)
            self.assertIn("計畫編號\nEXC-APPROVED-001", combined_word_text)
        finally:
            word_path.unlink(missing_ok=True)
            pdf_path.unlink(missing_ok=True)

    def test_pdf_report_excludes_page_only_status_overview(self) -> None:
        project = load_default_project().model_copy(deep=True)
        project.calculation_results = calculate_project(project)

        report_path = build_report(project)
        reader = PdfReader(str(report_path))
        text = "\n".join(page.extract_text() or "" for page in reader.pages)

        for needle in PAGE_ONLY_REPORT_STATUS_NEEDLES:
            self.assertNotIn(needle, text)
        report_path.unlink(missing_ok=True)

    def test_word_report_adds_header_footer_page_number_and_appendix_page_breaks(self) -> None:
        artifact = self.default_word_artifact()
        project = artifact["project"]
        section = artifact["section"]
        header_text = artifact["header_text"]
        footer_text = artifact["footer_text"]
        first_page_header_text = artifact["first_page_header_text"]
        first_page_footer_text = artifact["first_page_footer_text"]

        self.assertTrue(section.different_first_page_header_footer)
        self.assertIn("擋土支撐檢核計算書", header_text)
        self.assertIn(project.metadata.name, header_text)
        self.assertEqual(first_page_header_text.strip(), "")
        self.assertIn("第", footer_text)
        self.assertIn("頁", footer_text)
        self.assertIn("第", first_page_footer_text)
        self.assertIn("頁", first_page_footer_text)

        self.assertIn('w:type="page"', artifact["document_xml"])
        self.assertIn("PAGE", artifact["footer_xml"])

    def test_single_side_word_report_uses_generic_source_label(self) -> None:
        project = load_default_project().model_copy(deep=True)
        project.calculation_options.include_bottom_supports = False
        project.calculation_results = calculate_project(project)

        report_path = build_word_report(project)
        document = Document(str(report_path))
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        table_text = "\n".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)

        self.assertIn("支撐來源", text + "\n" + table_text)
        self.assertNotIn("上層來源", table_text)
        report_path.unlink(missing_ok=True)

    def test_word_report_summary_marks_section_names(self) -> None:
        artifact = self.default_word_artifact()
        project = artifact["project"]
        expected_section = next(
            (item.section_name for item in project.calculation_results.summary if item.section_name),
            "",
        )

        table_text = artifact["table_text"]

        self.assertTrue(expected_section)
        self.assertIn("型號：", table_text)
        self.assertIn(expected_section, table_text)

    def test_word_report_includes_appendix_sections_with_detail_checks(self) -> None:
        artifact = self.default_word_artifact()
        project = artifact["project"]
        support_label = project.calculation_results.support_checks[0].label
        column_label = project.calculation_results.column_checks[0].label

        combined_text = artifact["combined_text"]

        self.assertIn("附件一：支撐、橫擋、斜撐、大角撐細部計算結果", combined_text)
        self.assertIn("附件二：中間柱、共構柱細部計算結果", combined_text)
        self.assertIn("一、摘要", combined_text)
        self.assertIn("六、結構計算結果", combined_text)
        self.assertIn("內容提要", combined_text)
        self.assertIn("報告說明", combined_text)
        self.assertIn("設計規範與檢核依據", combined_text)
        self.assertIn("附件一", combined_text)
        self.assertIn("細部計算過程與檢核結果", combined_text)
        self.assertIn("附件一共用參數", combined_text)
        self.assertIn("附件二共用參數", combined_text)
        self.assertIn("附件一型鋼彙整表", combined_text)
        self.assertIn("附件二型鋼彙整表", combined_text)
        self.assertIn("支撐示意圖", combined_text)
        self.assertIn("檢核結論", combined_text)
        self.assertIn("納入檢討構件", combined_text)
        self.assertIn("整體判定", combined_text)
        self.assertIn("最不利構件", combined_text)
        self.assertIn("控制條件", combined_text)
        self.assertIn("最大利用率", combined_text)
        self.assertIn("應注意", combined_text)
        self.assertIn("支撐計算如下", combined_text)
        self.assertIn("檢核結果統計", combined_text)
        self.assertIn("橫檔計算如下", combined_text)
        self.assertIn("斜撐計算如下", combined_text)
        self.assertIn("大角撐計算如下", combined_text)
        self.assertLess(combined_text.find("附件一型鋼彙整表"), combined_text.find("支撐示意圖"))
        self.assertLess(combined_text.find("支撐示意圖"), combined_text.find("支撐計算如下"))
        self.assertIn(support_label, combined_text)
        self.assertIn(column_label, combined_text)
        self.assertIn("已知條件", combined_text)
        self.assertIn("斷面資料", combined_text)
        self.assertIn("檢核公式", combined_text)
        self.assertIn("代入計算", combined_text)
        self.assertIn("Fa 依 KL/r 與 Cc 之關係", combined_text)
        self.assertIn("本案 KL/r =", combined_text)
        self.assertIn("詳前述「附件一型鋼彙整表」", combined_text)
        self.assertIn("詳前述「附件二型鋼彙整表」", combined_text)
        self.assertIn("本節檢核摘要", combined_text)
        self.assertIn("構件名稱", combined_text)
        self.assertIn("判定統計", combined_text)
        self.assertIn("編排原則", combined_text)
        self.assertIn("Qc,allow", combined_text)
        self.assertIn("Qt,allow", combined_text)
        self.assertIn("壓入比 =", combined_text)
        self.assertIn("拉拔比 =", combined_text)
        self.assertIn("積載重 WL", combined_text)
        self.assertIn("構件折減係數 αs / αw / αb", combined_text)
        self.assertIn("柱構件折減係數 αp", combined_text)
        self.assertIn("上、下層支撐資料均採手動輸入", combined_text)
        self.assertIn("本案未匯入外部分析檔，支撐軸力、土層及施工階段資料均以人工整理後之輸入內容作為後續構件檢核之依據。", combined_text)
        self.assertIn("上層第 1 層", combined_text)
        self.assertIn("下層第 1 層", combined_text)
        self.assertLess(combined_text.find("上層第 1 層"), combined_text.find("下層第 1 層"))
        self.assertNotIn("附件一彙整表", combined_text)
        self.assertNotIn("附件二彙整表", combined_text)
        self.assertNotIn("條文索引 / 符號說明", combined_text)
        self.assertNotIn("利用率 0.", combined_text)
        self.assertNotIn("引用參考小工具", combined_text)
        self.assertNotIn("allowable_axial_stress", combined_text)
        self.assertNotIn("allowable_fbx", combined_text)
        self.assertNotIn("採用型號：", combined_text)
        self.assertNotIn("若 KL/r < Cc", combined_text)
        self.assertNotIn("若 KL/r >=", combined_text)
        self.assertNotIn("控制值 / 允許值 =", combined_text)
        self.assertNotIn("橫擋 / 斜撐示意圖", combined_text)
        self.assertNotIn("本節適用規範", combined_text)
        self.assertNotIn("Say~OK", combined_text)

    def test_word_report_concise_mode_summarizes_following_checks(self) -> None:
        project = load_default_project().model_copy(deep=True)
        project.calculation_results = calculate_project(project)

        report_path = build_word_report(project, concise_mode=True)
        document = Document(str(report_path))
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        table_text = "\n".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)
        combined_text = text + "\n" + table_text

        self.assertIn("本節檢核摘要", combined_text)
        self.assertIn("首筆列示完整驗算過程，其餘摘列關鍵值與判定", combined_text)
        self.assertIn("關鍵數值", combined_text)
        self.assertIn("檢核公式：", combined_text)
        self.assertLess(combined_text.count("檢核公式："), 10)
        self.assertIn("-concise-", report_path.name)
        report_path.unlink(missing_ok=True)

    def test_word_report_detailed_mode_keeps_full_detail(self) -> None:
        artifact = self.default_word_artifact()
        report_path = artifact["path"]
        combined_text = artifact["combined_text"]

        self.assertNotIn("本節採簡述版：首筆詳算，其餘僅列關鍵值與檢核結果。", combined_text)
        self.assertNotIn("其餘項目關鍵摘要", combined_text)
        self.assertGreater(combined_text.count("檢核公式："), 20)
        self.assertIn("附件編排方式", combined_text)
        self.assertIn("詳細版（逐筆列示完整驗算過程）", combined_text)
        self.assertIn("-detail-", report_path.name)

    def test_word_report_preserves_verified_decking_handoff_trace(self) -> None:
        project = load_default_project().model_copy(deep=True)
        for column in project.columns:
            column.enabled = column.variant == "composite_crane"
        column = next(item for item in project.columns if item.variant == "composite_crane")
        column.construction_stage_load_t = 64.32
        column.construction_stage_load_source = make_verified_handoff_source(64.32)
        project.calculation_results = calculate_project(project)

        report_path = build_word_report(project, concise_mode=False)
        document = Document(str(report_path))
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        table_text = "\n".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)
        combined_text = text + "\n" + table_text

        self.assertIn("Np = 64.32 tf", combined_text)
        self.assertIn("覆工板系統計算工具 v1.0", combined_text)
        self.assertIn("CF-0123456789ABCDEF", combined_text)
        self.assertIn(column.construction_stage_load_source.handoff_fingerprint, combined_text)
        self.assertIn("控制工況 = Pu1", combined_text)
        report_path.unlink(missing_ok=True)

    def test_word_report_lists_multi_stage_envelope_and_separate_controls(self) -> None:
        project = load_default_project().model_copy(deep=True)
        for column in project.columns:
            column.enabled = column.variant == "composite_crane"
        column = next(item for item in project.columns if item.variant == "composite_crane")
        initial = make_stage_adoption(
            column.column_id,
            "構台初設",
            40.0,
            eccentricity_x_m=0.75,
            eccentricity_y_m=-0.25,
            transfer_basis="施工配置圖 A-03",
        )
        crane = make_stage_adoption(column.column_id, "吊車作業", 80.0)
        column.construction_stage_load_t = 80.0
        column.construction_stage_load_source = crane.source
        column.construction_stage_loads = [initial, crane]
        project.calculation_results = calculate_project(project)

        report_path = build_word_report(project, concise_mode=False)
        document = Document(str(report_path))
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        table_text = "\n".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)
        combined_text = text + "\n" + table_text

        self.assertIn("施工階段包絡 = 3 案", combined_text)
        self.assertIn("施工階段逐案包絡結果如下", combined_text)
        self.assertIn("構台初設", combined_text)
        self.assertIn("吊車作業", combined_text)
        self.assertIn("ΔMx / ΔMy = 30.00 / -10.00 tf-m", combined_text)
        self.assertIn("附加偏心 = 已採用", combined_text)
        self.assertIn("施工配置圖 A-03", combined_text)
        self.assertIn("柱交互作用控制階段 = 構台初設", combined_text)
        self.assertIn("拉拔力控制階段 = 無施工構台荷重基準案", combined_text)
        self.assertIn(initial.source.handoff_fingerprint, combined_text)
        self.assertIn(crane.source.handoff_fingerprint, combined_text)
        report_path.unlink(missing_ok=True)

    def test_word_report_lists_source_reaction_distribution_and_adopted_load(self) -> None:
        project = load_default_project().model_copy(deep=True)
        columns = [item for item in project.columns if item.variant in {"composite_normal", "composite_crane"}]
        for column in project.columns:
            column.enabled = column in columns
        for column, factor in zip(columns, (0.4, 0.6)):
            column.construction_stage_loads = [make_stage_adoption(
                column.column_id,
                "構台滿載",
                100.0,
                distribution_factor=factor,
                distribution_basis="構台反力分配圖 S-05",
            )]
        project.calculation_results = calculate_project(project)

        report_path = build_word_report(project, concise_mode=False)
        document = Document(str(report_path))
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        table_text = "\n".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)
        combined_text = text + "\n" + table_text

        self.assertIn("來源 Np = 100.00 tf，分配比例 = 40.00%，本柱 Np = 40.00 tf", combined_text)
        self.assertIn("來源 Np = 100.00 tf，分配比例 = 60.00%，本柱 Np = 60.00 tf", combined_text)
        self.assertIn("反力分配依據 = 構台反力分配圖 S-05", combined_text)
        report_path.unlink(missing_ok=True)

    def test_word_report_lists_adopted_analysis_stage_mapping_without_review_warning(self) -> None:
        project = load_default_project().model_copy(deep=True)
        row = project.top_supports[0]
        row.force_source = "analysis_import"
        row.analysis_stage_cases = [
            AnalysisForceCase(stage_index=1, stage_label="第一階開挖", axial_force_t=row.axial_force_t - 10.0),
            AnalysisForceCase(stage_index=2, stage_label="第二階開挖", axial_force_t=row.axial_force_t),
        ]
        row.analysis_install_stage_index = 1
        row.analysis_install_stage_label = "第一道支撐安裝"
        row.analysis_control_stage_index = 2
        row.analysis_control_stage_label = "第二階開挖"
        row.analysis_removal_stage_index = 3
        row.analysis_removal_stage_label = "第一道支撐拆除"
        row.removal_transfer_mode = "floor"
        row.removal_transfer_target = "B2F 樓版 S1 區"
        row.removal_transfer_direction = "沿第一道支撐軸線向東"
        row.removal_transfer_basis = "拆撐順序圖 CS-04 與樓版階段分析 ST-12"
        row.removal_transfer_confirmed = True
        row.construction_step_label = "第二階開挖完成、第二層支撐施作前"
        row.analysis_mapping_basis = "施工順序圖 CS-02 與分析階段表逐項核對"
        row.analysis_mapping_confirmed = True
        project.calculation_results = calculate_project(project)

        report_path = build_word_report(project, concise_mode=False)
        document = Document(str(report_path))
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        table_text = "\n".join(cell.text for table in document.tables for table_row in table.rows for cell in table_row.cells)
        combined_text = text + "\n" + table_text

        self.assertIn("控制分析階段 = #2 第二階開挖", combined_text)
        self.assertIn("分析安裝階段 = #1 第一道支撐安裝", combined_text)
        self.assertIn("分析拆撐階段 = #3 第一道支撐拆除", combined_text)
        self.assertIn("拆撐後荷重處置 = 移轉至樓版", combined_text)
        self.assertIn("承接構造 = B2F 樓版 S1 區", combined_text)
        self.assertIn("承接設計需求 =", combined_text)
        self.assertIn("傳力方向 = 沿第一道支撐軸線向東", combined_text)
        self.assertIn("承接分配 = B2F 樓版 S1 區｜100%", combined_text)
        self.assertIn("處置依據 = 拆撐順序圖 CS-04 與樓版階段分析 ST-12", combined_text)
        self.assertIn("實際施工步驟 = 第二階開挖完成、第二層支撐施作前", combined_text)
        self.assertIn("對應依據 = 施工順序圖 CS-02 與分析階段表逐項核對", combined_text)
        self.assertIn("施工階段軸力時序（2 階段）", combined_text)
        self.assertIn("逐階段需求比包絡 = #1 第一階開挖", combined_text)
        self.assertIn("需求比控制階段 = #2 第二階開挖", combined_text)
        self.assertNotIn("尚未確認控制分析階段", combined_text)
        self.assertNotIn("尚未指定拆撐後荷重處置", combined_text)
        self.assertNotIn("處置待確認", combined_text)
        self.assertNotIn("請重新套用分析候選值", combined_text)
        report_path.unlink(missing_ok=True)

    def test_pdf_report_renders_confirmed_removal_transfer_disposition(self) -> None:
        project = load_default_project().model_copy(deep=True)
        row = project.top_supports[0]
        row.force_source = "analysis_import"
        row.analysis_stage_cases = [
            AnalysisForceCase(stage_index=2, stage_label="第二階開挖", axial_force_t=row.axial_force_t),
        ]
        row.analysis_install_stage_index = 1
        row.analysis_install_stage_label = "第一道支撐安裝"
        row.analysis_control_stage_index = 2
        row.analysis_control_stage_label = "第二階開挖"
        row.analysis_removal_stage_index = 3
        row.analysis_removal_stage_label = "第一道支撐拆除"
        row.removal_transfer_mode = "outside_scope"
        row.removal_transfer_direction = "沿第一道支撐軸線向東"
        row.removal_transfer_basis = "拆撐順序圖 CS-04，承接構造另案檢核"
        row.removal_transfer_confirmed = True
        row.construction_step_label = "第二階開挖完成、第二層支撐施作前"
        row.analysis_mapping_basis = "施工順序圖 CS-02 與分析階段表逐項核對"
        row.analysis_mapping_confirmed = True
        project.calculation_results = calculate_project(project)

        report_path = build_report(project)
        reader = PdfReader(str(report_path))
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
        support_check = project.calculation_results.support_checks[0]
        mapping_text = "\n".join(_analysis_mapping_lines(support_check.inputs))

        # Exact Chinese content is asserted through the shared calculation inputs
        # and the DOCX test above.  The PDF's embedded CJK font is not mapped
        # reliably by pypdf, so keep this assertion to render evidence and a
        # stable ASCII marker instead of coupling it to text-extraction details.
        self.assertEqual(support_check.inputs["拆撐後荷重處置"], "本構件檢核範圍外（另案檢核）")
        self.assertEqual(support_check.inputs["拆撐處置依據"], "拆撐順序圖 CS-04，承接構造另案檢核")
        self.assertIn("拆撐後荷重處置 = 本構件檢核範圍外（另案檢核）", mapping_text)
        self.assertNotIn("承接構造 = —", mapping_text)
        self.assertGreaterEqual(len(reader.pages), 20)
        self.assertIn("CS-04", text)
        report_path.unlink(missing_ok=True)

    def test_word_report_hides_non_included_structural_modules(self) -> None:
        project = load_default_project().model_copy(deep=True)
        project.calculation_options.include_top_wales = False
        project.calculation_options.include_bottom_wales = False
        project.calculation_options.include_top_braces = False
        project.calculation_options.include_bottom_braces = False
        project.calculation_options.include_corner_braces = False
        project.calculation_results = calculate_project(project)

        report_path = build_word_report(project, concise_mode=False)
        document = Document(str(report_path))
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        table_text = "\n".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)
        combined_text = text + "\n" + table_text

        self.assertIn("附件一：支撐細部計算結果", combined_text)
        self.assertNotIn("附件一：支撐、橫擋", combined_text)
        self.assertNotIn("橫檔計算如下", combined_text)
        self.assertNotIn("斜撐計算如下", combined_text)
        self.assertNotIn("大角撐計算如下", combined_text)
        self.assertNotIn("層別\n水平支撐\n橫擋", table_text)
        self.assertNotIn("層別\n水平支撐\n斜撐", table_text)
        self.assertNotIn("層別\n水平支撐\n大角撐", table_text)
        report_path.unlink(missing_ok=True)

    def test_word_report_uses_compact_manual_analysis_rows(self) -> None:
        table_text = self.default_word_artifact()["table_text"]

        self.assertIn("資料說明", table_text)
        self.assertIn("本案未匯入外部分析檔，後續檢核以手動輸入資料為準。", table_text)
        self.assertNotIn("來源檔案\n—", table_text)
        self.assertNotIn("來源格式\n—", table_text)

    def test_concise_metric_text_uses_multiline_layout(self) -> None:
        project = load_default_project().model_copy(deep=True)
        project.calculation_results = calculate_project(project)

        text = _concise_metric_text(project.calculation_results.support_checks[0])

        self.assertIn("\n", text)
        self.assertIn("交互作用比 R =", text)


if __name__ == "__main__":
    unittest.main()
