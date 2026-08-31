from __future__ import annotations

import base64
import hashlib
import json
import re
import shutil
import sys
from datetime import datetime, timezone
from pathlib import Path
from types import SimpleNamespace
from zipfile import ZipFile

from docx import Document
from pypdf import PdfReader

from backend.app.calculations import calculate_project
from backend.app.pdf_render_evidence import (
    build_pdf_canonical_render_evidence,
    build_pdf_formal_source_bundle,
)
from backend.app.project_store import ProjectStore
from backend.app.receiver_capacity import calculate_reshore_member_capacity
from backend.app.receiver_capacity_attachment import build_reshore_capacity_attachment
from backend.app.removal_transfer_handoff import build_removal_transfer_handoff
from backend.app.reporting import build_report, build_word_report, calculation_fingerprint
from backend.app.schemas import AnalysisForceCase, ProjectState, ReshoreMemberCapacityInput
from backend.app.workbook_loader import load_default_project


CALCULATION_BOOK_CONTENT_BOUNDARY = json.loads(
    (Path(__file__).resolve().parents[3] / "結構工具箱" / "tools" / "calculation-book-content-boundary.json").read_text(encoding="utf-8")
)
PAGE_ONLY_REPORT_STATUS_NEEDLES = tuple(dict.fromkeys(
    needle
    for category in CALCULATION_BOOK_CONTENT_BOUNDARY["forbiddenCategories"].values()
    for needle in category
))

REQUIRED_REPORT_TEXT = (
    "擋土支撐檢核計算書",
    "正式放行擋土支撐範例",
    "一、摘要",
    "二、設計依據",
    "三、結構分析使用之電腦程式",
    "六、結構計算結果",
    "附件一",
    "附件二",
)


def require(pass_condition: bool, message: str) -> None:
    if not pass_condition:
        raise RuntimeError(message)


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


RESULT_GROUPS = (
    "support_checks",
    "wale_checks",
    "brace_checks",
    "corner_brace_checks",
    "column_checks",
)

CHECK_FIELDS = (
    "module_name",
    "label",
    "formula_id",
    "computed_value",
    "allowable_value",
    "utilization_ratio",
    "status",
    "controlling_condition",
)


def canonical_results(project: ProjectState) -> dict:
    require(project.calculation_results is not None, "project must contain calculation results")
    return project.calculation_results.model_dump(mode="json", exclude={"generated_at"})


def canonical_sha256(payload: dict) -> str:
    encoded = json.dumps(payload, ensure_ascii=False, sort_keys=True, separators=(",", ":")).encode("utf-8")
    return hashlib.sha256(encoded).hexdigest()


def verify_replayed_results(expected: dict, actual: dict) -> tuple[int, int, dict[str, int]]:
    assertion_count = 0
    check_count = 0
    group_counts: dict[str, int] = {}
    for group_name in RESULT_GROUPS:
        expected_checks = expected.get(group_name) or []
        actual_checks = actual.get(group_name) or []
        require(len(actual_checks) == len(expected_checks), f"replayed {group_name} check count must match")
        assertion_count += 1
        group_counts[group_name] = len(actual_checks)
        check_count += len(actual_checks)
        for index, (expected_check, actual_check) in enumerate(zip(expected_checks, actual_checks, strict=True)):
            for field in CHECK_FIELDS:
                require(
                    actual_check.get(field) == expected_check.get(field),
                    f"replayed {group_name}[{index}].{field} must match",
                )
                assertion_count += 1

    expected_summary = expected.get("summary") or []
    actual_summary = actual.get("summary") or []
    require(len(actual_summary) == len(expected_summary), "replayed summary count must match")
    assertion_count += 1
    summary_fields = ("group", "label", "section_name", "status", "utilization_ratio")
    for index, (expected_item, actual_item) in enumerate(zip(expected_summary, actual_summary, strict=True)):
        for field in summary_fields:
            require(actual_item.get(field) == expected_item.get(field), f"replayed summary[{index}].{field} must match")
            assertion_count += 1

    require(actual.get("warnings") == expected.get("warnings"), "replayed warnings must match")
    assertion_count += 1
    return check_count, assertion_count, group_counts


def build_receiver_capacity_attachment_fixture() -> dict:
    project = load_default_project().model_copy(deep=True)
    project.metadata.id = "release-rsc-v4-rsb-v1"
    project.metadata.name = "正式放行 RSC v4／RSB v1 接收端附件範例"
    row = project.top_supports[0]
    row.force_source = "analysis_import"
    row.analysis_stage_cases = [
        AnalysisForceCase(
            stage_index=2,
            stage_label="第二階開挖",
            axial_force_t=row.axial_force_t,
        )
    ]
    row.analysis_install_stage_index = 1
    row.analysis_install_stage_label = "第一道支撐安裝"
    row.analysis_control_stage_index = 2
    row.analysis_control_stage_label = "第二階開挖"
    row.analysis_removal_stage_index = 3
    row.analysis_removal_stage_label = "第一道支撐拆除"
    row.removal_transfer_mode = "reshore"
    row.removal_transfer_target = "B2F R1 重撐群"
    row.removal_transfer_direction = "沿原支撐軸線傳至 B2F"
    row.removal_transfer_basis = "拆撐順序圖 CS-04"
    row.removal_transfer_confirmed = True
    row.construction_step_label = "第二階開挖完成"
    row.analysis_mapping_basis = "施工順序圖 CS-02"
    row.analysis_mapping_confirmed = True
    project.calculation_results = calculate_project(project)
    handoff = build_removal_transfer_handoff(
        project,
        calculation_fingerprint(project),
        generated_at="2026-09-01T08:00:00Z",
    )
    calculation_input = ReshoreMemberCapacityInput(
        analysis_mode="pure_axial",
        section_name="RH300X300X10X15",
        member_count=2,
        unbraced_length_x_m=3.0,
        unbraced_length_y_m=3.0,
        effective_length_factor_kx=1.0,
        effective_length_factor_ky=1.0,
        fy_tf_per_cm2=2.5,
        e_tf_per_cm2=2040.0,
        allowable_stress_increase_factor=1.0,
        imbalance_factor=1.1,
        additional_axial_load_tf_per_member=1.0,
        governing_load_combination="拆撐階段 LC-RM-01",
        effective_length_basis="上下端依施工詳圖視為鉸接，Kx=Ky=1.0",
        load_distribution_basis="兩支平均分配並以 1.10 不均勻係數放大",
        additional_load_basis="單支自重及預壓附加軸力 1.0 tf",
        pure_axial_no_eccentricity_confirmed=True,
        end_bearing_mode="centered_rectangular_plate",
        top_end_plate_fy_tf_per_cm2=2.5,
        top_end_plate_width_cm=45.0,
        top_end_plate_length_cm=45.0,
        top_end_plate_thickness_cm=3.0,
        top_support_material="concrete_full_area",
        top_support_material_strength_tf_per_cm2=0.28,
        bottom_end_plate_width_cm=50.0,
        bottom_end_plate_length_cm=50.0,
        bottom_end_plate_thickness_cm=3.5,
        bottom_end_plate_fy_tf_per_cm2=2.5,
        bottom_support_material="finished_steel",
        bottom_support_material_strength_tf_per_cm2=2.5,
        bottom_support_true_plane_confirmed=True,
        end_bearing_configuration_basis="端板配置與密貼依施工詳圖 EB-01。",
        end_plate_bending_method_basis="專案指定採均佈反力懸臂板模型及 Fb=0.60Fy。",
        top_support_bearing_basis="混凝土試驗報告 FC-01，f'c=0.28 tf/cm²；全面積承壓。",
        bottom_support_bearing_basis="加工鋼支承詳圖 SB-01，Fy=2.5 tf/cm²。",
        centered_full_contact_end_bearing_confirmed=True,
        h_section_end_finished_confirmed=True,
        unperforated_unstiffened_single_plate_confirmed=True,
    )
    response = calculate_reshore_member_capacity(
        handoff,
        handoff["transfers"][0]["transferId"],
        calculation_input,
        generated_at="2026-09-01T09:00:00Z",
    )
    return response


def main() -> int:
    if len(sys.argv) != 2:
        raise SystemExit("usage: release_report_artifacts.py <output-dir>")

    output_dir = Path(sys.argv[1]).resolve()
    output_dir.mkdir(parents=True, exist_ok=True)

    source_project = load_default_project().model_copy(deep=True)
    source_project.metadata.id = "release-evidence"
    source_project.metadata.name = "正式放行擋土支撐範例"
    source_project.metadata.project_code = "EXCAVATION-RELEASE"
    source_project.metadata.location = "本地正式放行驗證"
    source_project.calculation_results = None

    source_project_path = output_dir / "excavation-project-state.json"
    source_project_path.write_text(
        json.dumps(source_project.model_dump(mode="json"), ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )
    reloaded_project = ProjectState.model_validate_json(source_project_path.read_text(encoding="utf-8"))
    require(reloaded_project.calculation_results is None, "preserved ProjectState must contain inputs only")

    baseline_project = source_project.model_copy(deep=True)
    baseline_project.calculation_results = calculate_project(baseline_project)
    reloaded_project.calculation_results = calculate_project(reloaded_project)
    expected_results = canonical_results(baseline_project)
    replayed_results = canonical_results(reloaded_project)
    require(replayed_results == expected_results, "ProjectState replay must reproduce the complete result payload")
    verified_check_count, verified_assertion_count, verified_group_counts = verify_replayed_results(
        expected_results,
        replayed_results,
    )
    require(verified_check_count == 47, "release ProjectState replay must verify all 47 component checks")
    project = reloaded_project
    result_sha256 = canonical_sha256(replayed_results)
    fingerprint = calculation_fingerprint(project)
    require(re.fullmatch(r"CF-[0-9A-F]{16}", fingerprint) is not None, "release calculation fingerprint format")

    generated_pdf: Path | None = None
    generated_docx: Path | None = None
    generated_attachment_paths: list[Path] = []
    try:
        generated_pdf = build_report(project, concise_mode=False, approved=True)
        generated_docx = build_word_report(project, concise_mode=False, approved=True)

        pdf_path = output_dir / "excavation-report.pdf"
        docx_path = output_dir / "excavation-report.docx"
        shutil.copy2(generated_pdf, pdf_path)
        shutil.copy2(generated_docx, docx_path)

        store = ProjectStore.__new__(ProjectStore)
        store.settings = SimpleNamespace(projects_dir=output_dir / "latest-download")
        latest_pdf_path = store.save_report(project.metadata.id, pdf_path)
        latest_docx_path = store.save_report(
            project.metadata.id,
            docx_path,
            latest_name="latest-report.docx",
        )

        reader = PdfReader(str(pdf_path))
        pdf_page_texts = [page.extract_text() or "" for page in reader.pages]
        pdf_text = "\n".join(pdf_page_texts)
        require(len(reader.pages) >= 20, "release PDF must keep at least 20 pages")
        require(len(pdf_text) > 25_000, "release PDF must contain substantial text")
        require(
            all(len(page_text.strip()) > 100 for page_text in pdf_page_texts),
            "release PDF must not contain empty or sparse text pages",
        )

        document = Document(str(docx_path))
        paragraph_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        table_text = "\n".join(
            cell.text
            for table in document.tables
            for row in table.rows
            for cell in row.cells
        )
        docx_text = paragraph_text + "\n" + table_text
        with ZipFile(docx_path) as archive:
            archive_names = archive.namelist()
            document_xml = archive.read("word/document.xml").decode("utf-8")

        section_heading_count = len(re.findall(r"[一二三四五六七八九十]+、", docx_text))
        document_xml_text = re.sub(r"<[^>]+>", "", document_xml)
        xml_paragraph_count = len(re.findall(r"<w:p(?:\s|>)", document_xml))
        xml_table_count = len(re.findall(r"<w:tbl(?:\s|>)", document_xml))
        xml_section_count = len(re.findall(r"[一二三四五六七八九十]+、", document_xml_text))
        page_break_count = document_xml.count('w:type="page"')
        drawing_count = document_xml.count("<w:drawing")
        media_count = len([name for name in archive_names if name.startswith("word/media/")])
        require(len(docx_text) > 25_000, "release DOCX must contain substantial text")
        require(len(document.paragraphs) >= 500, "release DOCX must keep populated paragraphs")
        require(len(document.tables) >= 10, "release DOCX must keep populated tables")
        require(section_heading_count >= 8, "release DOCX must keep formal section headings")
        require(page_break_count >= 2, "release DOCX must keep appendix page breaks")
        require(drawing_count >= 1, "release DOCX must keep report drawings")
        require(media_count >= 1, "release DOCX must keep embedded media")

        for needle in REQUIRED_REPORT_TEXT:
            require(needle in pdf_text, f"release PDF missing required text: {needle}")
            require(needle in docx_text, f"release DOCX missing required text: {needle}")
        for needle in ("文件狀態：正式附件", "產出工具", "開挖擋土支撐計算書", "工具版本", "輸出時間", "核可時間"):
            require(needle in pdf_text, f"release PDF missing formal document metadata: {needle}")
            require(needle in docx_text, f"release DOCX missing formal document metadata: {needle}")
        require("文件狀態：內部審閱" not in pdf_text, "release PDF must not retain internal-review status")
        require("文件狀態：內部審閱" not in docx_text, "release DOCX must not retain internal-review status")
        for needle in PAGE_ONLY_REPORT_STATUS_NEEDLES:
            require(needle not in pdf_text, f"release PDF contains page-only text: {needle}")
            require(needle not in docx_text, f"release DOCX contains page-only text: {needle}")
        require(fingerprint in pdf_text, "release PDF must contain the replayed calculation fingerprint")
        require(fingerprint in docx_text, "release DOCX must contain the replayed calculation fingerprint")

        pdf_hash = sha256(pdf_path)
        docx_hash = sha256(docx_path)
        require(sha256(latest_pdf_path) == pdf_hash, "latest PDF must match generated PDF")
        require(sha256(latest_docx_path) == docx_hash, "latest DOCX must match generated DOCX")

        receiver_response = build_receiver_capacity_attachment_fixture()
        receiver_json_paths: dict[str, Path] = {}
        for key in ("evidence", "bearingEvidence"):
            envelope = receiver_response[key]
            evidence_bytes = base64.b64decode(
                envelope["contentBase64"],
                validate=True,
            )
            evidence_path = output_dir / envelope["fileName"]
            evidence_path.write_bytes(evidence_bytes)
            require(
                sha256(evidence_path) == envelope["fileSha256"],
                f"receiver {key} JSON SHA-256 must match its exact response envelope",
            )
            receiver_json_paths[key] = evidence_path
        attachment_output_time = datetime.now(timezone.utc)
        receiver_pdf_artifact = build_reshore_capacity_attachment(
            receiver_response,
            report_kind="pdf",
            output_at=attachment_output_time,
        )
        generated_attachment_paths.append(receiver_pdf_artifact.path)
        receiver_docx_artifact = build_reshore_capacity_attachment(
            receiver_response,
            report_kind="docx",
            output_at=attachment_output_time,
        )
        generated_attachment_paths.append(receiver_docx_artifact.path)
        receiver_render_evidence = build_pdf_canonical_render_evidence(
            receiver_pdf_artifact.path
        )
        generated_attachment_paths.append(receiver_render_evidence)
        receiver_source_bundle = build_pdf_formal_source_bundle(
            receiver_pdf_artifact.path,
            receiver_render_evidence,
        )
        generated_attachment_paths.append(receiver_source_bundle)

        receiver_release_paths = {
            "pdf": output_dir / receiver_pdf_artifact.path.name,
            "docx": output_dir / receiver_docx_artifact.path.name,
            "evidence": output_dir / receiver_render_evidence.name,
            "bundle": output_dir / receiver_source_bundle.name,
        }
        for source, destination in (
            (receiver_pdf_artifact.path, receiver_release_paths["pdf"]),
            (receiver_docx_artifact.path, receiver_release_paths["docx"]),
            (receiver_render_evidence, receiver_release_paths["evidence"]),
            (receiver_source_bundle, receiver_release_paths["bundle"]),
        ):
            if source.resolve() != destination.resolve():
                shutil.copy2(source, destination)

        receiver_pdf_reader = PdfReader(str(receiver_release_paths["pdf"]))
        receiver_pdf_text = "\n".join(
            page.extract_text() or "" for page in receiver_pdf_reader.pages
        )
        receiver_document = Document(str(receiver_release_paths["docx"]))
        receiver_docx_text = "\n".join(
            [paragraph.text for paragraph in receiver_document.paragraphs]
            + [
                cell.text
                for table in receiver_document.tables
                for row in table.rows
                for cell in row.cells
            ]
        )
        rsc_fingerprint = receiver_response["calculation"]["calculationFingerprint"]
        rsb_fingerprint = receiver_response["bearingEvidence"]["documentReference"]
        for needle in (
            "RSC v4",
            "RSB v1",
            rsc_fingerprint,
            rsb_fingerprint,
            receiver_response["evidence"]["fileSha256"],
            receiver_response["bearingEvidence"]["fileSha256"],
            "0.298816",
            "0.717457",
            "0.223352",
            "PASS",
            "https://ej.aisc.org/index.php/engj/article/view/214",
            "本獨立接收端附件不會寫入或改動來源專案的主 PDF/DOCX 計算書。",
        ):
            require(needle in receiver_pdf_text, f"receiver PDF missing required text: {needle}")
            require(needle in receiver_docx_text, f"receiver DOCX missing required text: {needle}")
        require(
            receiver_pdf_artifact.attachment_file_sha256
            == sha256(receiver_release_paths["pdf"]),
            "receiver PDF artifact SHA-256 must survive release copy",
        )
        require(
            receiver_docx_artifact.attachment_file_sha256
            == sha256(receiver_release_paths["docx"]),
            "receiver DOCX artifact SHA-256 must survive release copy",
        )
        receiver_evidence = json.loads(
            receiver_release_paths["evidence"].read_text(encoding="utf-8")
        )
        require(
            receiver_evidence.get("artifact") == receiver_release_paths["pdf"].name,
            "receiver canonical evidence must identify the exact PDF",
        )
        require(
            receiver_evidence.get("artifactSha256")
            == sha256(receiver_release_paths["pdf"]),
            "receiver canonical evidence must bind the exact PDF SHA-256",
        )
        require(
            receiver_evidence.get("pdf", {}).get("pageCount")
            == len(receiver_pdf_reader.pages),
            "receiver canonical evidence page count must match the PDF",
        )
        with ZipFile(receiver_release_paths["bundle"]) as receiver_bundle:
            require(
                receiver_bundle.namelist()
                == [
                    receiver_release_paths["pdf"].name,
                    receiver_release_paths["evidence"].name,
                ],
                "receiver formal source bundle must contain only the PDF/evidence pair",
            )
            require(
                hashlib.sha256(
                    receiver_bundle.read(receiver_release_paths["pdf"].name)
                ).hexdigest()
                == sha256(receiver_release_paths["pdf"]),
                "receiver formal source bundle PDF bytes must match",
            )

        receiver_attachment_summary = {
            "schemaVersion": 1,
            "kind": "rsc-v4-rsb-independent-calculation-attachment",
            "documentStatus": "formal-attachment",
            "rscCalculationFingerprint": rsc_fingerprint,
            "rscEvidenceFileName": receiver_response["evidence"]["fileName"],
            "rscEvidenceBytes": receiver_json_paths["evidence"].stat().st_size,
            "rscEvidenceFileSha256": receiver_response["evidence"]["fileSha256"],
            "rsbCalculationFingerprint": rsb_fingerprint,
            "rsbEvidenceFileName": receiver_response["bearingEvidence"]["fileName"],
            "rsbEvidenceBytes": receiver_json_paths["bearingEvidence"].stat().st_size,
            "rsbEvidenceFileSha256": receiver_response["bearingEvidence"]["fileSha256"],
            "artifact": receiver_release_paths["pdf"].name,
            "artifactBytes": receiver_release_paths["pdf"].stat().st_size,
            "artifactSha256": sha256(receiver_release_paths["pdf"]),
            "document": receiver_release_paths["docx"].name,
            "documentBytes": receiver_release_paths["docx"].stat().st_size,
            "documentSha256": sha256(receiver_release_paths["docx"]),
            "canonicalEvidence": receiver_release_paths["evidence"].name,
            "canonicalEvidenceBytes": receiver_release_paths["evidence"].stat().st_size,
            "canonicalEvidenceSha256": sha256(receiver_release_paths["evidence"]),
            "formalSourceBundle": receiver_release_paths["bundle"].name,
            "formalSourceBundleBytes": receiver_release_paths["bundle"].stat().st_size,
            "formalSourceBundleSha256": sha256(receiver_release_paths["bundle"]),
            "pdfPageCount": len(receiver_pdf_reader.pages),
            "pdfTextLength": len(receiver_pdf_text),
            "documentTextLength": len(receiver_docx_text),
            "pass": True,
        }

        summary = {
            "schemaVersion": 1,
            "family": "excavation-formal",
            "generatedAt": datetime.now(timezone.utc).isoformat(),
            "required": 1,
            "complete": ["excavation-report"],
            "pass": True,
            "records": [
                {
                    "key": "excavation-report",
                    "artifact": pdf_path.name,
                    "document": docx_path.name,
                    "latestArtifact": latest_pdf_path.relative_to(output_dir).as_posix(),
                    "latestDocument": latest_docx_path.relative_to(output_dir).as_posix(),
                    "latestArtifactBytes": latest_pdf_path.stat().st_size,
                    "latestDocumentBytes": latest_docx_path.stat().st_size,
                    "latestArtifactSha256": sha256(latest_pdf_path),
                    "latestDocumentSha256": sha256(latest_docx_path),
                    "artifactBytes": pdf_path.stat().st_size,
                    "documentBytes": docx_path.stat().st_size,
                    "documentXmlBytes": len(document_xml.encode("utf-8")),
                    "artifactSha256": pdf_hash,
                    "documentSha256": docx_hash,
                    "pdfPageCount": len(reader.pages),
                    "pdfTextLength": len(pdf_text),
                    "documentTextLength": len(docx_text),
                    "paragraphCount": len(document.paragraphs),
                    "tableCount": len(document.tables),
                    "sectionCount": section_heading_count,
                    "documentXmlTextLength": len(document_xml_text),
                    "xmlParagraphCount": xml_paragraph_count,
                    "xmlTableCount": xml_table_count,
                    "xmlSectionCount": xml_section_count,
                    "pageBreakCount": page_break_count,
                    "drawingCount": drawing_count,
                    "mediaCount": media_count,
                    "projectName": project.metadata.name,
                    "resultReconciliation": {
                        "schemaVersion": 1,
                        "strategy": "excavation-project-state-replay-to-pdf-docx-hash",
                        "caseId": project.metadata.id,
                        "sourceProject": source_project_path.name,
                        "sourceProjectSha256": sha256(source_project_path),
                        "resultSha256": result_sha256,
                        "calculationFingerprint": fingerprint,
                        "verifiedCheckCount": verified_check_count,
                        "verifiedAssertionCount": verified_assertion_count,
                        "verifiedGroupCounts": verified_group_counts,
                        "pdfSha256": pdf_hash,
                        "docxSha256": docx_hash,
                        "pass": True,
                    },
                    "receiverCapacityAttachment": receiver_attachment_summary,
                }
            ],
        }
        (output_dir / "rendered-delivery-evidence-summary.json").write_text(
            json.dumps(summary, ensure_ascii=False, indent=2) + "\n",
            encoding="utf-8",
        )
        print(
            "Excavation release artifacts OK "
            f"(pdfPages={len(reader.pages)}, pdfText={len(pdf_text)}, "
            f"docxText={len(docx_text)}, checks={verified_check_count}, "
            f"assertions={verified_assertion_count}, receiverAttachment=PDF+DOCX, "
            f"output={output_dir})"
        )
        return 0
    finally:
        if generated_pdf is not None:
            generated_pdf.unlink(missing_ok=True)
        if generated_docx is not None:
            generated_docx.unlink(missing_ok=True)
        for attachment_path in generated_attachment_paths:
            if attachment_path.resolve().parent != output_dir:
                attachment_path.unlink(missing_ok=True)


if __name__ == "__main__":
    raise SystemExit(main())
