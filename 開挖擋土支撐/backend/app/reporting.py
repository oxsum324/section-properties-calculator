from __future__ import annotations

from collections import Counter, defaultdict
from datetime import datetime
import math
from pathlib import Path
from zipfile import ZipFile
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.lib.utils import ImageReader
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.pdfbase.pdfmetrics import registerFont
from reportlab.platypus import Image as RLImage
from reportlab.platypus import KeepTogether, PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from .config import get_settings
from .schemas import BasicParameters, CheckResult, ProjectState, SummaryItem


def build_report(project: ProjectState, *, concise_mode: bool = False) -> Path:
    settings = get_settings()
    registerFont(UnicodeCIDFont("STSong-Light"))
    mode_slug = "concise" if concise_mode else "detail"
    report_path = settings.reports_dir / f"{project.metadata.id or 'draft'}-{mode_slug}-{datetime.now():%Y%m%d%H%M%S%f}.pdf"
    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="ZHBody",
            fontName="STSong-Light",
            fontSize=10.5,
            leading=16,
            alignment=TA_LEFT,
        )
    )
    styles.add(
        ParagraphStyle(
            name="ZHTitle",
            fontName="STSong-Light",
            fontSize=20,
            leading=28,
            alignment=TA_LEFT,
            textColor=colors.HexColor("#0f172a"),
        )
    )
    styles.add(
        ParagraphStyle(
            name="ZHHeading",
            fontName="STSong-Light",
            fontSize=14,
            leading=20,
            spaceBefore=10,
            spaceAfter=7,
            textColor=colors.HexColor("#0f4c81"),
        )
    )
    styles.add(
        ParagraphStyle(
            name="ZHCell",
            fontName="STSong-Light",
            fontSize=9,
            leading=11,
            alignment=TA_LEFT,
        )
    )
    styles.add(
        ParagraphStyle(
            name="ZHLead",
            fontName="STSong-Light",
            fontSize=11.5,
            leading=16,
            alignment=TA_LEFT,
            textColor=colors.HexColor("#0f172a"),
        )
    )
    elements = []
    elements.append(Paragraph("擋土支撐檢核計算書", styles["ZHTitle"]))
    elements.append(Spacer(1, 4 * mm))
    meta_rows = _report_meta_rows(project)
    meta_table = Table(meta_rows, colWidths=[35 * mm, 145 * mm])
    meta_table.setStyle(
        TableStyle(
            [
                ("FONTNAME", (0, 0), (-1, -1), "STSong-Light"),
                ("FONTSIZE", (0, 0), (-1, -1), 10),
                ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#eff6ff")),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#cbd5e1")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("PADDING", (0, 0), (-1, -1), 6),
            ]
        )
    )
    elements.append(meta_table)
    elements.append(Spacer(1, 6 * mm))

    elements.append(Paragraph("一、摘要", styles["ZHHeading"]))
    elements.append(
        Paragraph(
            (
                f"本計算書係就「{project.metadata.name}」之擋土支撐系統進行檢核，"
                f"{_analysis_source_narrative(project)}，"
                "並依輸入之材料、土層與構件條件，完成支撐系統與柱構件之整體驗算。"
            ),
            styles["ZHBody"],
        )
    )
    if project.calculation_results:
        elements.append(Paragraph("檢核結論", styles["ZHLead"]))
        executive_rows = _executive_summary_rows(project, project.calculation_results, concise_mode=concise_mode)
        executive_table = Table(executive_rows, colWidths=[40 * mm, 140 * mm])
        executive_table.setStyle(_table_style())
        elements.append(executive_table)
        for line in _summary_conclusion_lines(project, project.calculation_results):
            elements.append(Paragraph(line, styles["ZHBody"]))
    elements.append(Paragraph("內容提要", styles["ZHLead"]))
    outline_table = Table(_report_outline_rows(project), colWidths=[38 * mm, 142 * mm])
    outline_table.setStyle(_table_style())
    elements.append(outline_table)
    elements.append(Paragraph("報告說明", styles["ZHLead"]))
    for line in _report_scope_lines(project):
        elements.append(Paragraph(line, styles["ZHBody"]))

    elements.append(Paragraph("二、設計依據", styles["ZHHeading"]))
    elements.append(Paragraph("設計規範與檢核依據", styles["ZHLead"]))
    for index, line in enumerate(_design_basis_lines(project), start=1):
        elements.append(Paragraph(f"{index}. {line}", styles["ZHBody"]))

    elements.append(Paragraph("三、結構分析使用之電腦程式", styles["ZHHeading"]))
    elements.append(Paragraph(_analysis_program_description(project), styles["ZHBody"]))

    elements.append(Paragraph("四、材料性質", styles["ZHHeading"]))
    basic = project.basic_parameters
    elements.append(Paragraph("鋼材與混凝土基本設定", styles["ZHLead"]))
    material_rows = [
        ["鋼材彈性係數 E (tf/cm2)", f"{basic.e_tf_per_cm2:.2f}"],
        ["鋼材降伏應力 Fy (tf/cm2)", f"{basic.fy_tf_per_cm2:.2f}"],
        ["材料折減係數 ψ", f"{basic.psi_material:.2f}"],
        ["混凝土抗壓強度 fc' (kg/cm2)", f"{basic.wall_fc_kg_per_cm2:.1f}"],
    ]
    material_table = Table(material_rows, colWidths=[55 * mm, 125 * mm])
    material_table.setStyle(_table_style())
    elements.append(material_table)
    elements.append(Paragraph("擋土措施", styles["ZHLead"]))
    elements.append(
        Paragraph(
            f"擋土壁型式為 {basic.wall_type}，牆厚約 {basic.wall_thickness_cm:.1f} cm。",
            styles["ZHBody"],
        )
    )

    elements.append(Paragraph("五、輸入基本資料", styles["ZHHeading"]))
    elements.append(Paragraph("材料勁度", styles["ZHLead"]))
    basic_rows = [
        ["E (tf/cm2)", f"{basic.e_tf_per_cm2:.2f}"],
        ["Fy (tf/cm2)", f"{basic.fy_tf_per_cm2:.2f}"],
        ["Cm", f"{basic.cm_factor:.2f}"],
        ["積載重 WL (tf/m)", f"{basic.surcharge_wl_tf_per_m:.3f}"],
        ["αs / αw / αb / αp", f"{basic.alpha_support:.2f} / {basic.alpha_wale:.2f} / {basic.alpha_brace:.2f} / {basic.alpha_column:.2f}"],
        ["ψ", f"{basic.psi_material:.2f}"],
    ]
    basic_table = Table(basic_rows, colWidths=[55 * mm, 125 * mm])
    basic_table.setStyle(_table_style())
    elements.append(basic_table)

    analysis = project.analysis_import
    soils = _collect_word_soils(project)
    if soils:
        elements.append(Paragraph("土壤參數", styles["ZHLead"]))
        soil_table = Table(
            [["層次", "名稱", "深度 (m)", "厚度 (m)", "γ (t/m3)", "φ", "c", "Su", "Kh", "型態"], *_build_soil_plain_rows(soils)],
            colWidths=[12 * mm, 22 * mm, 18 * mm, 18 * mm, 18 * mm, 11 * mm, 11 * mm, 14 * mm, 14 * mm, 22 * mm],
            repeatRows=1,
        )
        soil_table.setStyle(_table_style(header=True))
        elements.append(soil_table)
    if analysis.stages:
        elements.append(Paragraph("開挖施工步驟概要", styles["ZHLead"]))
        stage_table = Table(
            [["階段", "開挖深度 (m)", "地下水位 (m)", "支撐數", "備註"], *_build_stage_plain_rows(analysis)],
            colWidths=[16 * mm, 28 * mm, 28 * mm, 16 * mm, 92 * mm],
            repeatRows=1,
        )
        stage_table.setStyle(_table_style(header=True))
        elements.append(stage_table)
    elements.append(Paragraph("分析匯入資訊", styles["ZHLead"]))
    analysis_rows = _analysis_input_rows(project)
    analysis_table = Table(analysis_rows, colWidths=[55 * mm, 125 * mm])
    analysis_table.setStyle(_table_style())
    elements.append(analysis_table)
    analysis_overview = _analysis_import_overview_text(project)
    if analysis_overview:
        elements.append(Paragraph(analysis_overview, styles["ZHBody"]))
    if analysis.warnings:
        elements.append(Paragraph(f"匯入判讀註記：{'；'.join(analysis.warnings)}", styles["ZHBody"]))

    results = project.calculation_results
    if results:
        elements.append(Paragraph("六、結構計算結果", styles["ZHHeading"]))
        elements.append(Paragraph("檢核結果統計", styles["ZHLead"]))
        elements.append(
            Paragraph(
                (
                    f"本次共完成 {len(results.summary)} 筆摘要項目檢核，"
                    f"其中 {_counted_status_phrase('OK', sum(1 for item in results.summary if item.status == 'OK'))}、"
                    f"{_counted_status_phrase('Say~OK', sum(1 for item in results.summary if item.status == 'Say~OK'))}、"
                    f"{_counted_status_phrase('NG', sum(1 for item in results.summary if item.status == 'NG'))}。"
                ),
                styles["ZHBody"],
            )
        )
        elements.append(Paragraph("支撐系統檢核摘要", styles["ZHLead"]))
        elements.append(
            Paragraph(
                "摘要表已將同層構件彙列於同一列，欄內依序列示控制比值、判定結果與採用型號，以利快速比對各層控制情形。",
                styles["ZHBody"],
            )
        )
        for line in _result_conclusion_lines(project, results):
            elements.append(Paragraph(line, styles["ZHBody"]))
        level_headers, level_data_rows, level_statuses = _build_level_summary_rows(results.summary, styles["ZHCell"])
        level_table = Table(
            [level_headers, *level_data_rows],
            colWidths=_level_summary_col_widths(len(level_headers)),
            repeatRows=1,
        )
        level_style = _table_style(header=True)
        _apply_status_styles(level_style, level_statuses, status_col=len(level_headers) - 1)
        level_table.setStyle(level_style)
        elements.append(level_table)
        elements.append(Spacer(1, 4 * mm))

        column_section: list[object] = [Paragraph("柱構件檢核摘要", styles["ZHLead"])]
        if results.column_checks:
            column_rows, column_statuses = _build_column_summary_rows(results.column_checks, styles["ZHCell"])
            column_table = Table(
                column_rows,
                colWidths=[42 * mm, 38 * mm, 20 * mm, 22 * mm, 58 * mm],
                repeatRows=1,
            )
            column_style = _table_style(header=True)
            _apply_status_styles(column_style, column_statuses, status_col=3)
            column_table.setStyle(column_style)
            column_section.append(column_table)
        else:
            column_section.append(
                Paragraph("本案未納入中間柱 / 共構柱檢討。", styles["ZHBody"])
            )
        control_section: list[object] = [Paragraph("主要控制項目彙整", styles["ZHLead"])]
        control_rows, control_statuses = _build_control_rows(results, styles["ZHCell"])
        control_table = Table(
            control_rows,
            colWidths=[28 * mm, 28 * mm, 42 * mm, 36 * mm, 18 * mm, 18 * mm],
            repeatRows=1,
        )
        control_style = _table_style(header=True)
        _apply_status_styles(control_style, control_statuses, status_col=5)
        control_table.setStyle(control_style)
        control_section.append(control_table)
        column_section.append(Spacer(1, 4 * mm))
        column_section.extend(control_section)
        elements.append(KeepTogether(column_section))

        appendix_one_groups = _appendix_one_groups(results)
        if any(checks for _, checks in appendix_one_groups):
            appendix_one_title = _appendix_one_title(appendix_one_groups)
            elements.append(PageBreak())
            elements.append(Paragraph(appendix_one_title, styles["ZHHeading"]))
            elements.append(
                Paragraph(
                    _appendix_one_intro_text(appendix_one_groups),
                    styles["ZHBody"],
                )
            )
            _append_pdf_common_parameter_table(
                elements,
                "附件一共用參數",
                _appendix_common_parameter_rows(
                    basic,
                    appendix="one",
                    wall_deduction=project.calculation_options.consider_wall_deduction_for_wales,
                ),
                styles,
            )
            appendix_one_checks = [check for _, checks in appendix_one_groups for check in checks]
            _append_pdf_section_summary_table(elements, appendix_one_checks, "附件一型鋼彙整表", styles)
            if _appendix_group_has_checks(appendix_one_groups, "水平支撐細部檢核"):
                _append_pdf_support_diagram(elements, styles)
            for title, checks in appendix_one_groups:
                if not checks:
                    continue
                _append_pdf_detail_group(
                    elements,
                    title,
                    checks,
                    basic,
                    styles,
                    section_reference_title="附件一型鋼彙整表",
                    concise_mode=concise_mode,
                )

        if results.column_checks:
            elements.append(PageBreak())
            elements.append(Paragraph("附件二：中間柱、共構柱細部計算結果", styles["ZHHeading"]))
            elements.append(
                Paragraph(
                    "本附件列示中間柱及共構柱之檢核內容，並依序整理軸力、彎矩交互作用與基礎壓入、拉拔驗算結果。",
                    styles["ZHBody"],
                )
            )
            _append_pdf_common_parameter_table(
                elements,
                "附件二共用參數",
                _appendix_common_parameter_rows(basic, appendix="two"),
                styles,
            )
            _append_pdf_section_summary_table(elements, results.column_checks, "附件二型鋼彙整表", styles)
            _append_pdf_detail_group(
                elements,
                "柱構件細部檢核",
                results.column_checks,
                basic,
                styles,
                section_reference_title="附件二型鋼彙整表",
                concise_mode=concise_mode,
            )

    if project.metadata.notes:
        elements.append(Paragraph("七、備註", styles["ZHHeading"]))
        elements.append(Paragraph(project.metadata.notes.replace("\n", "<br/>"), styles["ZHBody"]))

    doc = SimpleDocTemplate(
        str(report_path),
        pagesize=A4,
        leftMargin=15 * mm,
        rightMargin=15 * mm,
        topMargin=20 * mm,
        bottomMargin=18 * mm,
    )
    doc.build(
        elements,
        onFirstPage=lambda canvas, doc: _draw_pdf_page_frame(canvas, doc, project, concise_mode=concise_mode, include_header=False),
        onLaterPages=lambda canvas, doc: _draw_pdf_page_frame(canvas, doc, project, concise_mode=concise_mode),
    )
    return report_path


def build_word_report(project: ProjectState, *, concise_mode: bool = False) -> Path:
    settings = get_settings()
    mode_slug = "concise" if concise_mode else "detail"
    report_path = settings.reports_dir / f"{project.metadata.id or 'draft'}-{mode_slug}-{datetime.now():%Y%m%d%H%M%S%f}.docx"
    document = _new_word_document()
    _configure_document_styles(document)
    _configure_document_layout(document)
    _configure_document_header_footer(document, project, concise_mode=concise_mode)
    analysis = project.analysis_import
    basic = project.basic_parameters
    soils = _collect_word_soils(project)

    _add_report_title_block(document, project)
    _add_main_heading(document, _main_section_title("摘要"))
    _add_body_paragraph(
        document,
        (
            f"本計算書係就『{project.metadata.name}』之擋土支撐系統進行檢核，"
            f"{_analysis_source_narrative(project)}，"
            "並依輸入之材料、土層與構件條件，完成支撐系統與柱構件之整體驗算。"
        ),
    )
    if project.calculation_results:
        _add_subheading(document, "檢核結論")
        _add_key_value_table(document, _executive_summary_rows(project, project.calculation_results, concise_mode=concise_mode))
        for line in _summary_conclusion_lines(project, project.calculation_results):
            _add_body_paragraph(document, line)
    _add_subheading(document, "內容提要")
    _add_key_value_table(document, _report_outline_rows(project))
    _add_subheading(document, "報告說明")
    for line in _report_scope_lines(project):
        _add_body_paragraph(document, line)

    _add_main_heading(document, _main_section_title("設計依據"))
    _add_subheading(document, "設計規範與檢核依據")
    for index, line in enumerate(_design_basis_lines(project), start=1):
        _add_body_paragraph(document, f"{index}. {line}")

    _add_main_heading(document, _main_section_title("結構分析使用之電腦程式"))
    _add_body_paragraph(
        document,
        _analysis_program_description(project),
    )

    _add_main_heading(document, _main_section_title("材料性質"))
    _add_subheading(document, "鋼材與混凝土基本設定")
    _add_key_value_table(
        document,
        [
            ["鋼材彈性係數 E (tf/cm2)", f"{basic.e_tf_per_cm2:.2f}"],
            ["鋼材降伏應力 Fy (tf/cm2)", f"{basic.fy_tf_per_cm2:.2f}"],
            ["材料折減係數 ψ", f"{basic.psi_material:.2f}"],
            ["混凝土抗壓強度 fc' (kg/cm2)", f"{basic.wall_fc_kg_per_cm2:.1f}"],
        ],
    )
    _add_subheading(document, "擋土措施")
    _add_body_paragraph(
        document,
        f"擋土壁型式為 {basic.wall_type}，牆厚約 {basic.wall_thickness_cm:.1f} cm。",
    )

    _add_main_heading(document, _main_section_title("輸入基本資料"))
    _add_subheading(document, "材料勁度")
    _add_key_value_table(
        document,
        [
            ["Cm", f"{basic.cm_factor:.2f}"],
            ["積載重 WL (tf/m)", f"{basic.surcharge_wl_tf_per_m:.3f}"],
            [
                "αs / αw / αb / αp",
                f"{basic.alpha_support:.2f} / {basic.alpha_wale:.2f} / {basic.alpha_brace:.2f} / {basic.alpha_column:.2f}",
            ],
        ],
    )
    if soils:
        _add_subheading(document, "土壤參數")
        _add_grid_table(
            document,
            ["層次", "名稱", "深度 (m)", "厚度 (m)", "γ (t/m3)", "φ", "c", "Su", "Kh", "型態"],
            _build_soil_plain_rows(soils),
        )
    if analysis.stages:
        _add_subheading(document, "開挖施工步驟概要")
        _add_grid_table(
            document,
            ["階段", "開挖深度 (m)", "地下水位 (m)", "支撐數", "備註"],
            _build_stage_plain_rows(analysis),
        )
    _add_subheading(document, "分析匯入資訊")
    _add_key_value_table(
        document,
        _analysis_input_rows(project),
    )
    analysis_overview = _analysis_import_overview_text(project)
    if analysis_overview:
        _add_body_paragraph(document, analysis_overview)
    if analysis.warnings:
        _add_body_paragraph(document, f"匯入判讀註記：{'；'.join(analysis.warnings)}")

    results = project.calculation_results
    if results:
        _add_main_heading(document, _main_section_title("結構計算結果"))
        _add_subheading(document, "檢核結果統計")
        _add_body_paragraph(
            document,
            (
                f"本次共完成 {len(results.summary)} 筆摘要項目檢核，"
                f"其中 {_counted_status_phrase('OK', sum(1 for item in results.summary if item.status == 'OK'))}、"
                f"{_counted_status_phrase('Say~OK', sum(1 for item in results.summary if item.status == 'Say~OK'))}、"
                f"{_counted_status_phrase('NG', sum(1 for item in results.summary if item.status == 'NG'))}。"
            ),
        )
        for line in _result_conclusion_lines(project, results):
            _add_body_paragraph(document, line)
        _add_subheading(document, "支撐系統檢核摘要")
        level_headers, level_plain_rows = _build_level_summary_plain_rows(results.summary)
        _add_grid_table(
            document,
            level_headers,
            level_plain_rows,
            status_col=len(level_headers) - 1,
        )

        _add_subheading(document, "柱構件檢核摘要")
        if results.column_checks:
            _add_grid_table(
                document,
                ["構件", "型號", "利用率", "狀態", "備註"],
                _build_column_summary_plain_rows(results.column_checks),
                status_col=3,
            )
        else:
            _add_body_paragraph(document, "本案未納入中間柱 / 共構柱檢討。")

        _add_subheading(document, "主要控制項目彙整")
        _add_grid_table(
            document,
            ["模組", "標籤", "控制條件", "控制值 / 允許值", "利用率", "狀態"],
            _build_control_plain_rows(results),
            status_col=5,
        )

        appendix_one_groups = _appendix_one_groups(results)
        if any(checks for _, checks in appendix_one_groups):
            _add_page_break(document)
            _add_main_heading(document, _appendix_one_title(appendix_one_groups))
            _add_body_paragraph(
                document,
                _appendix_one_intro_text(appendix_one_groups),
            )
            _append_word_common_parameter_table(
                document,
                "附件一共用參數",
                _appendix_common_parameter_rows(
                    basic,
                    appendix="one",
                    wall_deduction=project.calculation_options.consider_wall_deduction_for_wales,
                ),
            )
            appendix_one_checks = [check for _, checks in appendix_one_groups for check in checks]
            _append_word_section_summary_table(document, appendix_one_checks, "附件一型鋼彙整表")
            if _appendix_group_has_checks(appendix_one_groups, "水平支撐細部檢核"):
                _append_word_support_diagram(document)
            for title, checks in appendix_one_groups:
                if not checks:
                    continue
                _append_word_detail_group(
                    document,
                    title,
                    checks,
                    basic,
                    section_reference_title="附件一型鋼彙整表",
                    concise_mode=concise_mode,
                )

        if results.column_checks:
            _add_page_break(document)
            _add_main_heading(document, "附件二：中間柱、共構柱細部計算結果")
            _add_body_paragraph(
                document,
                "本附件列示中間柱及共構柱之檢核內容，並依序整理軸力、彎矩交互作用與基礎壓入、拉拔驗算結果。",
            )
            _append_word_common_parameter_table(
                document,
                "附件二共用參數",
                _appendix_common_parameter_rows(basic, appendix="two"),
            )
            _append_word_section_summary_table(document, results.column_checks, "附件二型鋼彙整表")
            _append_word_detail_group(
                document,
                "柱構件細部檢核",
                results.column_checks,
                basic,
                section_reference_title="附件二型鋼彙整表",
                concise_mode=concise_mode,
            )

    if project.metadata.notes:
        _add_main_heading(document, _main_section_title("備註"))
        for line in project.metadata.notes.splitlines():
            _add_body_paragraph(document, line or " ")

    document.save(report_path)
    return report_path


def _new_word_document() -> Document:
    settings = get_settings()
    if settings.word_template_path.exists():
        document = Document(str(settings.word_template_path))
        _clear_document_body(document)
        _clear_headers_and_footers(document)
        return document
    return Document()


def _clear_document_body(document: Document) -> None:
    body = document._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def _clear_headers_and_footers(document: Document) -> None:
    for section in document.sections:
        for part in (section.header, section.footer, section.first_page_header, section.first_page_footer):
            for paragraph in part.paragraphs:
                paragraph.text = ""


def _configure_document_styles(document: Document) -> None:
    normal_style = document.styles["Normal"]
    normal_style.font.name = "Microsoft JhengHei"
    normal_style.font.size = Pt(10.5)
    _set_style_east_asia_font(normal_style, "Microsoft JhengHei")
    normal_style.paragraph_format.space_after = Pt(4)
    normal_style.paragraph_format.line_spacing = 1.28


def _configure_document_layout(document: Document) -> None:
    for section in document.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(1.9)
        section.left_margin = Cm(2.1)
        section.right_margin = Cm(2.1)
        section.header_distance = Cm(0.9)
        section.footer_distance = Cm(0.9)


def _configure_document_header_footer(document: Document, project: ProjectState, *, concise_mode: bool) -> None:
    mode_label = "簡述版" if concise_mode else "詳細版"
    for section in document.sections:
        section.different_first_page_header_footer = True
        header = section.header
        footer = section.footer
        first_page_header = section.first_page_header
        first_page_footer = section.first_page_footer
        if header.paragraphs:
            for paragraph in header.paragraphs:
                paragraph.text = ""
        else:
            header.add_paragraph()
        if footer.paragraphs:
            for paragraph in footer.paragraphs:
                paragraph.text = ""
        else:
            footer.add_paragraph()
        if first_page_header.paragraphs:
            for paragraph in first_page_header.paragraphs:
                paragraph.text = ""
        else:
            first_page_header.add_paragraph()
        if first_page_footer.paragraphs:
            for paragraph in first_page_footer.paragraphs:
                paragraph.text = ""
        else:
            first_page_footer.add_paragraph()

        title_paragraph = header.paragraphs[0]
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_paragraph.paragraph_format.space_after = Pt(1)
        _add_run(title_paragraph, "擋土支撐檢核計算書", bold=True, size=9.2, color="334155")

        subtitle_paragraph = header.add_paragraph()
        subtitle_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_paragraph.paragraph_format.space_after = Pt(0)
        _add_run(
            subtitle_paragraph,
            f"{project.metadata.name or '未命名專案'}｜{mode_label}",
            size=8.2,
            color="64748B",
        )

        footer_paragraph = footer.paragraphs[0]
        footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_paragraph.paragraph_format.space_after = Pt(0)
        _add_run(footer_paragraph, "第 ", size=8.2, color="64748B")
        _append_word_page_number_field(footer_paragraph)
        _add_run(footer_paragraph, " 頁", size=8.2, color="64748B")

        first_page_footer_paragraph = first_page_footer.paragraphs[0]
        first_page_footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        first_page_footer_paragraph.paragraph_format.space_after = Pt(0)
        _add_run(first_page_footer_paragraph, "第 ", size=8.2, color="64748B")
        _append_word_page_number_field(first_page_footer_paragraph)
        _add_run(first_page_footer_paragraph, " 頁", size=8.2, color="64748B")


def _add_report_title_block(document: Document, project: ProjectState) -> None:
    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_paragraph.paragraph_format.space_after = Pt(8)
    _add_run(title_paragraph, "擋土支撐檢核計算書", bold=True, size=18, color="0F172A")

    _add_key_value_table(document, _report_meta_rows(project))
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)


def _add_run(paragraph, text: str, *, bold: bool = False, size: float = 10, color: str | None = None) -> None:
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Microsoft JhengHei"
    run.font.size = Pt(size)
    _set_run_east_asia_font(run, "Microsoft JhengHei")
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def _add_heading(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(5)
    _add_run(paragraph, text, bold=True, size=14, color="0F4C81")


def _add_main_heading(document: Document, text: str) -> None:
    if _style_exists(document, "階層1"):
        paragraph = document.add_paragraph(style="階層1")
        paragraph.paragraph_format.space_before = Pt(12)
        paragraph.paragraph_format.space_after = Pt(5)
        _add_run(paragraph, text, bold=True, size=14, color="000000")
        return
    _add_heading(document, text)


def _add_subheading(document: Document, text: str) -> None:
    if _style_exists(document, "標題二"):
        paragraph = document.add_paragraph(style="標題二")
        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(2)
        _add_run(paragraph, text, bold=True, size=11)
        return
    _add_body_paragraph(document, text)


def _add_body_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="內文1" if _style_exists(document, "內文1") else None)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.28
    _add_run(paragraph, text, size=10.5)


def _add_page_break(document: Document) -> None:
    document.add_page_break()


def _add_indented_body_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="內文1" if _style_exists(document, "內文1") else None)
    paragraph.paragraph_format.left_indent = Cm(0.7)
    paragraph.paragraph_format.first_line_indent = 0
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.22
    _add_run(paragraph, text, size=10.2)


def _add_lead_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="內文1" if _style_exists(document, "內文1") else None)
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.line_spacing = 1.25
    _add_run(paragraph, text, bold=True, size=11.5)


def _add_detail_item_heading(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="內文1" if _style_exists(document, "內文1") else None)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.2
    _add_run(paragraph, text, bold=True, size=11.5)


def _add_detail_block_heading(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="內文1" if _style_exists(document, "內文1") else None)
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(1)
    paragraph.paragraph_format.line_spacing = 1.15
    _add_run(paragraph, text, bold=True, size=10.8)


def _add_key_value_table(document: Document, rows: list[list[str]]) -> None:
    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for key, value in rows:
        row_cells = table.add_row().cells
        _set_cell_text(row_cells[0], key, bold=True)
        _set_cell_text(row_cells[1], value)
        _shade_cell(row_cells[0], "EFF6FF")


def _add_grid_table(
    document: Document,
    headers: list[str],
    rows: list[list[str]],
    *,
    status_col: int | None = None,
    col_widths_cm: list[float] | None = None,
) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    if col_widths_cm:
        table.autofit = False
        _set_fixed_table_layout(table)
    for index, header in enumerate(headers):
        _set_cell_text(table.rows[0].cells[index], header, bold=True)
        _shade_cell(table.rows[0].cells[index], "DBEAFE")
        if col_widths_cm and index < len(col_widths_cm):
            table.rows[0].cells[index].width = Cm(col_widths_cm[index])
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            _set_cell_text(cells[index], value)
            if col_widths_cm and index < len(col_widths_cm):
                cells[index].width = Cm(col_widths_cm[index])
        if status_col is not None and status_col < len(row):
            bg, fg = _status_palette_hex(row[status_col])
            _shade_cell(cells[status_col], bg)
            _set_cell_text(cells[status_col], row[status_col], color=fg)


def _set_cell_text(cell, text: str, *, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _add_run(paragraph, str(text), bold=bold, color=color)


def _shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_fixed_table_layout(table) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")


def _style_exists(document: Document, style_name: str) -> bool:
    return any(style.name == style_name for style in document.styles)


def _set_style_east_asia_font(style, font_name: str) -> None:
    r_pr = style.element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), font_name)


def _set_run_east_asia_font(run, font_name: str) -> None:
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), font_name)


def _append_word_page_number_field(paragraph) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")

    text = OxmlElement("w:t")
    text.text = "1"

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    run = paragraph.add_run()
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(separate)
    run._r.append(text)
    run._r.append(end)
    _set_run_east_asia_font(run, "Microsoft JhengHei")


def _status_palette_hex(status: str) -> tuple[str, str]:
    if status == "NG":
        return "FEE2E2", "991B1B"
    if status in {"Say~OK", "應注意"}:
        return "FEF3C7", "92400E"
    return "DCFCE7", "166534"


def _fmt(value: object) -> str:
    if value is None or value == "":
        return "—"
    if isinstance(value, float):
        return f"{value:.3f}"
    return str(value)


def _fmt_short(value: object) -> str:
    if value is None or value == "":
        return "—"
    if isinstance(value, float):
        return f"{value:.2f}"
    return str(value)


def _main_section_title(title: str) -> str:
    numbered_titles = {
        "摘要": "一、摘要",
        "設計依據": "二、設計依據",
        "結構分析使用之電腦程式": "三、結構分析使用之電腦程式",
        "材料性質": "四、材料性質",
        "輸入基本資料": "五、輸入基本資料",
        "結構計算結果": "六、結構計算結果",
        "備註": "七、備註",
    }
    return numbered_titles.get(title, title)


def _report_meta_rows(project: ProjectState) -> list[list[str]]:
    return [
        ["工程名稱", project.metadata.name],
        ["專案代號", project.metadata.project_code or "—"],
        ["委託單位", project.metadata.client or "—"],
        ["設計人員", project.metadata.designer or "—"],
        ["校核人員", project.metadata.checker or "—"],
        ["工程位置", project.metadata.location or "—"],
        ["報告日期", datetime.now().strftime("%Y-%m-%d %H:%M")],
    ]


def _report_outline_rows(project: ProjectState) -> list[list[str]]:
    active_modules = "、".join(_active_module_labels(project)) or "本次未納入構件檢核"
    rows = [
        ["二、設計依據", "列示本計算書採用之設計規範、規則與專案規範包版本。"],
        ["三、結構分析使用之電腦程式", "說明分析資料來源、輸入整理方式及後續檢核依據。"],
        ["四、材料性質", "整理鋼材、混凝土與擋土壁之基本材料設定。"],
        ["五、輸入基本資料", "彙整材料勁度、土壤參數、施工步驟與分析匯入資訊。"],
        ["六、結構計算結果", f"綜整 {active_modules} 之主要控制結果、柱構件摘要與控制項目。"],
    ]
    if project.calculation_results:
        appendix_one_names = "、".join(_appendix_one_component_names(_appendix_one_groups(project.calculation_results)))
        if appendix_one_names:
            rows.append(["附件一", f"列示 {appendix_one_names} 之細部計算過程與檢核結果。"])
        if project.calculation_results.column_checks:
            rows.append(["附件二", "列示中間柱及共構柱之細部計算過程與基礎承載檢核結果。"])
    return rows


def _report_scope_lines(project: ProjectState) -> list[str]:
    active_modules = _active_module_labels(project)
    inactive_modules = _inactive_module_labels(project)
    lines: list[str] = []
    if active_modules:
        lines.append(
            f"本計算書之檢核範圍以{'、'.join(active_modules)}為限，各項控制比值與判定結果均依本次輸入資料與檢核條件計算。"
        )
    else:
        lines.append("本計算書目前未納入構件檢核項目，相關內容僅供輸入條件與章節編排確認之用。")
    if inactive_modules:
        lines.append(f"未納入檢討之{'、'.join(inactive_modules)}，不列入本章摘要與後附附件內容。")
    lines.append("本報告所列正文摘要與附件細部驗算，均以本次重新計算之結果為準。")
    return lines


def _design_basis_lines(project: ProjectState) -> list[str]:
    return [
        "最新版中華民國建築技術規則。",
        "鋼結構容許應力設計法規範及解說。",
        "混凝土結構設計規範。",
        f"專案規範包版本：{project.metadata.spec_pack_version or '標準包 v1'}。",
        "分析輸入與施工階段資料以匯入分析成果及人工調整資料為準。",
    ]


def _analysis_source_mode_label(mode: str) -> str:
    if mode == "import":
        return "匯入分析檔"
    if mode == "manual":
        return "手動輸入"
    return "不使用"


def _analysis_source_description(project: ProjectState, side: str) -> str:
    source = project.top_analysis_source if side == "top" else project.bottom_analysis_source
    mode_label = _analysis_source_mode_label(source.mode)
    if source.mode == "import":
        filename = source.import_result.source_name or "未提供檔名"
        return f"{mode_label}（{filename}）"
    return mode_label


def _analysis_source_rows(project: ProjectState) -> list[list[str]]:
    single_side = _active_single_support_side(project)
    if single_side:
        return [["支撐來源", _analysis_source_description(project, single_side)]]
    return [
        ["上層來源", _analysis_source_description(project, "top")],
        ["下層來源", _analysis_source_description(project, "bottom")],
    ]


def _analysis_input_rows(project: ProjectState) -> list[list[str]]:
    analysis = project.analysis_import
    rows = [*_analysis_source_rows(project)]
    has_import_payload = any(
        [
            analysis.source_name,
            analysis.source_type,
            analysis.project_title,
            analysis.excavation_depth_m is not None,
            analysis.ground_water_level_m is not None,
            analysis.wall_ei_tf_m2_per_m is not None,
            len(analysis.soils) > 0,
            len(analysis.stages) > 0,
        ]
    )
    if not has_import_payload:
        rows.append(["資料說明", "本案未匯入外部分析檔，後續檢核以手動輸入資料為準。"])
        return rows
    if analysis.source_name:
        rows.append(["來源檔案", analysis.source_name])
    if analysis.source_type:
        rows.append(["來源格式", analysis.source_type])
    if analysis.project_title:
        rows.append(["標題", analysis.project_title])
    if analysis.excavation_depth_m is not None:
        rows.append(["開挖深度 (m)", _fmt(analysis.excavation_depth_m)])
    if analysis.ground_water_level_m is not None:
        rows.append(["地下水位 (m)", _fmt(analysis.ground_water_level_m)])
    if analysis.wall_ei_tf_m2_per_m is not None:
        rows.append(["牆體 EI", _fmt(analysis.wall_ei_tf_m2_per_m)])
    rows.append(["土層筆數", str(len(analysis.soils))])
    rows.append(["施工階段數", str(len(analysis.stages))])
    return rows


def _analysis_source_summary_text(project: ProjectState) -> str:
    single_side = _active_single_support_side(project)
    if single_side:
        return f"支撐來源為{_analysis_source_description(project, single_side)}"
    return (
        f"上層為{_analysis_source_description(project, 'top')}，"
        f"下層為{_analysis_source_description(project, 'bottom')}"
    )


def _analysis_source_narrative(project: ProjectState) -> str:
    single_side = _active_single_support_side(project)
    if single_side:
        side_label = "上層" if single_side == "top" else "下層"
        return f"支撐資料採{side_label}{_analysis_source_description(project, single_side)}"
    top_description = _analysis_source_description(project, "top")
    bottom_description = _analysis_source_description(project, "bottom")
    if top_description == bottom_description:
        return f"上、下層支撐資料均採{top_description}"
    return f"上層支撐資料採{top_description}；下層支撐資料採{bottom_description}"


def _analysis_program_description(project: ProjectState) -> str:
    analysis = project.analysis_import
    source_narrative = _analysis_source_narrative(project)
    if analysis.source_type or analysis.source_name:
        detail_parts: list[str] = []
        if analysis.source_type:
            detail_parts.append(f"資料格式為 {analysis.source_type}")
        if analysis.source_name:
            detail_parts.append(f"來源檔名為 {analysis.source_name}")
        detail_text = "；".join(detail_parts)
        return (
            f"{source_narrative}。"
            f"本案另參酌外部分析成果，{detail_text}；"
            "匯入並整理後之支撐軸力、土層與施工階段資料，作為後續構件檢核之依據。"
        )
    return (
        f"{source_narrative}。"
        "本案未匯入外部分析檔，支撐軸力、土層及施工階段資料均以人工整理後之輸入內容作為後續構件檢核之依據。"
    )


def _active_module_labels(project: ProjectState, *, include_columns: bool = True) -> list[str]:
    labels: list[str] = []
    options = project.calculation_options
    if options.include_top_supports or options.include_bottom_supports:
        labels.append("水平支撐")
    if options.include_top_wales or options.include_bottom_wales:
        labels.append("橫擋")
    if options.include_top_braces or options.include_bottom_braces:
        labels.append("斜撐")
    if options.include_corner_braces:
        labels.append("大角撐")
    if include_columns and any(column.enabled for column in project.columns):
        labels.append("中間柱 / 共構柱")
    return labels


def _analysis_event_counts(project: ProjectState) -> dict[str, int]:
    analysis = project.analysis_import
    if analysis.events:
        counts = {"support": 0, "brace": 0, "floor": 0, "remove": 0, "other": 0}
        for event in analysis.events:
            if event.classification in counts:
                counts[event.classification] += 1
        return counts
    support_count = sum(len(stage.struts) for stage in analysis.stages)
    return {"support": support_count, "brace": 0, "floor": 0, "remove": 0, "other": 0}


def _analysis_import_overview_text(project: ProjectState) -> str:
    analysis = project.analysis_import
    if not (analysis.events or analysis.stages or analysis.warnings):
        return ""
    counts = _analysis_event_counts(project)
    parts = [
        f"本次匯入共整理水平支撐候選 {counts['support']} 筆",
        f"斜撐候選 {counts['brace']} 筆",
    ]
    if counts["floor"] > 0:
        parts.append(f"樓版事件 {counts['floor']} 筆")
    if counts["remove"] > 0:
        parts.append(f"拆撐事件 {counts['remove']} 筆")
    if counts["other"] > 0:
        parts.append(f"待人工判讀 {counts['other']} 筆")
    return "；".join(parts) + "。"


def _worst_check(results) -> CheckResult | None:
    all_checks = [
        *results.support_checks,
        *results.wale_checks,
        *results.brace_checks,
        *results.corner_brace_checks,
        *results.column_checks,
    ]
    if not all_checks:
        return None
    return max(all_checks, key=lambda item: item.utilization_ratio or -999)


def _summary_conclusion_lines(project: ProjectState, results) -> list[str]:
    module_labels = _active_module_labels(project)
    if not module_labels:
        return []
    status_counter = Counter(item.status for item in results.summary)
    lines = [
        f"本摘要表已就{'、'.join(module_labels)}之主要控制結果完成整理，相關細部驗算詳見後續章節及附件。",
        (
            f"本次檢核統計為 {_counted_status_phrase('OK', status_counter.get('OK', 0))}、"
            f"{_counted_status_phrase('Say~OK', status_counter.get('Say~OK', 0))}、"
            f"{_counted_status_phrase('NG', status_counter.get('NG', 0))}。"
        ),
    ]
    inactive_modules = _inactive_module_labels(project)
    if inactive_modules:
        lines.append(f"未納入檢討之{'、'.join(inactive_modules)}，不列入本次摘要與附件細部驗算。")
    return lines


def _inactive_module_labels(project: ProjectState) -> list[str]:
    labels: list[str] = []
    options = project.calculation_options
    if not (options.include_top_wales or options.include_bottom_wales):
        labels.append("橫擋")
    if not (options.include_top_braces or options.include_bottom_braces):
        labels.append("斜撐")
    if not options.include_corner_braces:
        labels.append("大角撐")
    if not any(column.enabled for column in project.columns):
        labels.append("中間柱 / 共構柱")
    return labels


def _overall_result_text(results) -> str:
    status_counter = Counter(item.status for item in results.summary)
    ng_count = status_counter.get("NG", 0)
    warn_count = status_counter.get("Say~OK", 0)
    if ng_count > 0:
        return f"本次檢核彙整尚有 NG {ng_count} 筆，宜優先修正後再行確認。"
    if warn_count > 0:
        return f"本次檢核彙整無 NG，惟尚有應注意 {warn_count} 筆，建議再行核對控制條件。"
    return "本次檢核彙整均為 OK，可作為後續整編與送審之依據。"


def _executive_summary_rows(project: ProjectState, results, *, concise_mode: bool) -> list[list[str]]:
    active_modules = _active_module_labels(project)
    inactive_modules = _inactive_module_labels(project)
    worst = _worst_check(results)
    rows = [
        ["納入檢討構件", "、".join(active_modules) if active_modules else "—"],
        ["整體判定", _overall_result_text(results)],
    ]
    if inactive_modules:
        rows.append(["未納入檢討構件", "、".join(inactive_modules)])
    if worst:
        rows.extend(
            [
                ["最不利構件", _detail_item_label(worst)],
                ["控制條件", worst.controlling_condition or "—"],
                [
                    "最大利用率",
                    (
                        f"{_fmt_short(worst.utilization_ratio)}（{_report_status_text(worst.status)}）"
                        if worst.utilization_ratio is not None
                        else _report_status_text(worst.status)
                    ),
                ],
            ]
        )
    rows.append(["附件編排方式", _report_mode_description(concise_mode)])
    return rows


def _result_conclusion_lines(project: ProjectState, results) -> list[str]:
    module_labels = _active_module_labels(project)
    lines: list[str] = []
    if module_labels:
        lines.append(f"本章已就{'、'.join(module_labels)}完成彙整，摘要表列示各層主要構件之控制結果。")
    inactive_modules = _inactive_module_labels(project)
    if inactive_modules:
        lines.append(f"未納入檢討之{'、'.join(inactive_modules)}，已自本章摘要與附件內容中略去。")
    worst = _worst_check(results)
    if worst and worst.utilization_ratio is not None:
        lines.append(
            (
                f"其中較不利之控制項為「{_detail_item_label(worst)}」之「{worst.controlling_condition}」，"
                f"最大利用率約為 {_fmt_short(worst.utilization_ratio)}，其細部計算詳後附附件。"
            )
        )
    if results.warnings:
        lines.append(f"另有判讀註記 {len(results.warnings)} 則，出具報表前宜再逐一核對。")
    return lines


SUMMARY_COLUMN_OPTIONS: list[tuple[str, str]] = [
    ("support", "水平支撐"),
    ("wale", "橫擋"),
    ("brace", "斜撐"),
    ("corner", "大角撐"),
]


def _available_summary_columns(summary_items: list[SummaryItem]) -> list[tuple[str, str]]:
    available = {
        _summary_bucket(item.group)
        for item in summary_items
        if item.group != "柱構件" and _summary_bucket(item.group) != "other"
    }
    return [(bucket, label) for bucket, label in SUMMARY_COLUMN_OPTIONS if bucket in available]


def _level_summary_col_widths(column_count: int) -> list[float]:
    fixed_widths = [16 * mm, 18 * mm, 20 * mm]
    if column_count <= 3:
        return fixed_widths
    dynamic_count = column_count - 3
    dynamic_width = max(28, min(38, 126 / max(dynamic_count, 1))) * mm
    return [16 * mm, *([dynamic_width] * dynamic_count), 18 * mm, 20 * mm]


def _build_level_summary_rows(
    summary_items: list[SummaryItem],
    cell_style: ParagraphStyle,
) -> tuple[list[object], list[list[object]], list[str]]:
    grouped: dict[str, dict[str, list[SummaryItem]]] = defaultdict(lambda: defaultdict(list))
    for item in summary_items:
        if item.group == "柱構件":
            continue
        grouped[item.label][_summary_bucket(item.group)].append(item)

    summary_columns = _available_summary_columns(summary_items)
    headers: list[object] = ["層別", *(label for _, label in summary_columns), "最差比值", "總評"]
    rows: list[list[object]] = []
    statuses: list[str] = []
    for label in sorted(grouped.keys(), key=_layer_sort_key):
        buckets = grouped[label]
        all_items = [item for bucket_items in buckets.values() for item in bucket_items]
        overall_status = _combine_status([item.status for item in all_items])
        worst_ratio = max(
            (item.utilization_ratio for item in all_items if item.utilization_ratio is not None),
            default=None,
        )
        rows.append(
            [
                label,
                *[_summary_cell(buckets.get(bucket, []), cell_style) for bucket, _ in summary_columns],
                _fmt_short(worst_ratio),
                _report_status_text(overall_status),
            ]
        )
        statuses.append(overall_status)
    return headers, rows, statuses


def _build_level_summary_plain_rows(summary_items: list[SummaryItem]) -> tuple[list[str], list[list[str]]]:
    grouped: dict[str, dict[str, list[SummaryItem]]] = defaultdict(lambda: defaultdict(list))
    for item in summary_items:
        if item.group == "柱構件":
            continue
        grouped[item.label][_summary_bucket(item.group)].append(item)

    summary_columns = _available_summary_columns(summary_items)
    headers = ["層別", *(label for _, label in summary_columns), "最差比值", "總評"]
    rows: list[list[str]] = []
    for label in sorted(grouped.keys(), key=_layer_sort_key):
        buckets = grouped[label]
        all_items = [item for bucket_items in buckets.values() for item in bucket_items]
        overall_status = _combine_status([item.status for item in all_items])
        worst_ratio = max(
            (item.utilization_ratio for item in all_items if item.utilization_ratio is not None),
            default=None,
        )
        rows.append(
            [
                label,
                *[_summary_text(buckets.get(bucket, [])) for bucket, _ in summary_columns],
                _fmt_short(worst_ratio),
                _report_status_text(overall_status),
            ]
        )
    return headers, rows


def _build_column_summary_rows(
    checks: list[CheckResult],
    cell_style: ParagraphStyle,
) -> tuple[list[list[object]], list[str]]:
    rows: list[list[object]] = [["構件", "型號", "利用率", "狀態", "備註"]]
    statuses: list[str] = []
    for check in checks:
        note = _column_warning_summary(check)
        rows.append(
            [
                check.label,
                str(check.inputs.get("型號", "—")),
                _fmt_short(check.utilization_ratio),
                _report_status_text(check.status),
                Paragraph(str(note), cell_style),
            ]
        )
        statuses.append(check.status)
    return rows, statuses


def _build_column_summary_plain_rows(checks: list[CheckResult]) -> list[list[str]]:
    rows: list[list[str]] = []
    for check in checks:
        note = _column_warning_summary(check)
        rows.append(
            [
                check.label,
                str(check.inputs.get("型號", "—")),
                _fmt_short(check.utilization_ratio),
                _report_status_text(check.status),
                str(note),
            ]
        )
    return rows


def _build_control_rows(
    results,
    cell_style: ParagraphStyle,
) -> tuple[list[list[object]], list[str]]:
    rows: list[list[object]] = [["模組", "標籤", "控制條件", "控制值 / 允許值", "利用率", "狀態"]]
    statuses: list[str] = []
    all_checks = [
        *results.support_checks,
        *results.wale_checks,
        *results.brace_checks,
        *results.corner_brace_checks,
        *results.column_checks,
    ]
    flagged = [check for check in all_checks if check.status != "OK"]
    if flagged:
        selected = sorted(flagged, key=lambda item: item.utilization_ratio or -999, reverse=True)
    else:
        selected = sorted(all_checks, key=lambda item: item.utilization_ratio or -999, reverse=True)[:10]
    for check in selected:
        rows.append(
            [
                check.module_name,
                check.label,
                Paragraph(check.controlling_condition, cell_style),
                f"{_fmt_short(check.computed_value)} / {_fmt_short(check.allowable_value)}",
                _fmt_short(check.utilization_ratio),
                _report_status_text(check.status),
            ]
        )
        statuses.append(check.status)
    return rows, statuses


def _build_control_plain_rows(results) -> list[list[str]]:
    rows: list[list[str]] = []
    all_checks = [
        *results.support_checks,
        *results.wale_checks,
        *results.brace_checks,
        *results.corner_brace_checks,
        *results.column_checks,
    ]
    flagged = [check for check in all_checks if check.status != "OK"]
    if flagged:
        selected = sorted(flagged, key=lambda item: item.utilization_ratio or -999, reverse=True)
    else:
        selected = sorted(all_checks, key=lambda item: item.utilization_ratio or -999, reverse=True)[:10]
    for check in selected:
        rows.append(
            [
                check.module_name,
                check.label,
                check.controlling_condition,
                f"{_fmt_short(check.computed_value)} / {_fmt_short(check.allowable_value)}",
                _fmt_short(check.utilization_ratio),
                _report_status_text(check.status),
            ]
        )
    return rows


def _appendix_one_groups(results) -> list[tuple[str, list[CheckResult]]]:
    return [
        ("水平支撐細部檢核", _ordered_appendix_checks(results.support_checks)),
        ("橫擋細部檢核", _ordered_appendix_checks(results.wale_checks)),
        ("斜撐細部檢核", _ordered_appendix_checks(results.brace_checks)),
        ("大角撐細部檢核", _ordered_appendix_checks(results.corner_brace_checks)),
    ]


def _appendix_group_has_checks(groups: list[tuple[str, list[CheckResult]]], title: str) -> bool:
    return any(group_title == title and checks for group_title, checks in groups)


def _appendix_group_display_name(title: str) -> str:
    if "水平支撐" in title:
        return "支撐"
    if "橫擋" in title:
        return "橫擋"
    if "斜撐" in title:
        return "斜撐"
    if "大角撐" in title:
        return "大角撐"
    return title.replace("細部檢核", "")


def _appendix_one_component_names(groups: list[tuple[str, list[CheckResult]]]) -> list[str]:
    return [_appendix_group_display_name(title) for title, checks in groups if checks]


def _appendix_one_title(groups: list[tuple[str, list[CheckResult]]]) -> str:
    names = _appendix_one_component_names(groups)
    if not names:
        return "附件一：細部計算結果"
    return f"附件一：{'、'.join(names)}細部計算結果"


def _appendix_one_intro_text(groups: list[tuple[str, list[CheckResult]]]) -> str:
    names = _appendix_one_component_names(groups)
    if not names:
        return "本附件依序列示各構件之檢核內容、計算過程及結果。"
    joined = "、".join(names)
    return f"本附件依{joined}之順序分節列示，逐項說明各構件之檢核內容、計算過程及結果。"


def _detail_group_display_title(title: str) -> str:
    return title.replace("細部檢核", "").strip()


def _detail_group_compilation_text(*, concise_mode: bool, check_count: int) -> str:
    if concise_mode and check_count > 1:
        return "首筆列示完整驗算過程，其餘摘列關鍵值與判定"
    return "逐筆列示完整驗算過程"


def _detail_group_summary_rows(title: str, checks: list[CheckResult], *, concise_mode: bool) -> list[list[str]]:
    if not checks:
        return []
    status_counter = Counter(check.status for check in checks)
    rows = [
        ["構件名稱", _detail_group_display_title(title)],
        ["檢核件數", f"{len(checks)} 筆"],
        [
            "判定統計",
            (
                f"{_counted_status_phrase('OK', status_counter.get('OK', 0))}；"
                f"{_counted_status_phrase('Say~OK', status_counter.get('Say~OK', 0))}；"
                f"{_counted_status_phrase('NG', status_counter.get('NG', 0))}"
            ),
        ],
        ["編排原則", _detail_group_compilation_text(concise_mode=concise_mode, check_count=len(checks))],
    ]
    source_lines = _detail_group_source_lines(checks)
    if source_lines:
        rows.insert(1, ["適用規範", "\n".join(source_lines)])
    return rows


def _append_pdf_detail_group(
    elements: list[object],
    title: str,
    checks: list[CheckResult],
    basic: BasicParameters,
    styles,
    *,
    section_reference_title: str | None = None,
    concise_mode: bool = False,
) -> None:
    header_block: list[object] = [Spacer(1, 2 * mm)]
    header_block.append(Paragraph(f"<b>{escape(_detail_group_display_title(title))}</b>", styles["ZHLead"]))
    intro_text = _detail_group_intro_text(title)
    if intro_text:
        header_block.append(Spacer(1, 1 * mm))
        header_block.append(Paragraph(f"<b>{escape(intro_text)}</b>", styles["ZHLead"]))
    summary_rows = _detail_group_summary_rows(title, checks, concise_mode=concise_mode)
    if summary_rows:
        _append_pdf_common_parameter_table(header_block, "本節檢核摘要", summary_rows, styles)
    elements.append(KeepTogether(header_block))
    detailed_checks = checks[:1] if concise_mode and checks else checks
    summary_checks = checks[1:] if concise_mode else []
    for index, check in enumerate(detailed_checks, start=1):
        block = (
            _build_check_detail_block(
                check,
                basic,
                section_reference_title=section_reference_title,
                include_source=False,
            )
        )
        elements.append(Spacer(1, 1.5 * mm))
        elements.append(
            Paragraph(
                f"<b>{escape(_detail_item_heading_text(check, index))}</b>",
                styles["ZHBody"],
            )
        )
        elements.append(Paragraph(escape(_detail_item_meta_text(check)), styles["ZHBody"]))
        for heading, lines in block:
            prepared_lines = _format_detail_block_lines(heading, lines)
            text = "<br/>".join(escape(line) for line in prepared_lines)
            elements.append(Paragraph(f"<b>{escape(heading)}：</b><br/>{text}", styles["ZHBody"]))
    if summary_checks:
        elements.append(Spacer(1, 2 * mm))
        elements.append(Paragraph("<b>其餘項目關鍵摘要</b>", styles["ZHBody"]))
        summary_rows, summary_statuses = _build_concise_summary_rows(summary_checks, styles["ZHCell"])
        summary_table = Table(
            summary_rows,
            colWidths=[10 * mm, 22 * mm, 28 * mm, 102 * mm, 18 * mm],
            repeatRows=1,
        )
        summary_style = _table_style(header=True)
        _apply_status_styles(summary_style, summary_statuses, status_col=4)
        summary_table.setStyle(summary_style)
        elements.append(summary_table)


def _append_word_detail_group(
    document: Document,
    title: str,
    checks: list[CheckResult],
    basic: BasicParameters,
    *,
    section_reference_title: str | None = None,
    concise_mode: bool = False,
) -> None:
    _add_subheading(document, _detail_group_display_title(title))
    intro_text = _detail_group_intro_text(title)
    if intro_text:
        _add_lead_paragraph(document, intro_text)
    summary_rows = _detail_group_summary_rows(title, checks, concise_mode=concise_mode)
    if summary_rows:
        _append_word_common_parameter_table(document, "本節檢核摘要", summary_rows)
    detailed_checks = checks[:1] if concise_mode and checks else checks
    summary_checks = checks[1:] if concise_mode else []
    for index, check in enumerate(detailed_checks, start=1):
        _add_detail_item_heading(
            document,
            _detail_item_heading_text(check, index),
        )
        _add_body_paragraph(document, _detail_item_meta_text(check))
        block = _build_check_detail_block(
            check,
            basic,
            section_reference_title=section_reference_title,
            include_source=False,
        )
        for heading, lines in block:
            _add_detail_block_heading(document, f"{heading}：")
            for line in _format_detail_block_lines(heading, lines):
                _add_indented_body_paragraph(document, line)
    if summary_checks:
        _add_body_paragraph(document, "其餘項目關鍵摘要")
        _add_grid_table(
            document,
            ["項次", "層次", "型號", "關鍵數值", "結果"],
            _build_concise_summary_plain_rows(summary_checks),
            status_col=4,
            col_widths_cm=[0.9, 2.0, 2.6, 8.2, 1.6],
        )


def _pdf_image_flowable(path: Path, *, width_mm: float) -> RLImage:
    width_px, height_px = ImageReader(str(path)).getSize()
    width = width_mm * mm
    height = width * (height_px / max(width_px, 1))
    image = RLImage(str(path), width=width, height=height)
    image.hAlign = "CENTER"
    return image


def _add_word_picture(document: Document, path: Path, *, width_cm: float) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(path), width=Cm(width_cm))


def _append_pdf_support_diagram(elements: list[object], styles) -> None:
    path = _report_diagram_path("支撐.png")
    if not path:
        return
    elements.append(Spacer(1, 1.5 * mm))
    elements.append(Paragraph("支撐示意圖", styles["ZHBody"]))
    elements.append(_pdf_image_flowable(path, width_mm=128))


def _append_word_support_diagram(document: Document) -> None:
    path = _report_diagram_path("支撐.png")
    if not path:
        return
    _add_body_paragraph(document, "支撐示意圖")
    _add_word_picture(document, path, width_cm=12.8)


def _detail_group_intro_text(title: str) -> str:
    if "水平支撐" in title:
        return "支撐計算如下"
    if "橫擋" in title:
        return "橫檔計算如下"
    if "斜撐" in title:
        return "斜撐計算如下"
    if "大角撐" in title:
        return "大角撐計算如下"
    if "柱構件" in title:
        return "柱構件計算如下"
    return ""


def _detail_group_source_lines(checks: list[CheckResult]) -> list[str]:
    seen: set[str] = set()
    lines: list[str] = []
    for check in checks:
        text = _formula_source_text(check.formula_id)
        if text in seen:
            continue
        seen.add(text)
        lines.append(text)
    return lines


def _detail_group_overview_text(checks: list[CheckResult], *, concise_mode: bool) -> str:
    if not checks:
        return ""
    status_counter = Counter(check.status for check in checks)
    parts = [
        f"本節共列 {len(checks)} 筆",
        _counted_status_phrase("OK", status_counter.get("OK", 0)),
        _counted_status_phrase("Say~OK", status_counter.get("Say~OK", 0)),
        _counted_status_phrase("NG", status_counter.get("NG", 0)),
    ]
    if concise_mode and len(checks) > 1:
        parts.append("編排方式為首筆詳算，其餘重點摘要")
    else:
        parts.append("編排方式為逐筆詳算")
    return "；".join(parts) + "。"


def _report_mode_description(concise_mode: bool) -> str:
    if concise_mode:
        return "簡述版（每節首筆列示完整驗算過程，其餘以關鍵數值摘要列示）"
    return "詳細版（逐筆列示完整驗算過程）"


def _column_warning_summary(check: CheckResult, *, default: str | None = None) -> str:
    warnings = check.details.get("warnings", [])
    warning_list = [str(item) for item in warnings] if isinstance(warnings, list) else []
    compression_failed = any("壓" in item for item in warning_list)
    tension_failed = any("拉" in item for item in warning_list)
    if compression_failed and tension_failed:
        return "壓入、拉拔檢核未通過"
    if compression_failed:
        return "壓入檢核未通過"
    if tension_failed:
        return "拉拔檢核未通過"
    if warning_list:
        return "；".join(warning_list)
    return default or check.controlling_condition


def _draw_pdf_page_frame(canvas, doc, project: ProjectState, *, concise_mode: bool, include_header: bool = True) -> None:
    canvas.saveState()
    canvas.setStrokeColor(colors.HexColor("#CBD5E1"))
    canvas.setLineWidth(0.6)
    if include_header:
        canvas.line(doc.leftMargin, A4[1] - 14 * mm, A4[0] - doc.rightMargin, A4[1] - 14 * mm)
    canvas.line(doc.leftMargin, 14 * mm, A4[0] - doc.rightMargin, 14 * mm)

    if include_header:
        canvas.setFont("STSong-Light", 9)
        canvas.setFillColor(colors.HexColor("#334155"))
        canvas.drawString(doc.leftMargin, A4[1] - 11 * mm, "擋土支撐檢核計算書")

        canvas.setFont("STSong-Light", 8)
        canvas.setFillColor(colors.HexColor("#64748B"))
        project_name = project.metadata.name or "未命名專案"
        mode_label = "簡述版" if concise_mode else "詳細版"
        canvas.drawRightString(A4[0] - doc.rightMargin, A4[1] - 11 * mm, f"{project_name}｜{mode_label}")
    canvas.drawCentredString(A4[0] / 2, 9.5 * mm, f"第 {canvas.getPageNumber()} 頁")
    canvas.restoreState()


def _paragraph_cell(text: str, style) -> Paragraph:
    return Paragraph("<br/>".join(escape(line) for line in str(text).splitlines()), style)


def _format_detail_block_lines(heading: str, lines: list[str]) -> list[str]:
    if heading in {"檢核公式", "代入計算"}:
        return [f"({index}) {line}" for index, line in enumerate(lines, start=1)]
    return lines


def _ordered_appendix_checks(checks: list[CheckResult]) -> list[CheckResult]:
    return sorted(
        checks,
        key=lambda check: (
            _detail_label_sort_key(check.label),
            _detail_side_order(check.module_name),
            _detail_section_name(check),
        ),
    )


def _detail_label_sort_key(label: str) -> tuple[int, object]:
    normalized = str(label).replace("第", "").replace("層", "").strip()
    return _layer_sort_key(normalized)


def _detail_side_order(module_name: str) -> int:
    if str(module_name).startswith("上層"):
        return 0
    if str(module_name).startswith("下層"):
        return 1
    return 2


def _detail_item_label(check: CheckResult) -> str:
    module_name = str(check.module_name)
    if module_name.startswith("上層"):
        return f"上層{check.label}"
    if module_name.startswith("下層"):
        return f"下層{check.label}"
    return check.label


def _detail_item_heading_text(check: CheckResult, index: int) -> str:
    return f"{index}. {_detail_item_label(check)}"


def _detail_item_meta_text(check: CheckResult) -> str:
    return f"型號：{_detail_section_name(check)}；判定：{_report_status_text(check.status)}"


def _collect_unique_sections(checks: list[CheckResult]) -> list[dict[str, object]]:
    sections: dict[str, dict[str, object]] = {}
    for check in checks:
        section_name = _detail_section_name(check)
        if not section_name or section_name == "未選型號":
            continue
        if section_name in sections:
            continue
        details = check.details
        sections[section_name] = {
            "name": section_name,
            "depth_cm": details.get("section_depth_cm"),
            "bf_cm": details.get("section_flange_width_cm"),
            "tw_cm": details.get("section_web_thickness_cm"),
            "tf_cm": details.get("section_flange_thickness_cm"),
            "area_cm2": details.get("section_area_cm2"),
            "ix_cm4": details.get("section_ix_cm4"),
            "iy_cm4": details.get("section_iy_cm4"),
            "sx_cm3": details.get("section_sx_cm3"),
            "sy_cm3": details.get("section_sy_cm3"),
            "rx_cm": details.get("section_rx_cm"),
            "ry_cm": details.get("section_ry_cm"),
            "rt_cm": details.get("section_rt_cm"),
            "unit_weight": details.get("section_unit_weight_kgf_per_m"),
        }
    return [sections[name] for name in sorted(sections)]


def _append_pdf_section_summary_table(
    elements: list[object],
    checks: list[CheckResult],
    title: str,
    styles,
) -> None:
    rows = _collect_unique_sections(checks)
    if not rows:
        return
    elements.append(Spacer(1, 1.5 * mm))
    elements.append(Paragraph(f"<b>{escape(title)}</b>", styles["ZHBody"]))
    table_rows = [["型號", "A", "Ix", "Iy", "Sx", "Sy", "rx", "ry", "rt", "單位重"]]
    for row in rows:
        table_rows.append([
            str(row["name"]),
            _fmt_short(row["area_cm2"]),
            _fmt_short(row["ix_cm4"]),
            _fmt_short(row["iy_cm4"]),
            _fmt_short(row["sx_cm3"]),
            _fmt_short(row["sy_cm3"]),
            _fmt_short(row["rx_cm"]),
            _fmt_short(row["ry_cm"]),
            _fmt_short(row["rt_cm"]),
            _fmt_short(row["unit_weight"]),
        ])
    table = Table(
        table_rows,
        colWidths=[34 * mm, 15 * mm, 18 * mm, 18 * mm, 16 * mm, 16 * mm, 14 * mm, 14 * mm, 14 * mm, 18 * mm],
        repeatRows=1,
    )
    table.setStyle(_table_style(header=True))
    elements.append(table)


def _append_word_section_summary_table(
    document: Document,
    checks: list[CheckResult],
    title: str,
) -> None:
    rows = _collect_unique_sections(checks)
    if not rows:
        return
    _add_body_paragraph(document, title)
    _add_grid_table(
        document,
        ["型號", "A", "Ix", "Iy", "Sx", "Sy", "rx", "ry", "rt", "單位重"],
        [[
            str(row["name"]),
            _fmt_short(row["area_cm2"]),
            _fmt_short(row["ix_cm4"]),
            _fmt_short(row["iy_cm4"]),
            _fmt_short(row["sx_cm3"]),
            _fmt_short(row["sy_cm3"]),
            _fmt_short(row["rx_cm"]),
            _fmt_short(row["ry_cm"]),
            _fmt_short(row["rt_cm"]),
            _fmt_short(row["unit_weight"]),
        ] for row in rows],
    )


def _append_pdf_common_parameter_table(
    elements: list[object],
    title: str,
    rows: list[list[str]],
    styles,
) -> None:
    if not rows:
        return
    elements.append(Spacer(1, 1.5 * mm))
    elements.append(Paragraph(f"<b>{escape(title)}</b>", styles["ZHBody"]))
    table = Table(
        [[_paragraph_cell(key, styles["ZHCell"]), _paragraph_cell(value, styles["ZHCell"])] for key, value in rows],
        colWidths=[58 * mm, 122 * mm],
    )
    table.setStyle(_table_style())
    elements.append(table)


def _append_word_common_parameter_table(
    document: Document,
    title: str,
    rows: list[list[str]],
) -> None:
    if not rows:
        return
    _add_body_paragraph(document, title)
    _add_key_value_table(document, rows)


def _appendix_common_parameter_rows(
    basic: BasicParameters,
    *,
    appendix: str,
    wall_deduction: bool = True,
) -> list[list[str]]:
    rows = [
        ["鋼材彈性係數 E (tf/cm2)", _fmt_short(basic.e_tf_per_cm2)],
        ["鋼材降伏應力 Fy (tf/cm2)", _fmt_short(basic.fy_tf_per_cm2)],
        ["彎矩分配係數 Cm", _fmt_short(basic.cm_factor)],
        ["折減係數 ψ", _fmt_short(basic.psi_material)],
    ]
    if appendix == "one":
        rows.extend(
            [
                ["積載重 WL (tf/m)", _fmt_short(basic.surcharge_wl_tf_per_m)],
                [
                    "構件折減係數 αs / αw / αb",
                    f"{_fmt_short(basic.alpha_support)} / {_fmt_short(basic.alpha_wale)} / {_fmt_short(basic.alpha_brace)}",
                ],
                [
                    "橫擋牆體扣底模式",
                    "考慮" if wall_deduction else "不考慮",
                ],
            ]
        )
    else:
        rows.append(["柱構件折減係數 αp", _fmt_short(basic.alpha_column)])
    return rows


def _report_diagram_path(filename: str) -> Path | None:
    settings = get_settings()
    direct = settings.root_dir / filename
    if direct.exists():
        return direct
    if filename == "支撐.png":
        return _ensure_excel_support_diagram_path()
    return None


def _ensure_excel_support_diagram_path() -> Path | None:
    settings = get_settings()
    diagram_dir = settings.app_data_dir / "excel_media"
    diagram_dir.mkdir(parents=True, exist_ok=True)
    target = diagram_dir / "image30.png"
    if target.exists():
        return target
    if not settings.workbook_path.exists():
        return None
    try:
        with ZipFile(settings.workbook_path) as archive:
            target.write_bytes(archive.read("xl/media/image30.png"))
    except Exception:
        return None
    return target if target.exists() else None


def _summary_bucket(group: str) -> str:
    if "水平支撐" in group:
        return "support"
    if "橫擋" in group:
        return "wale"
    if "斜撐" in group:
        return "brace"
    if "角撐" in group:
        return "corner"
    return "other"


def _summary_cell(items: list[SummaryItem], cell_style: ParagraphStyle) -> object:
    if not items:
        return "—"
    items = sorted(items, key=lambda item: _group_order(item.group))
    lines = []
    for item in items:
        lines.append(escape(_summary_headline(item)))
        lines.append(escape(_summary_section_name(item)))
    return Paragraph("<br/>".join(lines), cell_style)


def _summary_text(items: list[SummaryItem]) -> str:
    if not items:
        return "—"
    items = sorted(items, key=lambda item: _group_order(item.group))
    lines = []
    for item in items:
        lines.append(_summary_headline(item))
        lines.append(_summary_section_name(item))
    return "\n".join(lines)


def _summary_headline(item: SummaryItem) -> str:
    prefix = _group_prefix(item.group)
    return f"{prefix} {_fmt_short(item.utilization_ratio)} {_report_status_text(item.status)}".strip()


def _summary_section_name(item: SummaryItem) -> str:
    return f"型號：{item.section_name}" if item.section_name else "型號：未選型號"


def _detail_section_name(check: CheckResult) -> str:
    section_name = str(check.inputs.get("型號", "—"))
    return section_name or "未選型號"


def _detail_condition_text(check: CheckResult) -> str:
    warnings = check.details.get("warnings", [])
    if isinstance(warnings, list) and warnings:
        return f"{check.controlling_condition}；{'; '.join(str(item) for item in warnings)}"
    return check.controlling_condition


def _detail_condition_html(check: CheckResult) -> str:
    return escape(_detail_condition_text(check))


def _build_check_detail_block(
    check: CheckResult,
    basic: BasicParameters,
    *,
    section_reference_title: str | None = None,
    include_source: bool = True,
) -> list[tuple[str, list[str]]]:
    source = [_formula_source_text(check.formula_id)]
    conditions = _input_parameter_lines(check)
    section_data = _section_property_reference_lines(check, section_reference_title)
    formulas, substitutions, results = _formula_detail_content(check, basic)
    blocks = []
    if include_source:
        blocks.append(("規範", source))
    blocks.extend(
        [
            ("已知條件", conditions),
            ("斷面資料", section_data),
            ("檢核公式", formulas),
            ("代入計算", substitutions),
            ("檢核結果", results),
        ]
    )
    return blocks


def _build_check_concise_block(check: CheckResult) -> list[tuple[str, list[str]]]:
    return [
        ("已知條件", _input_parameter_lines(check)),
        ("關鍵數值", _concise_metric_lines(check)),
        ("檢核結果", _concise_result_lines(check)),
    ]


def _build_concise_summary_rows(checks: list[CheckResult], cell_style) -> tuple[list[list[object]], list[str]]:
    rows: list[list[object]] = [["項次", "層次", "型號", "關鍵數值", "結果"]]
    statuses: list[str] = []
    for index, check in enumerate(checks, start=2):
        rows.append(
            [
                str(index),
                _paragraph_cell(_detail_item_label(check), cell_style),
                _paragraph_cell(_detail_section_name(check), cell_style),
                _paragraph_cell(_concise_metric_text(check), cell_style),
                _report_status_text(check.status),
            ]
        )
        statuses.append(check.status)
    return rows, statuses


def _build_concise_summary_plain_rows(checks: list[CheckResult]) -> list[list[str]]:
    return [
        [
            str(index),
            _detail_item_label(check),
            _detail_section_name(check),
            _concise_metric_text(check),
            _report_status_text(check.status),
        ]
        for index, check in enumerate(checks, start=2)
    ]


def _formula_source_text(formula_id: str) -> str:
    mapping = {
        "support_interaction": "鋼結構容許應力設計法規範及解說，第六章受壓構材與第八章構材承受組合力及扭矩相關條文。",
        "wale_bending_shear": "鋼結構容許應力設計法規範及解說，第七章撓曲構材與剪力條文；若啟用牆體扣底，另扣除連續壁 Mwc、Vwc 之抵抗能力。",
        "brace_interaction": "鋼結構容許應力設計法規範及解說，第六章受壓構材與第八章構材承受組合力及扭矩相關條文。",
        "corner_brace_interaction": "鋼結構容許應力設計法規範及解說，第六章受壓構材與第八章構材承受組合力及扭矩相關條文。",
        "column_interaction": "鋼結構容許應力設計法規範及解說，第六章受壓構材、第八章組合力條文；基礎承載部分併用既有柱腳檢核公式。",
    }
    return mapping.get(formula_id, "依本案採用規範及既定檢核公式辦理。")


def _formula_detail_content(
    check: CheckResult,
    basic: BasicParameters,
) -> tuple[list[str], list[str], list[str]]:
    if check.controlling_condition == "資料未完整":
        message = str(check.details.get("message", "資料未完整，請先補齊輸入值。"))
        formulas = ["本項尚未進入正式檢核，需先補齊必要輸入資料。"]
        substitutions = _format_inputs_as_lines(check.inputs)
        results = [message, f"目前狀態：{_report_status_text(check.status)}"]
        return formulas, substitutions, results

    if check.formula_id == "support_interaction":
        return _support_detail_content(check, basic)
    if check.formula_id == "wale_bending_shear":
        return _wale_detail_content(check, basic)
    if check.formula_id == "brace_interaction":
        return _brace_detail_content(check, basic)
    if check.formula_id == "corner_brace_interaction":
        return _corner_brace_detail_content(check, basic)
    if check.formula_id == "column_interaction":
        return _column_detail_content(check, basic)
    return ["依本案既定檢核公式辦理。"], _format_inputs_as_lines(check.inputs), [_generic_result_line(check)]


def _support_detail_content(
    check: CheckResult,
    basic: BasicParameters,
) -> tuple[list[str], list[str], list[str]]:
    inputs = check.inputs
    details = check.details
    formulas = [
        "N = N1 + N2",
        "fa = N / A",
        "Cc = sqrt(2 x pi^2 x E / Fy)",
        "Fa 依 KL/r 與 Cc 之關係，按規範相應公式取值後乘以 alpha_s",
        "w = wself + WL",
        "M = (wself + WL) x SL^2 / 8",
        "fbx = M x 100 / Sx",
        "依 Lb、lc 與斷面分類選取 Fbx，再乘 alpha_s",
        "R = [fa/Fa + Cm x fbx / {(1 - fa/Fex) x Fbx}] / psi <= 1.0",
    ]
    substitutions = [
        f"N = {_fmt_short(inputs.get('軸力 N1'))} + {_fmt_short(inputs.get('溫度荷重 N2'))} = {_fmt_short(details.get('total_force_t'))} tf",
        f"fa = {_fmt_short(details.get('total_force_t'))} / {_fmt_short(details.get('area_cm2'))} = {_fmt_short(details.get('axial_stress'))} tf/cm2",
        f"Cc = sqrt(2 x pi^2 x {_fmt_short(basic.e_tf_per_cm2)} / {_fmt_short(basic.fy_tf_per_cm2)}) = {_fmt_short(_cc_value(basic))}",
        f"KL/r = {_fmt_short(inputs.get('水平間距 SL'))} x 100 / ry = {_fmt_short(details.get('klr'))}",
        _axial_allowable_substitution_line(_numeric(details.get("klr")), basic, basic.alpha_support, _numeric(details.get("fa_allow"))),
        f"wself = {_fmt_short(_section_self_weight(check))} tf/m，WL = {_fmt_short(basic.surcharge_wl_tf_per_m)} tf/m，w = {_fmt_short(details.get('line_load'))} tf/m",
        f"lc = min(20 x bf / sqrt(Fy), 1400 / ((d/(bf x tf)) x Fy)) = {_fmt_short(details.get('lc_cm'))} cm",
        f"M = {_fmt_short(details.get('line_load'))} x {_fmt_short(inputs.get('水平間距 SL'))}^2 / 8 = {_fmt_short(details.get('moment_tf_m'))} tf-m",
        f"fbx = {_fmt_short(details.get('moment_tf_m'))} x 100 / {_fmt_short(details.get('section_sx_cm3'))} = {_fmt_short(details.get('fbx_stress'))} tf/cm2",
        _fbx_allowable_substitution_line(
            _numeric(inputs.get("水平間距 SL")) * 100.0 if _numeric(inputs.get("水平間距 SL")) is not None else None,
            details,
            basic.fy_tf_per_cm2,
            basic.alpha_support,
            None,
            _numeric(details.get("fbx_allow")),
        ),
        f"Fby = {_fmt_short(details.get('fby_allow'))} tf/cm2",
        f"Fex = {_fmt_short(details.get('fex'))} tf/cm2，Fey = {_fmt_short(details.get('fey'))} tf/cm2",
        f"Cm = {_fmt_short(basic.cm_factor)}，psi = {_fmt_short(basic.psi_material)}",
    ]
    results = [
        f"交互作用比 R = {_fmt_short(check.utilization_ratio)}，允許值 = 1.000，判定 = {_report_status_text(check.status)}",
    ]
    return formulas, substitutions, results


def _wale_detail_content(
    check: CheckResult,
    basic: BasicParameters,
) -> tuple[list[str], list[str], list[str]]:
    inputs = check.inputs
    details = check.details
    formulas = [
        "M = max(Ww x Lw^2 / 10 - Mwc, 0)",
        "V = max(Ww x Lw / 2 - Vwc, 0)",
        "fbx = M x 100 / (Sx x n)",
        "fv = V / (d x tw x n)",
        "依 Lb、lc 與斷面分類選取 Fbx，再乘 alpha_w x psi",
        "Fv,allow = 0.4 x Fy x alpha_w x psi",
        "R = max[fbx/(Fbx x alpha_w x psi), fv/(0.4Fy x alpha_w x psi)] <= 1.0",
    ]
    substitutions = [
        f"Ww = {_fmt_short(inputs.get('線載重 Ww'))} tf/m，Lw = {_fmt_short(inputs.get('跨度 Lw'))} m，n = {_fmt_short(inputs.get('支數'))}",
        f"Mwc = {_fmt_short(details.get('wall_moment_strength'))} tf-m，Vwc = {_fmt_short(details.get('wall_shear_strength'))} tf",
        f"lc = min(20 x bf / sqrt(Fy), 1400 / ((d/(bf x tf)) x Fy)) = {_fmt_short(details.get('lc_cm'))} cm",
        f"M = max({_fmt_short(inputs.get('線載重 Ww'))} x {_fmt_short(inputs.get('跨度 Lw'))}^2 / 10 - {_fmt_short(details.get('wall_moment_strength'))}, 0) = {_fmt_short(details.get('moment_tf_m'))} tf-m",
        f"V = max({_fmt_short(inputs.get('線載重 Ww'))} x {_fmt_short(inputs.get('跨度 Lw'))} / 2 - {_fmt_short(details.get('wall_shear_strength'))}, 0) = {_fmt_short(details.get('shear_tf'))} tf",
        f"fbx = {_fmt_short(details.get('fbx_stress'))} tf/cm2",
        _fbx_allowable_substitution_line(
            _numeric(inputs.get("跨度 Lw")) * 100.0 if _numeric(inputs.get("跨度 Lw")) is not None else None,
            details,
            basic.fy_tf_per_cm2,
            basic.alpha_wale,
            basic.psi_material,
            (_numeric(details.get("fbx_allow")) or 0) * basic.alpha_wale * basic.psi_material,
        ),
        f"fv = {_fmt_short(details.get('fv_stress'))} tf/cm2，Fv,allow = 0.4 x {_fmt_short(basic.fy_tf_per_cm2)} x {_fmt_short(basic.alpha_wale)} x {_fmt_short(basic.psi_material)} = {_fmt_short((details.get('fv_allow') or 0) * basic.alpha_wale * basic.psi_material)} tf/cm2",
        f"彎矩比 = {_fmt_short(details.get('bending_ratio'))}，剪力比 = {_fmt_short(details.get('shear_ratio'))}，alpha_w = {_fmt_short(basic.alpha_wale)}，psi = {_fmt_short(basic.psi_material)}",
    ]
    results = [
        f"控制模式 = {check.controlling_condition}",
        f"利用率 R = {_fmt_short(check.utilization_ratio)}，判定 = {_report_status_text(check.status)}",
    ]
    return formulas, substitutions, results


def _brace_detail_content(
    check: CheckResult,
    basic: BasicParameters,
) -> tuple[list[str], list[str], list[str]]:
    inputs = check.inputs
    details = check.details
    formulas = [
        "L3 = (L1 + L2) / 2",
        "Lb = L1 / cos(theta)",
        "N = Ww x L3 / sin(theta)",
        "fa = N / A",
        "Cc = sqrt(2 x pi^2 x E / Fy)",
        "Fa 依 KL/r 與 Cc 之關係，按規範相應公式取值後乘以 alpha_b",
        "M = wself x Lb^2 / 8",
        "依 Lb、lc 與斷面分類選取 Fbx，再乘 alpha_b",
        "R = [fa/Fa + Cm x fbx / {(1 - fa/Fex) x Fbx}] / psi <= 1.0",
    ]
    substitutions = [
        f"L3 = ({_fmt_short(inputs.get('L1'))} + {_fmt_short(inputs.get('L2'))}) / 2 = {_fmt_short(details.get('l3_m'))} m",
        f"Lb = {_fmt_short(inputs.get('L1'))} / cos({_fmt_short(inputs.get('θ'))}) = {_fmt_short(details.get('lb_m'))} m",
        f"N = {_fmt_short(inputs.get('Ww'))} x {_fmt_short(details.get('l3_m'))} / sin({_fmt_short(inputs.get('θ'))}) = {_fmt_short(details.get('axial_force_t'))} tf",
        f"fa = {_fmt_short(details.get('axial_force_t'))} / {_fmt_short(details.get('section_area_cm2'))} = {_fmt_short(details.get('fa_value'))} tf/cm2",
        f"Cc = sqrt(2 x pi^2 x {_fmt_short(basic.e_tf_per_cm2)} / {_fmt_short(basic.fy_tf_per_cm2)}) = {_fmt_short(_cc_value(basic))}",
        _axial_allowable_substitution_line(_numeric(details.get("klr")), basic, basic.alpha_brace, _numeric(details.get("fa_allow"))),
        f"wself = {_fmt_short(details.get('self_weight_tf_per_m'))} tf/m，lc = {_fmt_short(details.get('lc_cm'))} cm",
        f"M = {_fmt_short(details.get('self_weight_tf_per_m'))} x {_fmt_short(details.get('lb_m'))}^2 / 8 = {_fmt_short(details.get('moment_tf_m'))} tf-m",
        f"fbx = {_fmt_short(details.get('fbx_stress'))} tf/cm2",
        _fbx_allowable_substitution_line(
            _numeric(details.get("lb_m")) * 100.0 if _numeric(details.get("lb_m")) is not None else None,
            details,
            basic.fy_tf_per_cm2,
            basic.alpha_brace,
            None,
            _numeric(details.get("fbx_allow")),
        ),
        f"Fex = {_fmt_short(details.get('fex'))} tf/cm2，Fey = {_fmt_short(details.get('fey'))} tf/cm2，psi = {_fmt_short(basic.psi_material)}",
    ]
    results = [
        f"交互作用比 R = {_fmt_short(check.utilization_ratio)}，允許值 = 1.000，判定 = {_report_status_text(check.status)}",
    ]
    return formulas, substitutions, results


def _corner_brace_detail_content(
    check: CheckResult,
    basic: BasicParameters,
) -> tuple[list[str], list[str], list[str]]:
    inputs = check.inputs
    details = check.details
    formulas = [
        "fa = Na / A",
        "Cc = sqrt(2 x pi^2 x E / Fy)",
        "Fa 依 KL/r 與 Cc 之關係，按規範相應公式取值後乘以 psi",
        "M = wself x La^2 / 8",
        "fbx = M x 100 / Sx",
        "依 Lb、lc 與斷面分類選取 Fbx，再乘 psi",
        "R = [fa/Fa + Cm x fbx / {(1 - fa/Fex) x Fbx}] <= 1.0",
    ]
    substitutions = [
        f"Na = {_fmt_short(inputs.get('軸力 Na'))} tf，La = {_fmt_short(inputs.get('長度 La'))} m",
        f"fa = {_fmt_short(details.get('axial_force_t'))} / {_fmt_short(details.get('section_area_cm2'))} = {_fmt_short(details.get('fa_value'))} tf/cm2",
        f"Cc = sqrt(2 x pi^2 x {_fmt_short(basic.e_tf_per_cm2)} / {_fmt_short(basic.fy_tf_per_cm2)}) = {_fmt_short(_cc_value(basic))}",
        _axial_allowable_substitution_line(_numeric(details.get("klr")), basic, basic.psi_material, _numeric(details.get("fa_allow")), factor_label="psi"),
        f"wself = {_fmt_short(details.get('self_weight_tf_per_m'))} tf/m，lc = {_fmt_short(details.get('lc_cm'))} cm",
        f"M = {_fmt_short(details.get('self_weight_tf_per_m'))} x {_fmt_short(inputs.get('長度 La'))}^2 / 8 = {_fmt_short(details.get('moment_tf_m'))} tf-m",
        f"fbx = {_fmt_short(details.get('fbx_stress'))} tf/cm2",
        _fbx_allowable_substitution_line(
            _numeric(inputs.get("長度 La")) * 100.0 if _numeric(inputs.get("長度 La")) is not None else None,
            details,
            basic.fy_tf_per_cm2,
            basic.psi_material,
            None,
            _numeric(details.get("fbx_allow")),
            factor_label="psi",
        ),
        f"Fby = {_fmt_short(details.get('fby_allow'))} tf/cm2，Fex = {_fmt_short(details.get('fex'))} tf/cm2，Fey = {_fmt_short(details.get('fey'))} tf/cm2",
    ]
    results = [
        f"交互作用比 R = {_fmt_short(check.utilization_ratio)}，允許值 = 1.000，判定 = {_report_status_text(check.status)}",
    ]
    return formulas, substitutions, results


def _column_detail_content(
    check: CheckResult,
    basic: BasicParameters,
) -> tuple[list[str], list[str], list[str]]:
    inputs = check.inputs
    details = check.details
    warnings = details.get("warnings", [])
    warning_text = _column_warning_summary(check, default="無額外註記")
    formulas = [
        "N = N1 + N2 + N3 + N4",
        "PT = max(N3 - N4 - N2 - N1, 0)",
        "beta = ((Kh x B) / (4 x E x I))^(1/4)，l0 = 1 / beta",
        "Cc = sqrt(2 x pi^2 x E / Fy)",
        "Fa 依 KL/r 與 Cc 之關係，按規範相應公式取值後乘以 alpha_p",
        "Fbx = 0.66 x Fy x alpha_p，Fby = 0.75 x Fy x alpha_p",
        "Mx = N3 x ex，My = N x ey",
        "R = [fa/Fa + Cmx x fbx/{(1 - fa/Fex) x Fbx} + Cmy x fby/{(1 - fa/Fey) x Fby}] / psi <= 1.0",
        "Qskin = Σ(qs x Li x 周長)，Qb = qb x Ab",
        "Qc,allow = (Qskin + Qb) / FS壓入",
        "Qt,allow = Qskin / FS拉拔 + Wpile",
        "壓入比 = N / Qc,allow，拉拔比 = PT / Qt,allow",
    ]
    substitutions = [
        f"N = {_fmt_short(inputs.get('N1'))} + {_fmt_short(inputs.get('N2'))} + {_fmt_short(inputs.get('N3'))} + {_fmt_short(inputs.get('N4'))} = {_fmt_short(inputs.get('N'))} tf",
        f"PT = max({_fmt_short(inputs.get('N3'))} - {_fmt_short(inputs.get('N4'))} - {_fmt_short(inputs.get('N2'))} - {_fmt_short(inputs.get('N1'))}, 0) = {_fmt_short(inputs.get('PT'))} tf",
        f"Ab = {_fmt_short(details.get('foundation_area_cm2'))} cm2，周長 U = {_fmt_short(details.get('foundation_perimeter_cm'))} cm，有效埋置深度 = {_fmt_short(details.get('effective_embedment_m'))} m",
        f"Cc = sqrt(2 x pi^2 x {_fmt_short(basic.e_tf_per_cm2)} / {_fmt_short(basic.fy_tf_per_cm2)}) = {_fmt_short(_cc_value(basic))}",
        f"fa = {_fmt_short(details.get('fa_value'))} tf/cm2",
        _axial_allowable_substitution_line(_numeric(details.get("klr_y")), basic, basic.alpha_column, _numeric(details.get("fa_allow"))),
        f"beta = (({_fmt_short(details.get('kh_kg_per_cm3'))} x {_fmt_short(details.get('foundation_size_x_m'))} x 100) / (4 x {_fmt_short(basic.e_tf_per_cm2)} x {_fmt_short(details.get('section_iy_cm4'))}))^(1/4) = {_fmt_short(details.get('beta'))}，l0 = {_fmt_short(details.get('l0_m'))} m，未支撐長度 = {_fmt_short(details.get('unsupported_length_m'))} m",
        f"KL/rx = {_fmt_short(details.get('klr_x'))}，KL/ry = {_fmt_short(details.get('klr_y'))}",
        f"Fbx = 0.66 x {_fmt_short(basic.fy_tf_per_cm2)} x {_fmt_short(basic.alpha_column)} = {_fmt_short(details.get('fbx_allow'))} tf/cm2；Fby = 0.75 x {_fmt_short(basic.fy_tf_per_cm2)} x {_fmt_short(basic.alpha_column)} = {_fmt_short(details.get('fby_allow'))} tf/cm2",
        f"ex = {_fmt_short(details.get('e_x_m'))} m，ey = {_fmt_short(details.get('e_y_m'))} m，Mx = {_fmt_short(details.get('mx_tf_m'))} tf-m，My = {_fmt_short(details.get('my_tf_m'))} tf-m",
        f"fbx = {_fmt_short(details.get('fbx_value'))} tf/cm2，Fbx = {_fmt_short(details.get('fbx_allow'))} tf/cm2，Fex = {_fmt_short(details.get('fex'))} tf/cm2",
        f"fby = {_fmt_short(details.get('fby_value'))} tf/cm2，Fby = {_fmt_short(details.get('fby_allow'))} tf/cm2，Fey = {_fmt_short(details.get('fey'))} tf/cm2",
        f"Qskin(壓入) = {_fmt_short(details.get('compression_skin_t'))} tf，Qb = {_fmt_short(details.get('compression_tip_t'))} tf，Qc,allow = ({_fmt_short(details.get('compression_skin_t'))} + {_fmt_short(details.get('compression_tip_t'))}) / {_fmt_short(details.get('compression_fs'))} = {_fmt_short(details.get('compression_capacity_t'))} tf",
        f"Qskin(拉拔) = {_fmt_short(details.get('tension_skin_t'))} tf，Wpile = {_fmt_short(details.get('tension_self_weight_t'))} tf，Qt,allow = {_fmt_short(details.get('tension_skin_t'))} / {_fmt_short(details.get('tension_fs'))} + {_fmt_short(details.get('tension_self_weight_t'))} = {_fmt_short(details.get('tension_capacity_t'))} tf",
        f"壓入比 = {_fmt_short(inputs.get('N'))} / {_fmt_short(details.get('compression_capacity_t'))} = {_fmt_short(details.get('compression_ratio'))}；拉拔比 = {_fmt_short(inputs.get('PT'))} / {_fmt_short(details.get('tension_capacity_t'))} = {_fmt_short(details.get('tension_ratio'))}",
    ]
    results = [
        f"柱交互作用比 R = {_fmt_short(check.utilization_ratio)}，允許值 = 1.000，判定 = {_report_status_text(check.status)}",
        f"壓入力比較：N = {_fmt_short(inputs.get('N'))} tf，Qc,allow = {_fmt_short(details.get('compression_capacity_t'))} tf，判定 = {'OK' if (details.get('compression_ratio') or 0) <= 1.0 else 'NG'}",
        f"拉拔力比較：PT = {_fmt_short(inputs.get('PT'))} tf，Qt,allow = {_fmt_short(details.get('tension_capacity_t'))} tf，判定 = {'OK' if (details.get('tension_ratio') or 0) <= 1.0 else 'NG'}",
        f"壓入 / 拉拔檢核說明：{warning_text}",
    ]
    return formulas, substitutions, results


def _format_inputs_as_lines(inputs: dict[str, object]) -> list[str]:
    if not inputs:
        return ["尚未提供可供檢核的輸入資料。"]
    return [f"{key} = {_fmt_short(value)}" for key, value in inputs.items()]


def _generic_result_line(check: CheckResult) -> str:
    return f"檢核結果：利用率 = {_fmt_short(check.utilization_ratio)}，狀態 = {_report_status_text(check.status)}"


def _cc_value(basic: BasicParameters) -> float:
    return (2.0 * 3.141592653589793**2 * basic.e_tf_per_cm2 / basic.fy_tf_per_cm2) ** 0.5


def _numeric(value: object) -> float | None:
    if value is None or value == "":
        return None
    try:
        return float(value)
    except (TypeError, ValueError):
        return None


def _axial_allowable_substitution_line(
    klr: float | None,
    basic: BasicParameters,
    factor: float,
    fa_allow: float | None,
    *,
    factor_label: str | None = None,
) -> str:
    if klr is None or fa_allow is None:
        return f"Fa = {_fmt_short(fa_allow)} tf/cm2"
    cc = _cc_value(basic)
    ratio = klr / cc if cc else 0.0
    factor_name = factor_label or f"{factor:.2f}"
    if klr < cc:
        return (
            f"本案 KL/r = {_fmt_short(klr)}，小於 Cc = {_fmt_short(cc)}，故 Fa 採短柱規定計算："
            f"Fa = [{{1 - ({_fmt_short(ratio)})^2 / 2}} x {_fmt_short(basic.fy_tf_per_cm2)}] / "
            f"[5/3 + 3 x {_fmt_short(ratio)} / 8 - ({_fmt_short(ratio)})^3 / 8] x {factor_name} = {_fmt_short(fa_allow)} tf/cm2"
        )
    return (
        f"本案 KL/r = {_fmt_short(klr)}，不小於 Cc = {_fmt_short(cc)}，故 Fa 採長柱規定計算："
        f"Fa = [12 x pi^2 x {_fmt_short(basic.e_tf_per_cm2)} / (32 x ({_fmt_short(ratio)})^2)] x {factor_name} = {_fmt_short(fa_allow)} tf/cm2"
    )


def _fbx_allowable_substitution_line(
    lb_cm: float | None,
    details: dict[str, object],
    fy: float,
    factor: float,
    psi: float | None,
    final_value: float | None,
    *,
    factor_label: str | None = None,
) -> str:
    if lb_cm is None or final_value is None:
        return f"Fbx = {_fmt_short(final_value)} tf/cm2"
    lc_cm = _numeric(details.get("lc_cm")) or 0.0
    bf = _numeric(details.get("section_flange_width_cm")) or 0.0
    tf = _numeric(details.get("section_flange_thickness_cm")) or 0.0
    d = _numeric(details.get("section_depth_cm")) or 0.0
    rt = _numeric(details.get("section_rt_cm")) or 0.0
    section_class = str(details.get("section_class", ""))
    factor_name = factor_label or f"{factor:.2f}"
    multiplier_text = f" x {factor_name}"
    if psi is not None:
        multiplier_text += f" x {_fmt_short(psi)}"

    if lb_cm < lc_cm:
        if section_class in {"塑性斷面", "結實斷面"}:
            return f"因 Lb = {_fmt_short(lb_cm)} cm < lc = {_fmt_short(lc_cm)} cm，且為{section_class}，採 Fbx = 0.66 x {_fmt_short(fy)}{multiplier_text} = {_fmt_short(final_value)} tf/cm2"
        if section_class == "半結實斷面":
            base = fy * (0.79 - 0.0075 * bf / max(2.0 * tf, 1e-6) * math.sqrt(fy))
            return f"因 Lb = {_fmt_short(lb_cm)} cm < lc = {_fmt_short(lc_cm)} cm，且為半結實斷面，採 Fbx = {_fmt_short(base)}{multiplier_text} = {_fmt_short(final_value)} tf/cm2"
        return f"因 Lb = {_fmt_short(lb_cm)} cm < lc = {_fmt_short(lc_cm)} cm，且為細長肢材斷面，採 Fbx = 0.60 x {_fmt_short(fy)}{multiplier_text} = {_fmt_short(final_value)} tf/cm2"

    if section_class == "細長肢材斷面":
        return f"因斷面為細長肢材斷面，採 Fbx = 0.60 x {_fmt_short(fy)}{multiplier_text} = {_fmt_short(final_value)} tf/cm2"

    lb_rt = lb_cm / max(rt, 1e-6)
    branch_limit = math.sqrt(35800.0 / max(fy, 1e-6))
    if lb_rt > branch_limit:
        cand1 = 12000.0 / max(lb_rt**2, 1e-6)
        cand2 = 840.0 / max(lb_cm * d / max(bf * tf, 1e-6), 1e-6)
        base = min(max(cand1, cand2), 0.6 * fy)
        return f"因 Lb/rt = {_fmt_short(lb_rt)} > sqrt(35800/Fy) = {_fmt_short(branch_limit)}，採 Fbx = min(max({_fmt_short(cand1)}, {_fmt_short(cand2)}), 0.60Fy){multiplier_text} = {_fmt_short(final_value)} tf/cm2"
    base = (2.0 / 3.0 - fy * (lb_rt**2) / 107600.0) * fy
    base = min(base, 0.6 * fy)
    return f"因 Lb/rt = {_fmt_short(lb_rt)} <= sqrt(35800/Fy) = {_fmt_short(branch_limit)}，採 Fbx = min((2/3 - Fy(Lb/rt)^2/107600)Fy, 0.60Fy){multiplier_text} = {_fmt_short(final_value)} tf/cm2"


def _section_self_weight(check: CheckResult) -> float:
    unit_weight = check.details.get("section_unit_weight_kgf_per_m")
    if unit_weight is None:
        return 0.0
    return float(unit_weight) / 1000.0


def _section_property_lines(check: CheckResult) -> list[str]:
    details = check.details
    lines = [
        f"d = {_fmt_short(details.get('section_depth_cm'))} cm，bf = {_fmt_short(details.get('section_flange_width_cm'))} cm",
        f"tw = {_fmt_short(details.get('section_web_thickness_cm'))} cm，tf = {_fmt_short(details.get('section_flange_thickness_cm'))} cm",
        f"A = {_fmt_short(details.get('section_area_cm2'))} cm2，Sx = {_fmt_short(details.get('section_sx_cm3'))} cm3，Sy = {_fmt_short(details.get('section_sy_cm3'))} cm3",
        f"Ix = {_fmt_short(details.get('section_ix_cm4'))} cm4，Iy = {_fmt_short(details.get('section_iy_cm4'))} cm4",
        f"rx = {_fmt_short(details.get('section_rx_cm'))} cm，ry = {_fmt_short(details.get('section_ry_cm'))} cm，rt = {_fmt_short(details.get('section_rt_cm'))} cm",
        f"單位重 = {_fmt_short(details.get('section_unit_weight_kgf_per_m'))} kgf/m",
    ]
    if details.get("section_class"):
        lines.append(f"斷面分類 = {details.get('section_class')}")
    return lines


def _section_property_reference_lines(
    check: CheckResult,
    section_reference_title: str | None,
) -> list[str]:
    if not section_reference_title:
        return _section_property_lines(check)
    return [f"型鋼型號：{_detail_section_name(check)}（斷面性質詳前述「{section_reference_title}」）"]


def _input_parameter_lines(check: CheckResult) -> list[str]:
    inputs = check.inputs
    if check.formula_id == "support_interaction":
        return [
            f"N1 = {_fmt_short(inputs.get('軸力 N1'))} tf，N2 = {_fmt_short(inputs.get('溫度荷重 N2'))} tf，SL = {_fmt_short(inputs.get('水平間距 SL'))} m",
        ]
    if check.formula_id == "wale_bending_shear":
        return [
            f"Lw = {_fmt_short(inputs.get('跨度 Lw'))} m，SS = {_fmt_short(inputs.get('支撐間距 SS'))} m，Ww = {_fmt_short(inputs.get('線載重 Ww'))} tf/m，n = {_fmt_short(inputs.get('支數'))}",
        ]
    if check.formula_id == "brace_interaction":
        return [
            f"L1 = {_fmt_short(inputs.get('L1'))} m，L2 = {_fmt_short(inputs.get('L2'))} m，θ = {_fmt_short(inputs.get('θ'))} deg，Ww = {_fmt_short(inputs.get('Ww'))} tf/m",
        ]
    if check.formula_id == "corner_brace_interaction":
        return [
            f"La = {_fmt_short(inputs.get('長度 La'))} m，Na = {_fmt_short(inputs.get('軸力 Na'))} tf",
        ]
    if check.formula_id == "column_interaction":
        return [
            f"N1 = {_fmt_short(inputs.get('N1'))} tf，N2 = {_fmt_short(inputs.get('N2'))} tf，N3 = {_fmt_short(inputs.get('N3'))} tf，N4 = {_fmt_short(inputs.get('N4'))} tf",
            f"N = {_fmt_short(inputs.get('N'))} tf，PT = {_fmt_short(inputs.get('PT'))} tf",
            f"基礎型式 = {inputs.get('基礎型式', '—')}，基礎形狀 = {inputs.get('基礎形狀', '—')}，Bx = {_fmt_short(inputs.get('基礎尺寸 Bx'))} m，By = {_fmt_short(inputs.get('基礎尺寸 By'))} m",
            f"埋置深度 = {_fmt_short(inputs.get('埋置深度'))} m，FS壓入 = {_fmt_short(inputs.get('FS壓入'))}，FS拉拔 = {_fmt_short(inputs.get('FS拉拔'))}，Kh = {_fmt_short(inputs.get('Kh'))}",
        ]
    return _format_inputs_as_lines(inputs)


def _concise_metric_lines(check: CheckResult) -> list[str]:
    details = check.details
    if check.formula_id == "support_interaction":
        return [
            f"N = {_fmt_short(details.get('total_force_t'))} tf；fa / Fa = {_fmt_short(details.get('axial_stress'))} / {_fmt_short(details.get('fa_allow'))} tf/cm2",
            f"M = {_fmt_short(details.get('moment_tf_m'))} tf-m；fbx / Fbx = {_fmt_short(details.get('fbx_stress'))} / {_fmt_short(details.get('fbx_allow'))} tf/cm2",
        ]
    if check.formula_id == "wale_bending_shear":
        return [
            f"M / V = {_fmt_short(details.get('moment_tf_m'))} tf-m / {_fmt_short(details.get('shear_tf'))} tf",
            f"fbx / fv = {_fmt_short(details.get('fbx_stress'))} / {_fmt_short(details.get('fv_stress'))} tf/cm2",
            f"彎矩比 / 剪力比 = {_fmt_short(details.get('bending_ratio'))} / {_fmt_short(details.get('shear_ratio'))}",
        ]
    if check.formula_id == "brace_interaction":
        return [
            f"N = {_fmt_short(details.get('axial_force_t'))} tf；fa / Fa = {_fmt_short(details.get('fa_value'))} / {_fmt_short(details.get('fa_allow'))} tf/cm2",
            f"M = {_fmt_short(details.get('moment_tf_m'))} tf-m；fbx / Fbx = {_fmt_short(details.get('fbx_stress'))} / {_fmt_short(details.get('fbx_allow'))} tf/cm2",
        ]
    if check.formula_id == "corner_brace_interaction":
        return [
            f"Na = {_fmt_short(details.get('axial_force_t'))} tf；fa / Fa = {_fmt_short(details.get('fa_value'))} / {_fmt_short(details.get('fa_allow'))} tf/cm2",
            f"M = {_fmt_short(details.get('moment_tf_m'))} tf-m；fbx / Fbx = {_fmt_short(details.get('fbx_stress'))} / {_fmt_short(details.get('fbx_allow'))} tf/cm2",
        ]
    if check.formula_id == "column_interaction":
        return [
            f"N / PT = {_fmt_short(check.inputs.get('N'))} / {_fmt_short(check.inputs.get('PT'))} tf",
            f"fa / Fa = {_fmt_short(details.get('fa_value'))} / {_fmt_short(details.get('fa_allow'))} tf/cm2",
            f"Qc / Qt = {_fmt_short(details.get('compression_capacity_t'))} / {_fmt_short(details.get('tension_capacity_t'))} tf",
            f"壓入比 / 拉拔比 = {_fmt_short(details.get('compression_ratio'))} / {_fmt_short(details.get('tension_ratio'))}",
        ]
    return [f"利用率 = {_fmt_short(check.utilization_ratio)}"]


def _concise_result_lines(check: CheckResult) -> list[str]:
    details = check.details
    if check.formula_id == "wale_bending_shear":
        return [
            f"控制模式 = {check.controlling_condition}",
            f"R = {_fmt_short(check.utilization_ratio)}；{_report_status_text(check.status)}",
        ]
    if check.formula_id == "column_interaction":
        warning_text = _column_warning_summary(check, default="無額外註記")
        return [
            f"柱交互作用比 R = {_fmt_short(check.utilization_ratio)}；{_report_status_text(check.status)}",
            f"壓入 / 拉拔檢核說明：{warning_text}",
        ]
    return [f"交互作用比 R = {_fmt_short(check.utilization_ratio)}；{_report_status_text(check.status)}"]


def _concise_metric_text(check: CheckResult) -> str:
    return "\n".join(_concise_metric_lines(check) + _concise_result_lines(check))


def _report_status_text(status: str) -> str:
    if status == "Say~OK":
        return "應注意"
    return status


def _counted_status_phrase(status: str, count: int) -> str:
    return f"{_report_status_text(status)} {count} 筆"


def _collect_word_soils(project: ProjectState) -> list[dict[str, object]]:
    if project.analysis_import.soils:
        return [
            {
                "index": soil.index,
                "name": soil.name,
                "depth_m": soil.depth_m,
                "thickness_m": soil.thickness_m,
                "unit_weight_t_per_m3": soil.unit_weight_t_per_m3,
                "phi_deg": soil.phi_deg,
                "cohesion_t_per_m2": soil.cohesion_t_per_m2,
                "su_t_per_m2": soil.su_t_per_m2,
                "kh_t_per_m3": soil.kh_t_per_m3,
                "soil_type": soil.soil_type,
            }
            for soil in project.analysis_import.soils
        ]
    if not project.columns:
        return []
    return [
        {
            "index": soil.index,
            "name": soil.name,
            "depth_m": soil.depth_m,
            "thickness_m": soil.thickness_m,
            "unit_weight_t_per_m3": None,
            "phi_deg": None,
            "cohesion_t_per_m2": None,
            "su_t_per_m2": soil.su_t_per_m2,
            "kh_t_per_m3": None,
            "soil_type": soil.soil_type,
        }
        for soil in project.columns[0].soil_layers
    ]


def _build_soil_plain_rows(soils: list[dict[str, object]]) -> list[list[str]]:
    rows: list[list[str]] = []
    for soil in soils:
        rows.append(
            [
                str(soil.get("index", "—")),
                str(soil.get("name", "—")),
                _fmt_short(soil.get("depth_m")),
                _fmt_short(soil.get("thickness_m")),
                _fmt_short(soil.get("unit_weight_t_per_m3")),
                _fmt_short(soil.get("phi_deg")),
                _fmt_short(soil.get("cohesion_t_per_m2")),
                _fmt_short(soil.get("su_t_per_m2")),
                _fmt_short(soil.get("kh_t_per_m3")),
                str(soil.get("soil_type", "—")),
            ]
        )
    return rows


def _build_stage_plain_rows(analysis) -> list[list[str]]:
    rows: list[list[str]] = []
    for stage in analysis.stages:
        note = []
        if stage.struts:
            note.append("含支撐分析")
        if stage.water_level_m is not None:
            note.append("含水位設定")
        rows.append(
            [
                stage.label,
                _fmt_short(stage.excavation_depth_m),
                _fmt_short(stage.water_level_m),
                str(len(stage.struts)),
                "、".join(note) if note else "—",
            ]
        )
    return rows


def _group_order(group: str) -> int:
    if group.startswith("上"):
        return 0
    if group.startswith("下"):
        return 1
    return 2


def _group_prefix(group: str) -> str:
    if group.startswith("上"):
        return "上"
    if group.startswith("下"):
        return "下"
    return ""


def _active_single_support_side(project: ProjectState) -> str | None:
    if project.calculation_options.include_top_supports and not project.calculation_options.include_bottom_supports:
        return "top"
    if project.calculation_options.include_bottom_supports and not project.calculation_options.include_top_supports:
        return "bottom"
    return None


def _combine_status(statuses: list[str]) -> str:
    if any(status == "NG" for status in statuses):
        return "NG"
    if any(status == "Say~OK" for status in statuses):
        return "Say~OK"
    return "OK"


def _layer_sort_key(label: str) -> tuple[int, object]:
    try:
        return (0, int(label))
    except (TypeError, ValueError):
        return (1, label)


def _apply_status_styles(style: TableStyle, statuses: list[str], status_col: int) -> None:
    for row_index, status in enumerate(statuses, start=1):
        bg, fg = _status_palette(status)
        style.add("BACKGROUND", (status_col, row_index), (status_col, row_index), bg)
        style.add("TEXTCOLOR", (status_col, row_index), (status_col, row_index), fg)
        style.add("FONTNAME", (status_col, row_index), (status_col, row_index), "STSong-Light")


def _status_palette(status: str) -> tuple[colors.Color, colors.Color]:
    if status == "NG":
        return colors.HexColor("#fee2e2"), colors.HexColor("#991b1b")
    if status in {"Say~OK", "應注意"}:
        return colors.HexColor("#fef3c7"), colors.HexColor("#92400e")
    return colors.HexColor("#dcfce7"), colors.HexColor("#166534")


def _table_style(header: bool = False) -> TableStyle:
    style = [
        ("FONTNAME", (0, 0), (-1, -1), "STSong-Light"),
        ("FONTSIZE", (0, 0), (-1, -1), 9.2),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#cbd5e1")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f8fafc")]),
    ]
    if header:
        style.extend(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#dbeafe")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#1e3a8a")),
                ("FONTNAME", (0, 0), (-1, 0), "STSong-Light"),
                ("FONTSIZE", (0, 0), (-1, 0), 9.4),
            ]
        )
    else:
        style.append(("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#f8fafc")))
    return TableStyle(style)
