from __future__ import annotations

import base64
from copy import deepcopy
from dataclasses import dataclass
from datetime import datetime
import hashlib
import json
import math
from pathlib import Path
import re
from tempfile import TemporaryDirectory
from typing import Any, Literal
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase.pdfmetrics import registerFont
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from .config import get_settings
from .reporting import REPORT_TOOL_NAME, REPORT_TOOL_VERSION, report_document_metadata


ATTACHMENT_SCHEMA_VERSION = 1
ATTACHMENT_KIND = "rsc-v4-rsb-independent-calculation-attachment"
ATTACHMENT_FILE_PREFIX = "rsc-v4-rsb"
RSC_KIND = "excavation-reshore-member-capacity-calculation"
RSB_KIND = "excavation-reshore-end-bearing-evidence"
RSC_SCHEMA_VERSION = 4
RSB_SCHEMA_VERSION = 1
_FINGERPRINT_RE = re.compile(r"^(RSC|RSB)-[0-9A-F]{20}$")
_ATTACHMENT_FILE_RE = re.compile(
    r"^rsc-v4-rsb-[0-9A-F]{20}-[0-9A-F]{20}-\d{20}\.(?:pdf|docx)$"
)
_PDF_FONT_NAME = "RscEmbeddedCjk"
_PDF_FONT_CANDIDATES = (
    Path(r"C:\Windows\Fonts\NotoSansTC-VF.ttf"),
    Path(r"C:\Windows\Fonts\ARIALUNI.TTF"),
    Path("/usr/share/fonts/truetype/noto/NotoSansTC-Regular.ttf"),
    Path("/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.otf"),
    Path("/Library/Fonts/Arial Unicode.ttf"),
)


@dataclass(frozen=True)
class ReshoreCapacityAttachmentArtifact:
    path: Path
    report_kind: Literal["pdf", "docx"]
    document_status: str
    output_time: str
    rsc_calculation_fingerprint: str
    rsc_evidence_file_name: str
    rsc_evidence_file_sha256: str
    rsb_calculation_fingerprint: str
    rsb_evidence_file_name: str
    rsb_evidence_file_sha256: str
    attachment_file_sha256: str
    attachment_size_bytes: int


def attachment_download_name_allowed(file_name: str) -> bool:
    safe_name = Path(file_name).name
    if safe_name != file_name:
        return False
    if _ATTACHMENT_FILE_RE.fullmatch(safe_name):
        return True
    for suffix in (".canonical-render.evidence.json", ".formal-source.zip"):
        if not safe_name.endswith(suffix):
            continue
        base_name = f"{safe_name[:-len(suffix)]}.pdf"
        return _ATTACHMENT_FILE_RE.fullmatch(base_name) is not None
    return False


def build_reshore_capacity_attachment(
    calculation_response: dict[str, Any],
    *,
    report_kind: Literal["pdf", "docx"],
    output_at: datetime | None = None,
) -> ReshoreCapacityAttachmentArtifact:
    """Build one receiver-side formal attachment without mutating a source project."""
    if report_kind not in {"pdf", "docx"}:
        raise ValueError("接收端獨立計算附件只支援 PDF 或 DOCX。")
    context = _validated_attachment_context(calculation_response)
    output_at = output_at or datetime.now().astimezone()
    if output_at.tzinfo is None:
        output_at = output_at.astimezone()
    document_metadata = report_document_metadata(approved=True, output_at=output_at)
    settings = get_settings()
    output_dir = settings.reports_dir.resolve()
    output_dir.mkdir(parents=True, exist_ok=True)
    rsc_token = context["rscFingerprint"].removeprefix("RSC-")
    rsb_token = context["rsbFingerprint"].removeprefix("RSB-")
    file_name = (
        f"{ATTACHMENT_FILE_PREFIX}-{rsc_token}-{rsb_token}-"
        f"{output_at:%Y%m%d%H%M%S%f}.{report_kind}"
    )
    final_path = output_dir / file_name
    if final_path.exists():
        raise FileExistsError(f"接收端附件已存在，不得覆寫：{file_name}")

    try:
        with TemporaryDirectory(prefix=".rsc-v4-rsb-", dir=output_dir) as temp_dir:
            temp_path = Path(temp_dir) / file_name
            if report_kind == "pdf":
                _write_pdf(temp_path, context, document_metadata)
            else:
                _write_docx(temp_path, context, document_metadata)
            if not temp_path.is_file() or temp_path.stat().st_size <= 0:
                raise RuntimeError("接收端獨立計算附件沒有產生有效檔案。")
            temp_path.replace(final_path)
        artifact_bytes = final_path.read_bytes()
        return ReshoreCapacityAttachmentArtifact(
            path=final_path,
            report_kind=report_kind,
            document_status=document_metadata["document_status"],
            output_time=document_metadata["output_time"],
            rsc_calculation_fingerprint=context["rscFingerprint"],
            rsc_evidence_file_name=context["rscEvidence"]["fileName"],
            rsc_evidence_file_sha256=context["rscEvidence"]["fileSha256"],
            rsb_calculation_fingerprint=context["rsbFingerprint"],
            rsb_evidence_file_name=context["rsbEvidence"]["fileName"],
            rsb_evidence_file_sha256=context["rsbEvidence"]["fileSha256"],
            attachment_file_sha256=hashlib.sha256(artifact_bytes).hexdigest(),
            attachment_size_bytes=len(artifact_bytes),
        )
    except Exception:
        final_path.unlink(missing_ok=True)
        raise


def _canonical_json(value: Any) -> str:
    return json.dumps(value, ensure_ascii=False, sort_keys=True, separators=(",", ":"), allow_nan=False)


def _record_fingerprint(record: dict[str, Any], prefix: str) -> str:
    payload = deepcopy(record)
    payload.pop("calculationFingerprint", None)
    digest = hashlib.sha256(_canonical_json(payload).encode("utf-8")).hexdigest()[:20].upper()
    return f"{prefix}-{digest}"


def _decode_json_evidence(
    envelope: Any,
    *,
    expected_kind: str,
    expected_schema_version: int,
    fingerprint_prefix: str,
) -> tuple[dict[str, Any], bytes]:
    if not isinstance(envelope, dict):
        raise ValueError("計算附件缺少 JSON 證據封套。")
    if envelope.get("mediaType") != "application/json" or envelope.get("contentEncoding") != "base64":
        raise ValueError("計算附件的 JSON 證據格式不受支援。")
    file_name = str(envelope.get("fileName", ""))
    if not file_name or Path(file_name).name != file_name:
        raise ValueError("計算附件的 JSON 證據檔名不安全。")
    try:
        evidence_bytes = base64.b64decode(str(envelope.get("contentBase64", "")), validate=True)
    except Exception as exc:
        raise ValueError("計算附件的 JSON 證據不是有效 Base64。") from exc
    actual_sha256 = hashlib.sha256(evidence_bytes).hexdigest()
    if actual_sha256 != str(envelope.get("fileSha256", "")).lower():
        raise ValueError("計算附件的 JSON 證據 SHA-256 不符。")
    try:
        record = json.loads(evidence_bytes)
    except (UnicodeDecodeError, json.JSONDecodeError) as exc:
        raise ValueError("計算附件的 JSON 證據無法解析。") from exc
    if not isinstance(record, dict):
        raise ValueError("計算附件的 JSON 證據根節點必須是物件。")
    if evidence_bytes != (_canonical_json(record) + "\n").encode("utf-8"):
        raise ValueError("計算附件的 JSON 證據不是 canonical bytes。")
    if record.get("kind") != expected_kind or record.get("schemaVersion") != expected_schema_version:
        raise ValueError("計算附件引用了不受支援的證據版本或種類。")
    fingerprint = str(record.get("calculationFingerprint", ""))
    if not _FINGERPRINT_RE.fullmatch(fingerprint) or not fingerprint.startswith(f"{fingerprint_prefix}-"):
        raise ValueError("計算附件的證據指紋格式不正確。")
    if fingerprint != _record_fingerprint(record, fingerprint_prefix):
        raise ValueError("計算附件的證據指紋無法重播。")
    if envelope.get("documentReference") != fingerprint:
        raise ValueError("計算附件的證據文件參照與指紋不符。")
    return record, evidence_bytes


def _validated_attachment_context(calculation_response: Any) -> dict[str, Any]:
    if not isinstance(calculation_response, dict):
        raise ValueError("接收端附件需要受控的 RSC v4 計算回應。")
    rsc_record, _ = _decode_json_evidence(
        calculation_response.get("evidence"),
        expected_kind=RSC_KIND,
        expected_schema_version=RSC_SCHEMA_VERSION,
        fingerprint_prefix="RSC",
    )
    if calculation_response.get("calculation") != rsc_record:
        raise ValueError("接收端附件的 RSC 回應與精確 JSON bytes 不符。")
    if rsc_record.get("input", {}).get("endBearingMode") != "centered_rectangular_plate":
        raise ValueError("只有啟用上下端直接承壓的 RSC v4 才能產生獨立計算附件。")
    rsb_record, _ = _decode_json_evidence(
        calculation_response.get("bearingEvidence"),
        expected_kind=RSB_KIND,
        expected_schema_version=RSB_SCHEMA_VERSION,
        fingerprint_prefix="RSB",
    )
    source = rsc_record.get("source")
    rsb_source = rsb_record.get("source")
    if not isinstance(source, dict) or not isinstance(rsb_source, dict):
        raise ValueError("接收端附件缺少 RSC/RSB 來源追溯資料。")
    rsc_fingerprint = str(rsc_record["calculationFingerprint"])
    if rsb_source.get("reshoreCalculationFingerprint") != rsc_fingerprint:
        raise ValueError("接收端附件的 RSB 沒有連結同一筆 RSC v4。")
    if (
        rsb_source.get("reshoreCalculationEvidenceFileName")
        != calculation_response["evidence"].get("fileName")
        or rsb_source.get("reshoreCalculationEvidenceFileSha256")
        != calculation_response["evidence"].get("fileSha256")
    ):
        raise ValueError("接收端附件的 RSB 父 RSC JSON 檔名或 SHA-256 不符。")
    for key in ("handoffFingerprint", "sourceCalculationFingerprint", "transferId"):
        if source.get(key) != rsb_source.get(key):
            raise ValueError(f"接收端附件的 RSC/RSB 來源欄位 {key} 不一致。")
    results = rsc_record.get("results")
    rsb_results = rsb_record.get("results")
    if not isinstance(results, dict) or not isinstance(rsb_results, dict):
        raise ValueError("接收端附件缺少 RSC/RSB 計算結果。")
    end_bearing = results.get("endBearing")
    if not isinstance(end_bearing, dict):
        raise ValueError("接收端附件缺少上下端承壓結果。")
    if (
        rsb_results.get("top") != end_bearing.get("top")
        or rsb_results.get("bottom") != end_bearing.get("bottom")
        or rsb_results.get("totalDemandPerMemberTf") != results.get("totalDemandPerMemberTf")
        or rsb_results.get("endBearingTransferCapacityTf") != results.get("endBearingTransferCapacityTf")
    ):
        raise ValueError("接收端附件的 RSB 與 RSC v4 承壓結果不一致。")
    rsb_input = rsb_record.get("input")
    if not isinstance(rsb_input, dict):
        raise ValueError("接收端附件的 RSB 缺少可重播輸入。")
    for key in (
        "memberCount",
        "imbalanceFactor",
        "additionalAxialLoadTfPerMember",
        "fyTfPerCm2",
        "governingLoadCombination",
        "loadDistributionBasis",
        "additionalLoadBasis",
    ):
        if key not in rsb_input:
            raise ValueError(f"接收端附件的 RSB 缺少自足輸入 {key}。")
    member_count = rsb_input["memberCount"]
    imbalance_factor = rsb_input["imbalanceFactor"]
    additional_axial = rsb_input["additionalAxialLoadTfPerMember"]
    source_demand = rsb_source.get("receiverTransferDemandTf")
    numeric_values = (member_count, imbalance_factor, additional_axial, source_demand)
    if any(
        isinstance(item, bool)
        or not isinstance(item, (int, float))
        or not math.isfinite(float(item))
        for item in numeric_values
    ) or float(member_count) <= 0:
        raise ValueError("接收端附件的 RSB 需求重播輸入不是有效有限數值。")
    expected_transfer_per_member = (
        float(source_demand) * float(imbalance_factor) / float(member_count)
    )
    expected_total_per_member = expected_transfer_per_member + float(additional_axial)
    transfer_result = rsb_results.get("transferDemandPerMemberTf")
    total_result = rsb_results.get("totalDemandPerMemberTf")
    if any(
        isinstance(item, bool)
        or not isinstance(item, (int, float))
        or not math.isfinite(float(item))
        for item in (transfer_result, total_result)
    ):
        raise ValueError("接收端附件的 RSB 需求結果不是有效有限數值。")
    if not math.isclose(
        float(transfer_result),
        expected_transfer_per_member,
        rel_tol=0.0,
        abs_tol=1e-6,
    ) or not math.isclose(
        float(total_result),
        expected_total_per_member,
        rel_tol=0.0,
        abs_tol=1e-6,
    ):
        raise ValueError("接收端附件的 RSB 無法以單檔輸入重播單支需求 P。")
    for end_name in ("top", "bottom"):
        _validate_end_result(end_bearing.get(end_name), end_name)
    code_basis = rsb_record.get("codeBasis")
    h_section_basis = (
        code_basis.get("hSectionFinishedEndBearing")
        if isinstance(code_basis, dict)
        else None
    )
    if not isinstance(h_section_basis, dict) or not str(
        h_section_basis.get("criterion", "")
    ).strip():
        raise ValueError("接收端附件的 RSB 缺少 H 型鋼精平端面承壓公式依據。")
    return {
        "rsc": rsc_record,
        "rsb": rsb_record,
        "rscFingerprint": rsc_fingerprint,
        "rsbFingerprint": str(rsb_record["calculationFingerprint"]),
        "rscEvidence": calculation_response["evidence"],
        "rsbEvidence": calculation_response["bearingEvidence"],
    }


def _validate_end_result(value: Any, end_name: str) -> None:
    if not isinstance(value, dict) or value.get("status") not in {"passed", "failed"}:
        raise ValueError(f"接收端附件的 {end_name} 端承壓結果不完整。")
    for key in (
        "supportBearingUtilizationRatio",
        "plateBendingUtilizationRatio",
        "hSectionEndBearingUtilizationRatio",
        "supportCapacityPerMemberTf",
        "plateBendingCapacityPerMemberTf",
        "hSectionEndBearingCapacityPerMemberTf",
        "governingCapacityPerMemberTf",
    ):
        child = value.get(key)
        if isinstance(child, bool) or not isinstance(child, (int, float)) or not math.isfinite(float(child)):
            raise ValueError(f"接收端附件的 {end_name} 端欄位 {key} 不是有限數值。")


def _fmt(value: Any, digits: int = 6) -> str:
    if value is None:
        return "-"
    if isinstance(value, bool):
        return "是" if value else "否"
    if isinstance(value, (int, float)):
        number = float(value)
        if not math.isfinite(number):
            return "-"
        return f"{number:.{digits}f}".rstrip("0").rstrip(".")
    return str(value)


def _status(value: Any) -> str:
    return "PASS" if value == "passed" else "NG"


def _support_material(value: Any) -> str:
    return "研磨或加工鋼面" if value == "finished_steel" else "混凝土全面積承壓"


def _identity_rows(context: dict[str, Any], document_metadata: dict[str, str]) -> list[list[str]]:
    rsc = context["rsc"]
    source = rsc["source"]
    rsb_source = context["rsb"]["source"]
    return [
        ["文件狀態", document_metadata["document_status_label"]],
        ["輸出時間", document_metadata["output_time"]],
        ["產出工具", REPORT_TOOL_NAME],
        ["工具版本", REPORT_TOOL_VERSION],
        ["附件種類", ATTACHMENT_KIND],
        ["RSC 版本／指紋", f"v{rsc['schemaVersion']} / {context['rscFingerprint']}"],
        ["RSC JSON／SHA-256", f"{context['rscEvidence']['fileName']} / {context['rscEvidence']['fileSha256']}"],
        ["RSB 版本／指紋", f"v{context['rsb']['schemaVersion']} / {context['rsbFingerprint']}"],
        ["RSB JSON／SHA-256", f"{context['rsbEvidence']['fileName']} / {context['rsbEvidence']['fileSha256']}"],
        [
            "RSB 宣告父 RSC JSON／SHA-256",
            f"{rsb_source.get('reshoreCalculationEvidenceFileName')} / "
            f"{rsb_source.get('reshoreCalculationEvidenceFileSha256')}",
        ],
        ["ERH 指紋", _fmt(source.get("handoffFingerprint"))],
        ["來源計算指紋", _fmt(source.get("sourceCalculationFingerprint"))],
        ["ERT 承接列", _fmt(source.get("transferId"))],
        ["承接對象", _fmt(source.get("receiverTarget"))],
    ]


def _input_rows(context: dict[str, Any]) -> list[list[str]]:
    value = {**context["rsc"]["input"], **context["rsb"]["input"]}
    return [
        ["分析模式", _fmt(value.get("analysisMode"))],
        ["H 型鋼斷面", _fmt(value.get("sectionName"))],
        ["共同承接支數", _fmt(value.get("memberCount"))],
        ["支數不均勻係數", _fmt(value.get("imbalanceFactor"))],
        ["附加單支軸力 Padd (tf)", _fmt(value.get("additionalAxialLoadTfPerMember"))],
        ["H 型鋼 Fy (tf/cm2)", _fmt(value.get("fyTfPerCm2"))],
        ["容許應力提高係數", _fmt(value.get("allowableStressIncreaseFactor"))],
        ["控制載重組合", _fmt(value.get("governingLoadCombination"))],
        ["移轉需求分配依據", _fmt(value.get("loadDistributionBasis"))],
        ["附加軸力依據", _fmt(value.get("additionalLoadBasis"))],
        ["端部承壓模式", _fmt(value.get("endBearingMode"))],
        ["頂端端板 B x L x t (cm)", f"{_fmt(value.get('topEndPlateWidthCm'))} x {_fmt(value.get('topEndPlateLengthCm'))} x {_fmt(value.get('topEndPlateThicknessCm'))}"],
        ["頂端端板 Fy (tf/cm2)", _fmt(value.get("topEndPlateFyTfPerCm2"))],
        ["頂端支承材料／強度", f"{_support_material(value.get('topSupportMaterial'))} / {_fmt(value.get('topSupportMaterialStrengthTfPerCm2'))} tf/cm2"],
        ["底端端板 B x L x t (cm)", f"{_fmt(value.get('bottomEndPlateWidthCm'))} x {_fmt(value.get('bottomEndPlateLengthCm'))} x {_fmt(value.get('bottomEndPlateThicknessCm'))}"],
        ["底端端板 Fy (tf/cm2)", _fmt(value.get("bottomEndPlateFyTfPerCm2"))],
        ["底端支承材料／強度", f"{_support_material(value.get('bottomSupportMaterial'))} / {_fmt(value.get('bottomSupportMaterialStrengthTfPerCm2'))} tf/cm2"],
        ["構造與密貼依據", _fmt(value.get("endBearingConfigurationBasis"))],
        ["端板局部彎曲方法依據", _fmt(value.get("endPlateBendingMethodBasis"))],
        ["頂端支承承壓依據", _fmt(value.get("topSupportBearingBasis"))],
        ["底端支承承壓依據", _fmt(value.get("bottomSupportBearingBasis"))],
        ["中心全接觸直接承壓確認", _fmt(value.get("centeredFullContactEndBearingConfirmed"))],
        ["H 型鋼端面精平確認", _fmt(value.get("hSectionEndFinishedConfirmed"))],
        ["單片無孔無加勁端板確認", _fmt(value.get("unperforatedUnstiffenedSinglePlateConfirmed"))],
    ]


def _demand_rows(context: dict[str, Any]) -> list[list[str]]:
    rsb = context["rsb"]
    source = rsb["source"]
    results = rsb["results"]
    return [
        ["來源端移轉需求 PERT (tf)", _fmt(source.get("receiverTransferDemandTf"))],
        ["不均勻放大後單支移轉需求 (tf)", _fmt(results.get("transferDemandPerMemberTf"))],
        ["附加單支軸力 Padd (tf)", _fmt(rsb["input"].get("additionalAxialLoadTfPerMember"))],
        ["單支端部總需求 P (tf)", _fmt(results.get("totalDemandPerMemberTf"))],
    ]


def _summary_rows(context: dict[str, Any]) -> list[list[str]]:
    results = context["rsc"]["results"]
    rsb_results = context["rsb"]["results"]
    return [
        ["RSB 上下端承壓結論", _status(rsb_results.get("status"))],
        ["RSC 構件與承壓綜合結論", _status(results.get("status"))],
        ["承壓可移轉容量 (tf)", _fmt(results.get("endBearingTransferCapacityTf"))],
        ["構件名義可移轉容量 (tf)", _fmt(results.get("nominalTransferCapacityTf"))],
        ["採用可移轉容量 (tf)", _fmt(results.get("adoptableTransferCapacityTf"))],
        ["控制容量模式", _fmt(results.get("governingCapacityMode"))],
        ["採用容量利用比", _fmt(results.get("capacityUtilizationRatio"))],
    ]


def _end_check_rows(end: dict[str, Any]) -> list[list[str]]:
    return [
        [
            "支承面承壓",
            _fmt(end.get("supportBearingUtilizationRatio")),
            _fmt(end.get("supportCapacityPerMemberTf")),
            _status(end.get("checks", {}).get("supportBearing")),
        ],
        [
            "端板局部彎曲",
            _fmt(end.get("plateBendingUtilizationRatio")),
            _fmt(end.get("plateBendingCapacityPerMemberTf")),
            _status(end.get("checks", {}).get("endPlateLocalBending")),
        ],
        [
            "H 型鋼精平端面承壓",
            _fmt(end.get("hSectionEndBearingUtilizationRatio")),
            _fmt(end.get("hSectionEndBearingCapacityPerMemberTf")),
            _status(end.get("checks", {}).get("hSectionFinishedEndBearing")),
        ],
        ["單端控制容量", "-", _fmt(end.get("governingCapacityPerMemberTf")), _status(end.get("status"))],
    ]


def _formula_lines(context: dict[str, Any]) -> list[str]:
    code_basis = context["rsb"].get("codeBasis", {})
    support = code_basis.get("supportBearing", {})
    plate = code_basis.get("endPlateLocalBending", {})
    h_section = code_basis.get("hSectionFinishedEndBearing", {})
    connection = code_basis.get("compressionBearingConnectionBoundary", {})
    support_criteria = support.get("criteria", [])
    return [
        "檢核公式：P = PERT x imbalance / memberCount + Padd.",
        *[str(item) for item in support_criteria],
        str(h_section.get("criterion", "")),
        str(plate.get("criterion", "")),
        f"端板局部彎曲依據：{plate.get('authority', '')}；{plate.get('method', '')}",
        f"端板局部彎曲力學參考：{plate.get('mechanicsReference', '')}",
        f"承壓接合未涵蓋邊界：{connection.get('criterion', '')}",
    ]


def _boundary_lines(context: dict[str, Any]) -> list[str]:
    rsc = context["rsc"]
    lines = [str(item) for item in rsc.get("verificationScope", {}).get("uncoveredChecks", [])]
    labels = {
        "receivingStructureStrengthNotDerived": "未由本附件推導承接結構本體強度。",
        "connectionWeldsBoltsAnchorsNotChecked": "焊道、螺栓、錨定與定位接合未檢核。",
        "stiffenersAndLocalReinforcementNotChecked": "加勁肋與 H 型鋼局部補強未檢核。",
        "constructionSequenceAndPreloadNotChecked": "施工順序、預壓與卸載程序未檢核。",
        "doesNotAutoApproveReceiverReceipt": "本附件不會自動核可 RVR。",
    }
    for key, label in labels.items():
        if context["rsb"].get("boundary", {}).get(key) is True:
            lines.append(label)
    lines.append("本獨立接收端附件不會寫入或改動來源專案的主 PDF/DOCX 計算書。")
    return list(dict.fromkeys(line for line in lines if line))


def _register_embedded_pdf_font() -> str:
    failures: list[str] = []
    for candidate in _PDF_FONT_CANDIDATES:
        if not candidate.is_file():
            continue
        try:
            registerFont(TTFont(_PDF_FONT_NAME, str(candidate)))
            return _PDF_FONT_NAME
        except Exception as exc:
            failures.append(f"{candidate}: {exc}")
    detail = f"（{'; '.join(failures)}）" if failures else ""
    raise RuntimeError(
        "找不到可嵌入的繁中文字型，接收端 PDF 已 fail closed；請安裝 Noto Sans TC 或 Arial Unicode。"
        + detail
    )


def _write_pdf(path: Path, context: dict[str, Any], document_metadata: dict[str, str]) -> None:
    pdf_font = _register_embedded_pdf_font()
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="RscTitle", fontName=pdf_font, fontSize=19, leading=26, textColor=colors.HexColor("#0f172a"), alignment=TA_LEFT))
    styles.add(ParagraphStyle(name="RscHeading", fontName=pdf_font, fontSize=13, leading=18, textColor=colors.HexColor("#0f4c81"), spaceBefore=8, spaceAfter=5))
    styles.add(ParagraphStyle(name="RscBody", fontName=pdf_font, fontSize=9.5, leading=14, textColor=colors.HexColor("#0f172a"), alignment=TA_LEFT))
    styles.add(ParagraphStyle(name="RscCell", fontName=pdf_font, fontSize=8.4, leading=10.5, textColor=colors.HexColor("#0f172a"), alignment=TA_LEFT))
    elements: list[Any] = [
        Paragraph("RSC v4 / RSB v1 接收端獨立計算附件", styles["RscTitle"]),
        Paragraph(f"文件狀態：{escape(document_metadata['document_status_label'])}", styles["RscBody"]),
        Spacer(1, 4 * mm),
        Paragraph("一、文件身分與追溯", styles["RscHeading"]),
        _pdf_table(_identity_rows(context, document_metadata), [43 * mm, 127 * mm], styles),
        Paragraph("二、採用輸入", styles["RscHeading"]),
        _pdf_table(
            [["項目", "採用值"], *_input_rows(context)],
            [51 * mm, 119 * mm],
            styles,
            header=True,
        ),
        Paragraph("三、需求 P 與控制容量", styles["RscHeading"]),
        _pdf_table(_demand_rows(context), [72 * mm, 98 * mm], styles),
        Spacer(1, 2 * mm),
        _pdf_table(_summary_rows(context), [72 * mm, 98 * mm], styles, status_rows=True),
        Paragraph("四、上下端三項檢核", styles["RscHeading"]),
    ]
    end_bearing = context["rsc"]["results"]["endBearing"]
    for title, key in (("頂端", "top"), ("底端", "bottom")):
        end = end_bearing[key]
        elements.append(Paragraph(f"{title}：{_status(end.get('status'))}", styles["RscBody"]))
        elements.append(_pdf_table(
            [["檢核項目", "DCR", "容量 (tf/支)", "結果"], *_end_check_rows(end)],
            [60 * mm, 30 * mm, 45 * mm, 35 * mm],
            styles,
            header=True,
            status_column=3,
        ))
        elements.append(Spacer(1, 3 * mm))
    elements.extend([
        Paragraph("五、公式與依據", styles["RscHeading"]),
        *[Paragraph(f"{index}. {escape(line)}", styles["RscBody"]) for index, line in enumerate(_formula_lines(context), start=1) if line],
        Paragraph("六、未涵蓋邊界", styles["RscHeading"]),
        *[Paragraph(f"{index}. {escape(line)}", styles["RscBody"]) for index, line in enumerate(_boundary_lines(context), start=1)],
    ])
    document = SimpleDocTemplate(
        str(path),
        pagesize=A4,
        rightMargin=20 * mm,
        leftMargin=20 * mm,
        topMargin=18 * mm,
        bottomMargin=18 * mm,
        title="RSC v4 / RSB v1 接收端獨立計算附件",
        author="開挖擋土支撐工具",
    )
    document.build(elements, onFirstPage=_draw_pdf_frame, onLaterPages=_draw_pdf_frame)


def _pdf_table(
    rows: list[list[str]],
    widths: list[float],
    styles,
    *,
    header: bool = False,
    status_column: int | None = None,
    status_rows: bool = False,
) -> Table:
    data = [[Paragraph(escape(str(cell)), styles["RscCell"]) for cell in row] for row in rows]
    table = Table(data, colWidths=widths, repeatRows=1 if header else 0)
    style = TableStyle([
        ("FONTNAME", (0, 0), (-1, -1), _PDF_FONT_NAME),
        ("FONTSIZE", (0, 0), (-1, -1), 8.4),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#cbd5e1")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("ROWBACKGROUNDS", (0, 1 if header else 0), (-1, -1), [colors.white, colors.HexColor("#f8fafc")]),
    ])
    if header:
        style.add("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#dbeafe"))
        style.add("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#1e3a8a"))
    else:
        style.add("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#eff6ff"))
    for row_index, row in enumerate(rows):
        candidates: list[int] = []
        if status_column is not None and status_column < len(row):
            candidates.append(status_column)
        if status_rows and len(row) > 1:
            candidates.append(1)
        for column in candidates:
            if row[column] == "NG":
                style.add("BACKGROUND", (column, row_index), (column, row_index), colors.HexColor("#fee2e2"))
                style.add("TEXTCOLOR", (column, row_index), (column, row_index), colors.HexColor("#991b1b"))
            elif row[column] == "PASS":
                style.add("BACKGROUND", (column, row_index), (column, row_index), colors.HexColor("#dcfce7"))
                style.add("TEXTCOLOR", (column, row_index), (column, row_index), colors.HexColor("#166534"))
    table.setStyle(style)
    return table


def _draw_pdf_frame(canvas, document) -> None:
    canvas.saveState()
    canvas.setFont(_PDF_FONT_NAME, 8)
    canvas.setFillColor(colors.HexColor("#64748b"))
    canvas.drawString(20 * mm, 10 * mm, "RSC v4 / RSB v1 接收端獨立計算附件")
    canvas.drawRightString(A4[0] - 20 * mm, 10 * mm, f"第 {document.page} 頁")
    canvas.restoreState()


def _write_docx(path: Path, context: dict[str, Any], document_metadata: dict[str, str]) -> None:
    document = Document()
    _configure_word_document(document)
    bullet_numbering_id = _add_compact_bullet_numbering(document)
    kicker = document.add_paragraph()
    kicker.paragraph_format.space_before = Pt(16)
    kicker.paragraph_format.space_after = Pt(3)
    _word_run(kicker, "RSC v4 / RSB v1", bold=True, size=11, color="1F4D78")
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(4)
    _word_run(title, "接收端獨立計算附件", bold=True, size=23, color="000000")
    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(10)
    _word_run(
        subtitle,
        "H 型鋼上下端直接承壓＋無加勁端板局部彎曲",
        size=13,
        color="373737",
    )

    _word_heading(document, "一、文件身分與追溯")
    _word_table(document, _identity_rows(context, document_metadata), [2700, 6660])
    _word_heading(document, "二、採用輸入")
    _word_table(
        document,
        [["項目", "採用值"], *_input_rows(context)],
        [2700, 6660],
        header=True,
    )
    _word_heading(document, "三、需求 P 與控制容量")
    _word_table(document, _demand_rows(context), [2700, 6660])
    document.add_paragraph().paragraph_format.space_after = Pt(1)
    _word_table(document, _summary_rows(context), [2700, 6660], status_rows=True)
    _word_heading(document, "四、上下端三項檢核")
    end_bearing = context["rsc"]["results"]["endBearing"]
    for title_text, key in (("頂端", "top"), ("底端", "bottom")):
        end = end_bearing[key]
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(4)
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.keep_with_next = True
        _word_run(paragraph, f"{title_text}：{_status(end.get('status'))}", bold=True, size=11)
        _word_table(
            document,
            [["檢核項目", "DCR", "容量 (tf/支)", "結果"], *_end_check_rows(end)],
            [3300, 1500, 2460, 2100],
            header=True,
            status_column=3,
        )
    _word_heading(document, "五、公式與依據")
    for line in _formula_lines(context):
        if line:
            _word_list_paragraph(document, line, bullet_numbering_id)
    _word_heading(document, "六、未涵蓋邊界")
    for line in _boundary_lines(context):
        _word_list_paragraph(document, line, bullet_numbering_id)
    document.save(path)


def _configure_word_document(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    _set_style_fonts(normal, latin="Calibri", east_asia="Microsoft JhengHei")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for style_name, size, color, before, after in (
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ):
        heading = document.styles[style_name]
        heading.font.name = "Calibri"
        heading.font.size = Pt(size)
        heading.font.bold = True
        heading.font.color.rgb = RGBColor.from_string(color)
        _set_style_fonts(heading, latin="Calibri", east_asia="Microsoft JhengHei")
        heading.paragraph_format.space_before = Pt(before)
        heading.paragraph_format.space_after = Pt(after)
        heading.paragraph_format.keep_with_next = True
    for section in document.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)
        _configure_word_header_footer(section)


def _configure_word_header_footer(section) -> None:
    header = section.header
    footer = section.footer
    header_paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    header_paragraph.text = ""
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _word_run(header_paragraph, "RSC v4 / RSB v1 接收端獨立計算附件", size=8.5, color="64748B")
    footer_paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_paragraph.text = ""
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _word_run(footer_paragraph, "第 ", size=8, color="64748B")
    _append_page_number(footer_paragraph)
    _word_run(footer_paragraph, " 頁", size=8, color="64748B")


def _word_heading(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="Heading 1")
    _word_run(paragraph, text, bold=True, size=16, color="2E74B5")


def _word_list_paragraph(document: Document, text: str, numbering_id: int) -> None:
    paragraph = document.add_paragraph(style="Normal")
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph.paragraph_format.left_indent = Inches(0.375)
    paragraph.paragraph_format.first_line_indent = Inches(-0.188)
    num_pr = paragraph._p.get_or_add_pPr().get_or_add_numPr()
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), str(numbering_id))
    num_pr.extend([ilvl, num_id])
    _word_run(paragraph, text, size=11)


def _word_table(
    document: Document,
    rows: list[list[str]],
    widths_dxa: list[int],
    *,
    header: bool = False,
    status_column: int | None = None,
    status_rows: bool = False,
) -> None:
    if sum(widths_dxa) != 9360:
        raise ValueError("DOCX 表格欄寬必須精確合計 9360 DXA。")
    table = document.add_table(rows=0, cols=len(widths_dxa))
    table.style = "Table Grid"
    table.autofit = False
    _set_word_table_geometry(table, widths_dxa)
    for row_index, row in enumerate(rows):
        cells = table.add_row().cells
        _prevent_word_row_split(table.rows[-1])
        for column, value in enumerate(row):
            cell = cells[column]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _set_word_cell_margins(cell, top=80, bottom=80, start=120, end=120)
            _set_word_cell_text(cell, str(value), bold=header and row_index == 0)
            if column == 0 and not header:
                _shade_word_cell(cell, "F2F4F7")
        if header and row_index == 0:
            _repeat_word_header(table.rows[0])
            for cell in cells:
                _shade_word_cell(cell, "E8EEF5")
        candidates: list[int] = []
        if status_column is not None and status_column < len(row):
            candidates.append(status_column)
        if status_rows and len(row) > 1:
            candidates.append(1)
        for column in candidates:
            if row[column] == "NG":
                _shade_word_cell(cells[column], "FEE2E2")
                _set_word_cell_text(cells[column], "NG", bold=True, color="991B1B")
            elif row[column] == "PASS":
                _shade_word_cell(cells[column], "DCFCE7")
                _set_word_cell_text(cells[column], "PASS", bold=True, color="166534")
    _set_word_table_geometry(table, widths_dxa)


def _set_word_table_geometry(table, widths_dxa: list[int]) -> None:
    table_width = sum(widths_dxa)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(table_width))
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "120")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            cell.width = Inches(width / 1440.0)
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))


def _set_word_cell_text(cell, text: str, *, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.08
    _word_run(paragraph, text, bold=bold, size=10, color=color)


def _word_run(paragraph, text: str, *, bold: bool = False, size: float = 10, color: str | None = None) -> None:
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    _set_run_fonts(run, latin="Calibri", east_asia="Microsoft JhengHei")
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def _shade_word_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_word_cell_margins(
    cell,
    *,
    top: int,
    bottom: int,
    start: int,
    end: int,
) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.find(qn("w:tcMar"))
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side, value in (("top", top), ("left", start), ("bottom", bottom), ("right", end)):
        element = margins.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def _repeat_word_header(row) -> None:
    row_properties = row._tr.get_or_add_trPr()
    header = row_properties.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        row_properties.append(header)
    header.set(qn("w:val"), "true")


def _prevent_word_row_split(row) -> None:
    row_properties = row._tr.get_or_add_trPr()
    cant_split = row_properties.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        row_properties.append(cant_split)
    cant_split.set(qn("w:val"), "true")


def _add_compact_bullet_numbering(document: Document) -> int:
    numbering = document.part.numbering_part.element
    abstract_ids = [
        int(element.get(qn("w:abstractNumId")))
        for element in numbering.findall(qn("w:abstractNum"))
    ]
    number_ids = [
        int(element.get(qn("w:numId")))
        for element in numbering.findall(qn("w:num"))
    ]
    abstract_id = max(abstract_ids, default=-1) + 1
    number_id = max(number_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi_level = OxmlElement("w:multiLevelType")
    multi_level.set(qn("w:val"), "singleLevel")
    abstract.append(multi_level)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_format = OxmlElement("w:numFmt")
    num_format.set(qn("w:val"), "bullet")
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "•")
    level_justification = OxmlElement("w:lvlJc")
    level_justification.set(qn("w:val"), "left")
    paragraph_properties = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    indentation = OxmlElement("w:ind")
    indentation.set(qn("w:left"), "540")
    indentation.set(qn("w:hanging"), "270")
    paragraph_properties.extend([tabs, indentation])
    run_properties = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Calibri")
    fonts.set(qn("w:hAnsi"), "Calibri")
    run_properties.append(fonts)
    level.extend(
        [
            start,
            num_format,
            level_text,
            level_justification,
            paragraph_properties,
            run_properties,
        ]
    )
    abstract.append(level)
    numbering.append(abstract)

    number = OxmlElement("w:num")
    number.set(qn("w:numId"), str(number_id))
    abstract_reference = OxmlElement("w:abstractNumId")
    abstract_reference.set(qn("w:val"), str(abstract_id))
    number.append(abstract_reference)
    numbering.append(number)
    return number_id


def _set_style_fonts(style, *, latin: str, east_asia: str) -> None:
    r_pr = style.element.get_or_add_rPr()
    fonts = r_pr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        r_pr.append(fonts)
    fonts.set(qn("w:ascii"), latin)
    fonts.set(qn("w:hAnsi"), latin)
    fonts.set(qn("w:eastAsia"), east_asia)


def _set_run_fonts(run, *, latin: str, east_asia: str) -> None:
    r_pr = run._element.get_or_add_rPr()
    fonts = r_pr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        r_pr.append(fonts)
    fonts.set(qn("w:ascii"), latin)
    fonts.set(qn("w:hAnsi"), latin)
    fonts.set(qn("w:eastAsia"), east_asia)


def _append_page_number(paragraph) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()
    run._r.extend([begin, instruction, separate, text, end])
