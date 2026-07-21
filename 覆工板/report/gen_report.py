# -*- coding: utf-8 -*-
"""
覆工板系統結構計算書 — Word 產報腳本
讀取 HTML 工具匯出的 JSON，輸出正式 .docx 計算書
用法：
    python gen_report.py <input.json> [output.docx]
    若未指定 output，會存在 JSON 同目錄，檔名：覆工板計算書_{案名}_{日期}.docx
"""
import sys, os, json, io, hashlib
# Windows 主控台 UTF-8 輸出
if sys.platform.startswith('win'):
    try:
        sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
        sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8')
    except Exception:
        pass
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ========== 樣式與工具 ==========
FONT_ZH = '微軟正黑體'
FONT_EN = 'Times New Roman'

def set_cell_bg(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)

def set_run_font(run, size=10, bold=False, color=None):
    run.font.name = FONT_EN
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_ZH)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    sizes = {1: 16, 2: 13, 3: 11}
    set_run_font(run, size=sizes.get(level, 11), bold=True, color='1A4480')
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_para(doc, text, size=10, bold=False, align=None, indent=0):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    if align is not None:
        p.alignment = align
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    return p

def add_table(doc, headers, rows, col_widths=None, hilite_last=False):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = 'Light Grid Accent 1'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ''
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        set_run_font(r, size=10, bold=True, color='FFFFFF')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(hdr[i], '1A4480')
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for ri, row in enumerate(rows):
        tr = tbl.rows[ri + 1].cells
        for ci, v in enumerate(row):
            tr[ci].text = ''
            p = tr[ci].paragraphs[0]
            r = p.add_run(str(v))
            set_run_font(r, size=10)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            tr[ci].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if hilite_last and ri == len(rows) - 1:
                set_cell_bg(tr[ci], 'FFF3CD')
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in tbl.rows:
                row.cells[i].width = Cm(w)
    return tbl

def fmt(v, d=3):
    if v is None: return '—'
    try:
        if abs(v) >= 1e4: return f'{v:,.{d}f}'
        return f'{v:.{d}f}'
    except Exception:
        return str(v)

def ok_ng(b):
    if b is None: return '—'
    return 'OK' if b else 'NG'


# ========== 主流程 ==========
def generate(json_path, out_path=None):
    with open(json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)
    p = data['project']
    g = data['global']
    R = data.get('results', {})
    inp = data.get('inputs', {})
    document_state = data.get('document', {})
    document_mode = document_state.get('mode') or 'attachment'
    approved = document_state.get('approved') is True
    approved_at = str(document_state.get('approvedAt') or '').strip()
    calculation_fingerprint = str(document_state.get('calculationFingerprint') or '').strip()
    if not calculation_fingerprint:
        calculation_source = json.dumps({'global': g, 'results': R}, ensure_ascii=False, sort_keys=True, separators=(',', ':'))
        calculation_fingerprint = f"CF-{hashlib.sha256(calculation_source.encode('utf-8')).hexdigest()[:16].upper()}"

    doc = Document()
    # 版面：A4、邊界 2cm
    for section in doc.sections:
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
    # 預設字型
    style = doc.styles['Normal']
    style.font.name = FONT_EN
    style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_ZH)
    style.font.size = Pt(10)

    # ----- 封面 -----
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.paragraph_format.space_before = Pt(80)
    r = t.add_run('覆工板系統結構計算書')
    set_run_font(r, size=26, bold=True, color='1A4480')

    project_rows = [
        ('案　　　名', str(p.get('name') or '').strip()),
        ('案件編號', str(p.get('no') or '').strip()),
        ('日　　　期', str(p.get('date') or '').strip()),
    ]
    project_rows = [(label, value) for label, value in project_rows if value]
    if project_rows:
        add_para(doc, '', size=10)
        for label, value in project_rows:
            add_para(doc, f"{label}：{value}", size=13, align=WD_ALIGN_PARAGRAPH.CENTER)

    for _ in range(6):
        add_para(doc, '', size=11)

    add_para(doc, '弘一工程顧問有限公司', size=15, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_page_break()

    # ----- 一、計算結果總表 -----
    add_heading(doc, '一、計算結果總表', 1)
    add_para(doc, '本案依 HS 20-44 卡車、PC400 履帶車、指定噸位吊車三組載重，檢核覆工板面、覆工小梁、覆工大梁及共構柱。各項檢核摘要如下：')

    rows = []
    if R.get('deck'):
        d = R['deck']
        all_ok = d['ok_fb'] and d['ok_fv'] and d['ok_d']
        rows.append(['覆工版面', f"{d['n']}片 (Sx={d['Sx']/d['n']:.0f}cm³/片)",
                     fmt(d['Mmax'],2), fmt(d['Vmax'],2), ok_ng(all_ok)])
    if R.get('stringer'):
        s = R['stringer']
        all_ok = s['ok_fb'] and s['ok_fv'] and s['ok_d']
        rows.append(['覆工小梁', f"{s['sec']['name']}×{s['n']}支",
                     fmt(s['Mmax'],2), fmt(s['Vmax'],2), ok_ng(all_ok)])
    if R.get('girder'):
        gi = R['girder']
        all_ok = gi['ok_fb'] and gi['ok_fv'] and gi['ok_d']
        rows.append(['覆工大梁', f"{gi['sec']['name']}×{gi['n']}支",
                     fmt(gi['Mmax'],2), fmt(gi['Vmax'],2), ok_ng(all_ok)])
    if R.get('column'):
        c = R['column']
        rows.append(['共構柱', f"{c['sec']['name']}  L={c['L']}cm",
                     f"Mx={fmt(c['Mx'],2)}", f"N={fmt(c['N'],2)}", ok_ng(c['ok'])])
    if R.get('bond'):
        b = R['bond']
        rows.append(['H型鋼握裹', f"L={b['L']}cm, fc'={b['fc']}", f"F={fmt(b['F'],2)}tf", f"N={fmt(b['Nc'],2)}tf", ok_ng(b['ok'])])
    if R.get('pile'):
        pl = R['pile']
        rows.append(['樁基承載', f"D={pl['D']}cm, Lb={pl['Lb']}cm", f"Qa={fmt(pl['Qa'],2)}tf", f"P={pl['P']}tf", ok_ng(pl['ok'])])

    add_table(doc, ['構件', '規格', 'Mmax / 其他', 'Vmax / 其他', '結果'], rows)

    # 整體結論
    overall = True
    for k, v in R.items():
        if v is None: continue
        if k in ('deck','stringer','girder'):
            if not (v['ok_fb'] and v['ok_fv'] and v['ok_d']):
                overall = False
        elif 'ok' in v:
            if not v['ok']:
                overall = False
    add_para(doc, '')
    conclu = f"整體結論：{'整體系統檢核「通過」，符合規範規定。' if overall else '部分項目不通過，請依上表 NG 項目調整構件規格或補強配置。'}"
    add_para(doc, conclu, size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_page_break()

    # ----- 二、設計準則 -----
    add_heading(doc, '二、設計準則', 1)

    add_heading(doc, '2.1 適用規範', 2)
    for txt in [
        '1. 鋼構造建築物鋼結構設計技術規範（內政部）',
        '2. AASHTO Standard Specifications for Highway Bridges (HS 20-44 載重)',
        '3. 美國鋼結構學會 AISC ASD（容許應力設計法）',
        '4. CNS 國家標準 H 型鋼規格（JIS G 3192 / CNS 550）',
    ]:
        add_para(doc, txt, indent=0.5)

    add_heading(doc, '2.2 材料性質', 2)
    add_table(doc,
        ['項目', '數值', '備註'],
        [
            ['鋼料降伏強度 Fy', f"{g['Fy']} kgf/cm²", 'A36 鋼料'],
            ['鋼料彈性模數 E', f"{g['E']:,.0f} kgf/cm²", ''],
            ['短期應力放大係數 β', f"{g['beta']}", '一般 1.33、吊車 1.25'],
            ['衝擊係數 i', f"{g['imp']}", '僅行駛車輛適用'],
            ['覆工板自重 Wp', f"{g['Wp']} kgf/m²", '標準 H190 集成板'],
            ['活載重 Wl', f"{g['Wl']} kgf/m²", '一般工況'],
            ['容許撓度', f"L / {g['defl']}", ''],
        ])

    add_heading(doc, '2.3 設計載重規格', 2)
    add_table(doc,
        ['類別', '型式', '載重', '幾何', '衝擊'],
        [
            ['HS 20-44 卡車', '集中', 'P = 7.3 t/輪', '輪心距 1.2/1.8/1.2 m', f"×{1+g['imp']}"],
            ['PC400 履帶車', '均佈(履帶)', 'Wc = 8.244 t/m', '履帶長 4.35 m、輪心距 2.74 m', f"×{1+g['imp']}"],
            ['指定吊車', '集中', f"P = 由工況載入", '依機型選配', '不計（固定吊裝）'],
        ])

    doc.add_page_break()

    # ----- 三、覆工版面檢核 -----
    if R.get('deck'):
        d = R['deck']
        add_heading(doc, '三、覆工版面檢核', 1)

        add_heading(doc, '3.1 斷面性質與幾何', 2)
        add_table(doc, ['項目','值','單位'], [
            ['板寬 B', d['B'], 'm'],
            ['板跨 L', d['L'], 'm'],
            ['並排板數 n', d['n'], '片'],
            ['Sx (總)', f"{d['Sx']}", 'cm³'],
            ['Ix (總)', f"{d['Ix']}", 'cm⁴'],
            ['Aw (總)', f"{d['Aw']}", 'cm²'],
        ])

        add_heading(doc, '3.2 載重計算', 2)
        add_para(doc, f"線均佈 WT = (Wp+Wl)·B/1000 = ({g['Wp']}+{g['Wl']})×{d['B']}/1000 = {fmt(d['WT'],4)} Tf/m")
        add_para(doc, f"HS20 含衝擊 P_HS = 7.3×(1+{g['imp']}) = {fmt(d['P_HS'],3)} tf")
        add_para(doc, f"PC400 含衝擊 Wc = 8.244×(1+{g['imp']}) = {fmt(d['Wc'],4)} tf/m")
        add_para(doc, f"指定吊車 P = {d['Pc']} tf（不衝擊）")

        add_heading(doc, '3.3 載重組合彎矩與剪力', 2)
        add_table(doc, ['工況','公式','Mmax (Tf-m)','Vmax (Tf)'], [
            ['① HS 20-44', '⅛·WT·L² + P_HS·L/4', fmt(d['M1'],3), fmt(d['V1'],3)],
            ['② PC400 履帶', '⅛·(WT+Wc)·L² (板跨<履帶長)', fmt(d['M2'],3), fmt(d['V2'],3)],
            ['③ 吊車', '⅛·WT·L² + P·L/4', fmt(d['M3'],3), fmt(d['V3'],3)],
            [f"控制（{d['ctrl']}）", 'max', fmt(d['Mmax'],3), fmt(d['Vmax'],3)],
        ], hilite_last=True)

        add_heading(doc, '3.4 應力與撓度檢核', 2)
        add_table(doc, ['檢核項目', '計算', '結果'], [
            ['彎矩 fb = Mmax×10⁵/Sx', f"{fmt(d['Mmax'],3)}×10⁵/{d['Sx']}", f"{fmt(d['fb'],1)} kgf/cm² < Fb={fmt(d['Fb'],1)} … {ok_ng(d['ok_fb'])}"],
            ['剪力 fv = Vmax×10³/Aw', f"{fmt(d['Vmax'],3)}×10³/{d['Aw']}", f"{fmt(d['fv'],1)} kgf/cm² < Fv={fmt(d['Fv'],1)} … {ok_ng(d['ok_fv'])}"],
            ['撓度 δmax（三工況取大）', f"控制 = {d['def3']['ctrl']}", f"{fmt(d['def3']['dmax'],3)} cm < L/{g['defl']}={fmt(d['def_allow'],3)} cm … {ok_ng(d['ok_d'])}"],
        ])
        add_para(doc, '（撓度嚴謹解：HS20/吊車以 5wL⁴/384EI + PL³/48EI；PC400 履帶以部分 UDL 閉式解 wc(8L³−4c²L+c³)/384EI）', size=9)

        doc.add_page_break()

    # ----- 四、覆工小梁檢核 -----
    if R.get('stringer'):
        s = R['stringer']
        add_heading(doc, '四、覆工小梁檢核', 1)

        add_heading(doc, '4.1 採用斷面與幾何', 2)
        add_table(doc, ['項目','值'], [
            ['型鋼規格', f"{s['sec']['name']} × {s['n']} 支"],
            ['小梁間距 B (= 板跨)', f"{s['B']} m"],
            ['小梁跨距 L', f"{s['L']} m"],
            ['斷面 Sx (總)', f"{s['Sx']} cm³"],
            ['斷面 Ix (總)', f"{s['Ix']} cm⁴"],
            ['斷面 Aw (總)', f"{s['Aw']:.1f} cm²"],
            ['翼板', f"B/(2tf)={s['sec']['B']/(2*s['sec']['tf']):.2f} → {s['flange']['label']}"],
            ['腹板', f"H/tw={s['sec']['H']/s['sec']['tw']:.2f} → {s['web']['label']}"],
            ['側撐', f"Lc={s['brace']['Lc']:.1f} cm / Lt={s['Lt']} cm → {s['brace']['label']}"],
        ])

        add_heading(doc, '4.2 載重與組合', 2)
        add_para(doc, f"WT = (Wp·B + Wl·B + Wb)/1000 = {fmt(s['WT'],4)} Tf/m　(小梁自重 Wb={s['Wb']} kgf/m)")
        add_table(doc, ['工況','Mmax (Tf-m)','Vmax (Tf)'], [
            ['① HS 20-44', fmt(s['M1'],3), fmt(s['V1'],3)],
            ['② PC400 履帶', fmt(s['M2'],3), fmt(s['V2'],3)],
            ['③ 吊車', fmt(s['M3'],3), fmt(s['V3'],3)],
            [f"控制（{s['ctrl']}）", fmt(s['Mmax'],3), fmt(s['Vmax'],3)],
        ], hilite_last=True)

        add_heading(doc, '4.3 應力與撓度檢核', 2)
        add_table(doc, ['檢核項目','計算','結果'], [
            ['彎矩 fb', f"{fmt(s['Mmax'],3)}×10⁵/{s['Sx']}", f"{fmt(s['fb'],1)} < {fmt(s['Fb'],1)} … {ok_ng(s['ok_fb'])}"],
            ['剪力 fv', f"{fmt(s['Vmax'],3)}×10³/{s['Aw']:.1f}", f"{fmt(s['fv'],1)} < {fmt(s['Fv'],1)} … {ok_ng(s['ok_fv'])}"],
            ['撓度 δmax', f"控制 = {s['def3']['ctrl']}", f"{fmt(s['def3']['dmax'],3)} cm < {fmt(s['def_allow'],3)} cm … {ok_ng(s['ok_d'])}"],
        ])
        doc.add_page_break()

    # ----- 五、覆工大梁檢核 -----
    if R.get('girder'):
        gi = R['girder']
        add_heading(doc, '五、覆工大梁檢核', 1)

        add_heading(doc, '5.1 採用斷面與幾何', 2)
        add_table(doc, ['項目','值'], [
            ['型鋼規格', f"{gi['sec']['name']} × {gi['n']} 支"],
            ['大梁間距 B', f"{gi['B']} m"],
            ['大梁跨距 L', f"{gi['L']} m"],
            ['小梁傳來線載 W2', f"{gi['W2']} kgf/m"],
            ['Sx (總)', f"{gi['Sx']} cm³"],
            ['Ix (總)', f"{gi['Ix']} cm⁴"],
        ])

        add_heading(doc, '5.2 載重組合（雙 WT 制）', 2)
        add_para(doc, f"WT  (一般) = (Wp·B+Wl·B+Wb+W2)/1000 = {fmt(gi['WT'],4)} Tf/m")
        add_para(doc, f"WT₂ (吊車) = (Wp·B+Wb+W2)/1000 = {fmt(gi['WT2'],4)} Tf/m  （吊車工況不疊堆料、活載）")
        add_table(doc, ['工況','Mmax (Tf-m)','Vmax (Tf)'], [
            ['① HS (WT)',  fmt(gi['M1'],3), fmt(gi['V1'],3)],
            ['② PC400 (WT)', fmt(gi['M2'],3), fmt(gi['V2'],3)],
            ['③ 吊車 (WT₂)', fmt(gi['M3'],3), fmt(gi['V3'],3)],
            [f"控制（{gi['ctrl']}）", fmt(gi['Mmax'],3), fmt(gi['Vmax'],3)],
        ], hilite_last=True)

        add_heading(doc, '5.3 應力與撓度檢核', 2)
        add_table(doc, ['檢核項目','計算','結果'], [
            ['彎矩 fb', f"{fmt(gi['Mmax'],3)}×10⁵/{gi['Sx']}", f"{fmt(gi['fb'],1)} < {fmt(gi['Fb'],1)} … {ok_ng(gi['ok_fb'])}"],
            ['剪力 fv', f"{fmt(gi['Vmax'],3)}×10³/{gi['Aw']:.1f}", f"{fmt(gi['fv'],1)} < {fmt(gi['Fv'],1)} … {ok_ng(gi['ok_fv'])}"],
            ['撓度 δmax', f"控制 = {gi['def3']['ctrl']}", f"{fmt(gi['def3']['dmax'],3)} cm < {fmt(gi['def_allow'],3)} cm … {ok_ng(gi['ok_d'])}"],
        ])

        add_heading(doc, '5.4 傳至共構柱之軸力 Pu（三情境）', 2)
        add_table(doc, ['情境','公式','Pu (tf)'], [
            ['① PC400 固定支座位於柱上', 'Pu = Vmax', fmt(gi['Pu1'],3)],
            ['② 吊車支座距柱 0.5 m', '½·WT₂·L + P·(L−0.5)/L', fmt(gi['Pu2'],3)],
            ['③ 吊車支座於相鄰跨間', '½·WT₂·L + 2P·(L−4)/L', fmt(gi['Pu3'],3)],
            ['控制 PuMax', 'max', fmt(gi['PuMax'],3)],
        ], hilite_last=True)

        doc.add_page_break()

    # ----- 六、共構柱檢核 -----
    if R.get('column'):
        c = R['column']
        add_heading(doc, '六、共構柱檢核（AISC 軸壓+雙軸彎交互式）', 1)

        add_table(doc, ['項目','值'], [
            ['柱斷面', f"{c['sec']['name']} (A={c['A']} cm²)"],
            ['柱長 L₁ / 有效長度係數 K', f"{c['L']} cm / {c['K']}"],
            ['軸力 N (壓)', f"{fmt(c['N'],2)} tf"],
            ['偏心 ex, ey', f"{c['ex']}, {c['ey']} cm"],
            ['舊料折減', f"{c['old']}"],
            ['fa = N/A', f"{fmt(c['fa'],2)} kgf/cm²"],
            ['Mx, My', f"{fmt(c['Mx'],3)}, {fmt(c['My'],3)} tf-m"],
            ['fbx, fby', f"{fmt(c['fbx'],2)}, {fmt(c['fby'],2)} kgf/cm²"],
            ['KL/rx, KL/ry', f"{fmt(c['KLrx'],2)}, {fmt(c['KLry'],2)}"],
            ['Cc', f"{fmt(c['Cc'],2)}"],
            ['Fa, Fa₁', f"{fmt(c['Fa'],1)} → {fmt(c['Fa1'],1)} kgf/cm²"],
            ['Fbx, Fby', f"{fmt(c['Fbx'],1)}, {fmt(c['Fby'],1)} kgf/cm²"],
        ])

        add_para(doc, '')
        add_para(doc, '交互式檢核：', bold=True)
        add_para(doc, f"式①: fa/Fa₁ + Cm·fbx/(1−fa/Fex)/Fbx + Cm·fby/(1−fa/Fey)/Fby = {fmt(c['chk1'],3)} ≤ 1.0 … {ok_ng(c['chk1']<=1.0)}", indent=0.5)
        add_para(doc, f"式②: fa/(0.6Fy) + fbx/Fbx + fby/Fby = {fmt(c['chk2'],3)} ≤ 1.0 … {ok_ng(c['chk2']<=1.0)}", indent=0.5)

        doc.add_page_break()

    # ----- 七、H 型鋼握裹 -----
    if R.get('bond'):
        b = R['bond']
        add_heading(doc, '七、H 型鋼貫入 PC 之握裹力檢核', 1)
        add_table(doc, ['項目','計算','結果'], [
            ["混凝土強度 fc'", '—', f"{b['fc']} kgf/cm²"],
            ['貫入深度 L', '—', f"{b['L']} cm"],
            ['型鋼周長 ls', '2(B+H)', f"{fmt(b['ls'],2)} cm"],
            ["τ = 0.03·fc'", f"0.03×{b['fc']}", f"{fmt(b['tau'],3)} kgf/cm²"],
            ['拉拔 F = τ·L·ls', '—', f"{fmt(b['F'],2)} tf > T={b['T']} tf … {ok_ng(b['ok_T'])}"],
            ["壓入 N = τ·L·ls + 0.35·fc'·As'", '—', f"{fmt(b['Nc'],2)} tf > P={b['P']} tf … {ok_ng(b['ok_P'])}"],
        ])
        add_para(doc, '')

    # ----- 八、樁基承載 -----
    if R.get('pile'):
        pl = R['pile']
        add_heading(doc, '八、共構樁基承載力檢核', 1)
        add_table(doc, ['項目','計算','結果'], [
            ['樁徑 D / 樁長 Lb', '—', f"{pl['D']} cm / {pl['Lb']} cm"],
            ['樁底面積 Ab', 'πD²/4', f"{fmt(pl['Ab'],1)} cm²"],
            ['端點承載 qb', '7.5·N (砂)', f"{fmt(pl['qb'],2)} tf/m²"],
            ['側摩擦 fs', 'N/3 (砂)', f"{fmt(pl['fs'],2)} tf/m²"],
            ['Qb', 'qb·Ab', f"{fmt(pl['Qb'],2)} tf"],
            ['Qs', 'fs·πD·Lb', f"{fmt(pl['Qs'],2)} tf"],
            ['Qa', f"Qb/{pl['FSb']} + Qs/{pl['FSs']}", f"{fmt(pl['Qa'],2)} tf"],
            ['承載檢核', f"Qa > P", f"{fmt(pl['Qa'],2)} > {pl['P']} tf … {ok_ng(pl['ok'])}"],
        ])
        doc.add_page_break()

    # ----- 九、結論與建議 -----
    add_heading(doc, '九、結論與建議', 1)
    if overall:
        add_para(doc, '一、本案覆工板系統各構件（覆工版面、覆工小梁、覆工大梁）之彎矩、剪力、撓度檢核均符合規範規定。')
        add_para(doc, '二、共構柱、H 型鋼貫入 PC 握裹力、及樁基承載力檢核亦均通過。')
        add_para(doc, '三、建議：施工時應確實按本計算書所示構件規格、斷面配置及支撐間距施作，並禁止超載堆積。')
    else:
        add_para(doc, '一、經檢核結果，部分項目未能通過容許值：')
        for nm, lbl in [('deck','覆工版面'), ('stringer','覆工小梁'), ('girder','覆工大梁')]:
            v = R.get(nm)
            if v and not (v['ok_fb'] and v['ok_fv'] and v['ok_d']):
                ng_items = []
                if not v['ok_fb']: ng_items.append('彎矩應力')
                if not v['ok_fv']: ng_items.append('剪應力')
                if not v['ok_d']:  ng_items.append('撓度')
                add_para(doc, f"  ・{lbl}：{'、'.join(ng_items)} 超過容許值", indent=0.5)
        add_para(doc, '二、建議：NG 項目需放大斷面、增加並排支數、或於固定支座下加鋪補強小梁。')

    add_para(doc, '')
    add_para(doc, f"本計算書係依本案設計載重（HS 20-44 卡車、PC400 履帶車、指定噸位吊車）為依據，若實際施工載重或機具規格變更，應另行檢核。", size=9)

    # 獨立報告才保留簽章欄；附件模式由主文承接簽認。
    if document_mode == 'standalone':
        add_para(doc, '')
        add_para(doc, '')
        add_table(doc, ['設　計','覆　核','核　定'], [['','',''],[' ',' ',' ']], col_widths=[5.6, 5.6, 5.6])
    add_para(doc, '')
    status_parts = ['文件狀態：正式附件' if approved else '文件狀態：內部審閱']
    if approved and approved_at:
        status_parts.append(f'核可時間：{approved_at}')
    status_parts.append(f'計算指紋：{calculation_fingerprint}')
    add_para(doc, '｜'.join(status_parts), size=9, bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_para(doc, f"報告產出時間：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", size=9, align=WD_ALIGN_PARAGRAPH.RIGHT)

    # ----- 存檔 -----
    if out_path is None:
        base = os.path.dirname(os.path.abspath(json_path))
        safe_name = (p.get('name') or 'CASE').replace('/', '_').replace('\\', '_')[:30]
        out_path = os.path.join(base, f"覆工板計算書_{safe_name}_{p.get('date') or datetime.now().strftime('%Y%m%d')}.docx")
    doc.save(out_path)
    return out_path


if __name__ == '__main__':
    if len(sys.argv) < 2:
        # 自動找同目錄內最新 JSON
        here = os.path.dirname(os.path.abspath(__file__))
        parent = os.path.abspath(os.path.join(here, '..'))
        jsons = [os.path.join(parent, f) for f in os.listdir(parent) if f.endswith('.json') and f.startswith('覆工板_')]
        if not jsons:
            print('找不到 JSON 輸入檔，請將 HTML 工具匯出的 .json 放到覆工板資料夾。')
            sys.exit(1)
        jsons.sort(key=os.path.getmtime, reverse=True)
        inp = jsons[0]
        print(f'使用最新 JSON：{os.path.basename(inp)}')
    else:
        inp = sys.argv[1]
    out = sys.argv[2] if len(sys.argv) >= 3 else None
    result = generate(inp, out)
    print(f'\n✓ 計算書已產出：\n  {result}\n')
    if not os.environ.get('COVER_SLAB_NO_OPEN'):
        try:
            os.startfile(result)
        except Exception:
            pass
